"""Compose evidence-grounded annual-audit attachments from template blueprints."""

from __future__ import annotations

import json
import hashlib
import mimetypes
import re
import tempfile
import zipfile
from collections import defaultdict
from concurrent.futures import ThreadPoolExecutor, as_completed
from datetime import date, datetime
from decimal import Decimal, InvalidOperation
from difflib import SequenceMatcher
from io import BytesIO
from pathlib import Path
from typing import Any

from ai_hunter.app.services.minio_service import get_minio_service
from ai_hunter.app.settings import Settings, get_settings
from ai_hunter.platform_core import ArtifactRef

from .storage import mysql_connection
from .attachment_fact_policy import filter_unavailable_statement_sources
from .attachment_quality_service import (
    evaluate_attachment_quality,
    preflight_attachment_quality_blockers,
)
from .generic_template_repository import get_active_template_catalog, template_version_ref


_TEMPLATE_PLAN_VERSION = "complete-attachment-composition-plan-v7"
# All attachment types and future business templates enter this contract before
# their format-specific writer runs.  ``template_type`` is intentionally not a
# part of the placement algorithm: it only selects business release rules and
# user-facing names.  Values are always resolved through the same periodised,
# source-prioritised fact registry and are written only into slots discovered
# from the uploaded template structure.
_TEMPLATE_RENDER_CONTRACT_VERSION = "template-structure-fact-registry-v2"
_CASE_MATERIAL_MAX_LABELS = 2500
_MATERIAL_NOISE_LABELS = frozenset(
    {
        "增加",
        "减少",
        "调整",
        "期初",
        "期末",
        "本期",
        "上期",
        "本年",
        "上年",
        "未审数",
        "审定数",
        "金额",
        "余额",
        "借方",
        "贷方",
        "直接",
        "间接",
        "合计数",
    }
)
_CASE_REFERENCE_FILE_MARKERS = (
    "case-notes-reference",
    "案例回放",
    "案例产出",
    "标准无保留意见",
)

# These are accounting-presentation aliases, not template-specific rules.  A
# template can call the same financial-statement fact by a different standard
# label (for example, the cash-flow statement calls the line "现金的期末余额"
# while the statement of financial position calls it "货币资金").  The source
# of every value is retained in the fact registry built for the engagement.
_STANDARD_FACT_ALIASES: dict[str, tuple[str, int]] = {
    "实收资本或股本": ("股本", 1),
    "其他应收款项": ("其他应收款", 1),
    "可随时用于支付的银行存款": ("银行存款", 1),
    "可随时用于支付的其他货币资金": ("其他货币资金", 1),
    "现金的期末余额": ("货币资金", 1),
    "现金的期初余额": ("货币资金", 1),
    "现金期末余额": ("货币资金", 1),
    "现金期初余额": ("货币资金", 1),
    "三期末现金及现金等价物余额": ("货币资金", 1),
    "期末现金及现金等价物余额": ("货币资金", 1),
    "期初现金及现金等价物余额": ("货币资金", 1),
    "固定资产折旧油气资产折耗生产性生物资产折旧": ("固定资产折旧", 1),
    "投资损失收益以号填列": ("投资收益", -1),
}

# A financial-statement presentation may use the same short label as a gross
# asset schedule.  Prefer the statement's net carrying amount where that more
# specific fact exists; retain the gross amount separately as ``固定资产原值``.
_PREFERRED_PRESENTATION_FACTS = {
    "固定资产": "固定资产净值",
}

_COMPOSITE_TABLE_FACTS: dict[str, dict[str, str]] = {
    "主营业务": {"收入": "主营业务收入", "成本": "主营业务成本"},
    "其他业务": {"收入": "其他业务收入", "成本": "其他业务成本"},
    "合计": {"收入": "营业收入", "成本": "营业成本"},
}

_TEMPLATE_FIELD_SOURCE_LABELS = {
    "entity_name": "审计项目主数据：被审计单位名称",
    "fiscal_year": "审计项目主数据：财务年度",
    "period_start": "审计项目主数据：审计期间起始日",
    "period_end": "审计项目主数据：资产负债表日",
    "report_number": "审计项目主数据/项目名称中的审计报告文号",
    "audit_result": "年度审计事实快照：审计程序、发现、复核门禁",
    "financial_statement_line_items": "主底稿/科目余额表/经审计财务报表工作表",
    "notes_disclosure": "主底稿 296 个工作表及财务报表附注相关证据",
}


# The case reference produces three core deliverables: the audit report,
# audited financial statements, and notes.  A management letter is a separate
# optional deliverable and must never make the core package partial merely
# because its template is not configured.
DEFAULT_ATTACHMENT_TYPES = (
    "annual_report",
    "financial_statements",
    "notes",
)
OPTIONAL_ATTACHMENT_TYPES = ("management_letter", "confirmations")

_TOKEN_RE = re.compile(r"\{\{\s*([\w.-]+)\s*\}\}|\[\[\s*([\w.-]+)\s*\]\]|\$\{\s*([\w.-]+)\s*\}")
_SINGLE_BRACKET_TOKEN_RE = re.compile(r"\[([\w.-]+)\]")
_BRACKET_PLACEHOLDER_RE = re.compile(r"[\[【](?:被审计单位名称|项目合伙人姓名|注册会计师签名|日期|日期|XXXX|XX)[\]】]")
_INVALID_FILENAME_CHARS_RE = re.compile(r'[<>:"/\\|?*\x00-\x1f]')
_CLIENT_DELIVERY_TRACE_PATTERNS = (
    re.compile(r"(?i)\b(?:evidence(?:[_ -]?id|[_ -]?entries)?|source[_ -]?(?:chunk|file|page|locator)|chunk:[\w.-]+|claim:[\w.-]+)\b"),
    re.compile(r"(?:项目证据|结构化事实|证据(?:编号|id|ID|定位|来源|条目)|来源(?:文件|工作表|页码|行号|单元格)|内部程序编号|审计程序编号)"),
    re.compile(r"(?:工作底稿(?:编号)?\s*[:：]?\s*[A-Za-z]{1,5}\d{1,4}(?:-\d+)?|\b(?:[A-Z]{1,5}\d{1,4}(?:-\d+)?|TBA\d{2,})\b)"),
    re.compile(r"(?i)\b(?:fact|statement|audit|material|engagement|finding):[a-z0-9_.-]+\b"),
)

_ATTACHMENT_NAME_LABELS = {
    "annual_report": "财务报表审计报告",
    "financial_statements": "财务报表",
    "notes": "财务报表附注",
    "management_letter": "管理建议书",
    "confirmations": "函证",
    "audit_workpaper": "审计工作底稿",
}

_CONTENT_TYPES_BY_SUFFIX = {
    ".doc": "application/msword",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".xls": "application/vnd.ms-excel",
    ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    ".xlsm": "application/vnd.ms-excel.sheet.macroEnabled.12",
    ".pdf": "application/pdf",
    ".csv": "text/csv",
    ".md": "text/markdown",
    ".markdown": "text/markdown",
    ".txt": "text/plain",
}

_RESULT_TOKEN_KEYS = frozenset(
    {
        "report_text",
        "audit_result",
        "审计结果",
        "findings_summary",
        "审计发现摘要",
    }
)

_GENERIC_NOTE_INSTRUCTION_MARKERS = (
    "披露要求",
    "提示：",
    "本参考格式",
    "具体描述指定的情况",
    "公司应评价自报告期末",
    "公司应制定与实际生产经营特点",
    "根据实际情况从性质和金额",
    "企业需根据其自身实际情况",
    "上表应完整列示",
    "各类固定资产折旧方法、折旧年限、残值率和年折旧率如下",
)

_BLOCKED_RENDER_STATUSES = frozenset(
    {
        "copied_legacy_or_unknown_format",
        "copied_legacy_unmodified",
        "requires_office_automation",
        "copied_no_matching_placeholders",
        "copied_no_matching_form_fields",
        "copied_pdf",
        "metadata_filled_no_result_mapping",
        "no_template_mapping",
    }
)


def _json_default(value: Any) -> Any:
    if isinstance(value, (date, datetime)):
        return value.isoformat()
    return str(value)


def _exact_decimal_text(value: Any) -> str:
    """Preserve database DECIMAL values for deterministic accounting checks."""

    if value in (None, ""):
        return "0"
    try:
        return format(Decimal(str(value).replace(",", "").strip()), "f")
    except (InvalidOperation, TypeError, ValueError):
        return str(value)


def _load_trial_balance(engagement_id: int, settings: Settings) -> dict[str, Any]:
    """从数据库加载科目余额表，作为财务报表重建的结构化事实来源。

    返回结构：
    {
        "accounts": [ {account_code, account_name, opening_debit, opening_credit,
                        period_debit, period_credit, closing_debit, closing_credit}, ... ],
        "by_name": { "货币资金": {closing_debit, closing_credit, opening_debit, ...}, ... },
    }
    """

    try:
        with mysql_connection(settings) as connection:
            with connection.cursor() as cursor:
                cursor.execute(
                    """
                    SELECT balances.account_code, balances.account_name,
                           balances.opening_debit, balances.opening_credit,
                           balances.period_debit, balances.period_credit,
                           balances.closing_debit, balances.closing_credit,
                           batches.id AS import_batch_id,
                           batches.source_ref, batches.source_sha256
                    FROM annual_account_balance AS balances
                    JOIN annual_import_batch AS batches
                      ON batches.id = balances.import_batch_id
                    JOIN (
                      SELECT latest.source_sha256,
                             SUBSTRING_INDEX(latest.source_ref, '#', 1) AS source_object_ref
                      FROM annual_import_batch AS latest
                      WHERE latest.engagement_id = %s
                        AND latest.source_type = 'account_balance'
                        AND latest.status = 'completed'
                      ORDER BY latest.completed_at DESC, latest.id DESC
                      LIMIT 1
                    ) AS active_revision
                      ON active_revision.source_sha256 = batches.source_sha256
                     AND active_revision.source_object_ref =
                         SUBSTRING_INDEX(batches.source_ref, '#', 1)
                    WHERE balances.engagement_id = %s
                      AND batches.status = 'completed'
                    ORDER BY balances.account_code
                    """,
                    (engagement_id, engagement_id),
                )
                rows = cursor.fetchall()
    except Exception:
        return {
            "accounts": [],
            "exact_accounts": [],
            "by_name": {},
            "authoritative_batch": {},
        }

    accounts: list[dict[str, Any]] = []
    exact_accounts: list[dict[str, Any]] = []
    by_name: dict[str, dict[str, Any]] = {}
    for row in rows:
        exact_accounts.append(
            {
                "account_code": str(row.get("account_code") or ""),
                "account_name": str(row.get("account_name") or "").strip(),
                **{
                    key: _exact_decimal_text(row.get(key))
                    for key in (
                        "opening_debit",
                        "opening_credit",
                        "period_debit",
                        "period_credit",
                        "closing_debit",
                        "closing_credit",
                    )
                },
            }
        )
        item = {
            "account_code": str(row.get("account_code") or ""),
            "account_name": str(row.get("account_name") or "").strip(),
            "opening_debit": float(row.get("opening_debit") or 0),
            "opening_credit": float(row.get("opening_credit") or 0),
            "period_debit": float(row.get("period_debit") or 0),
            "period_credit": float(row.get("period_credit") or 0),
            "closing_debit": float(row.get("closing_debit") or 0),
            "closing_credit": float(row.get("closing_credit") or 0),
        }
        accounts.append(item)
        name = item["account_name"]
        if name:
            # 如果有多个同名科目，累加余额
            if name in by_name:
                for key in ("opening_debit", "opening_credit", "period_debit", "period_credit", "closing_debit", "closing_credit"):
                    by_name[name][key] += item[key]
            else:
                by_name[name] = dict(item)

    first = rows[0] if rows else {}
    batch_ids = sorted(
        {int(row.get("import_batch_id") or 0) for row in rows}
        - {0}
    )
    return {
        "accounts": accounts,
        "exact_accounts": exact_accounts,
        "by_name": by_name,
        "authoritative_batch": {
            "import_batch_id": batch_ids[-1] if batch_ids else 0,
            "import_batch_ids": batch_ids,
            "source_ref": str(first.get("source_ref") or ""),
            "source_sha256": str(first.get("source_sha256") or ""),
        },
    }


def _normalize_material_label(value: Any) -> str:
    """Normalize a workbook label for deterministic template matching."""

    text = str(value or "").strip()
    text = re.sub(
        r"[\s　\t\r\n:：()（）\[\]【】\"'“”‘’\-—–+，,。．.；;、/\\]",
        "",
        text,
    )
    text = text.replace("合计", "").replace("小计", "")
    return text


def _numeric_material_values(row: list[Any]) -> list[float]:
    values: list[float] = []
    for value in row:
        if isinstance(value, bool) or value in (None, ""):
            continue
        if isinstance(value, (int, float)):
            values.append(float(value))
            continue
        text = str(value).strip().replace(",", "")
        if not text or text.startswith("="):
            continue
        try:
            values.append(float(text))
        except (TypeError, ValueError):
            continue
    return values


def _extract_case_entity_profile(sheets: list[Any]) -> dict[str, Any]:
    """Extract only explicit entity facts from the supplied workpaper pack.

    This is intentionally a small semantic profile, not an LLM guess.  Each
    value is taken from a labelled row in the workbook and is retained with a
    sheet/row locator so it can be used to fill the generic note template and
    audited later.
    """

    profile: dict[str, Any] = {}

    def put(key: str, value: Any, sheet_name: str, row_number: int) -> None:
        text = str(value or "").strip()
        if key == "entity_nature":
            # The client-profile sheet stores this field in a fixed-width
            # cell and the trailing zero is a formatting residue, not part of
            # the legal entity type.
            text = re.sub(r"0$", "", text).strip()
        if text and key not in profile:
            profile[key] = text
            profile[f"{key}_evidence"] = {
                "sheet_name": sheet_name,
                "row_number": row_number,
            }

    for sheet in sheets:
        name = str(getattr(sheet, "name", "") or "")
        rows = list(getattr(sheet, "rows", []) or [])
        for row_number, raw_row in enumerate(rows, start=1):
            row = list(raw_row)
            texts = [str(value).strip() for value in row if value not in (None, "")]
            row_text = " | ".join(texts)
            if not row_text:
                continue
            for value in texts:
                if value.startswith("被审计单位：") and "entity_masked" not in profile:
                    put("entity_masked", value.split("：", 1)[1], name, row_number)
                if value.startswith("财务报表截止日/期间：") and "period_label" not in profile:
                    put("period_label", value.split("：", 1)[1], name, row_number)

            if name == "客户基本情况":
                if row_number == 5 and len(row) > 2:
                    put("registered_address", row[2], name, row_number)
                elif row_number == 7 and len(row) > 2:
                    put("business_scope", row[2], name, row_number)
                elif row_number == 9 and len(row) > 11:
                    put("entity_nature", row[11], name, row_number)
                elif row_number == 11 and len(row) > 11:
                    put("registered_capital", row[11], name, row_number)
                elif row_number == 17:
                    if len(row) > 0:
                        put("shareholder", row[0], name, row_number)
                    if len(row) > 12:
                        put("paid_in_capital", row[12], name, row_number)

            if "主要经营活动" in row_text and "main_activity" not in profile:
                for index, value in enumerate(texts):
                    if "主要经营活动" in value and index + 1 < len(texts):
                        put("main_activity", texts[index + 1], name, row_number)
                        break
            if "主要为：" in row_text and "main_activity" not in profile:
                put("main_activity", row_text.split("主要为：", 1)[1], name, row_number)

    return profile


def _extract_case_statement_values(sheets: list[Any]) -> dict[str, dict[str, Any]]:
    """Extract canonical statement/detail values used by note tables."""

    values: dict[str, dict[str, Any]] = {}

    def put(key: str, current: Any = None, opening: Any = None, source: str = "", row_number: int = 0) -> None:
        if not key or all(
            isinstance(item, bool) or not isinstance(item, (int, float))
            for item in (current, opening)
            if item is not None
        ):
            return
        item = values.setdefault(
            key,
            {
                "current": None,
                "opening": None,
                "source_sheet": source,
                "source_row": row_number,
            },
        )
        if isinstance(current, (int, float)) and not isinstance(current, bool):
            item["current"] = float(current)
        if isinstance(opening, (int, float)) and not isinstance(opening, bool):
            item["opening"] = float(opening)
        if source:
            item["source_sheet"] = source
            item["source_row"] = row_number

    def statement_key(label: str) -> str:
        clean = _normalize_material_label(label)
        if not clean:
            return ""
        if "主营业务收入" in clean:
            return "主营业务收入"
        if "主营业务成本" in clean:
            return "主营业务成本"
        if "营业收入" in clean and "主营业务" not in clean:
            return "营业收入"
        if "营业成本" in clean and "主营业务" not in clean:
            return "营业成本"
        if "营业利润" in clean:
            return "营业利润"
        if "投资收益" in clean and "其中" not in clean:
            return "投资收益"
        if "所得税费用" in clean:
            return "所得税费用"
        if clean.startswith("五、净利润") or ("净利润" in clean and not any(term in clean for term in ("被合并方", "归属于", "少数"))):
            return "净利润"
        for token, key in (
            ("营业税金及附加", "税金及附加"),
            ("税金及附加", "税金及附加"),
            ("销售费用", "销售费用"),
            ("管理费用", "管理费用"),
            ("财务费用", "财务费用"),
            ("其他收益", "其他收益"),
            ("营业外收入", "营业外收入"),
            ("营业外支出", "营业外支出"),
            ("利润总额", "利润总额"),
        ):
            if token in clean:
                return key
        return ""

    def number(value: Any) -> float | None:
        """Read a literal workbook amount; formulas are handled explicitly."""

        if isinstance(value, (int, float)) and not isinstance(value, bool):
            return float(value)
        text = str(value or "").strip().replace(",", "")
        if not text or text.startswith("="):
            return None
        try:
            return float(text)
        except (TypeError, ValueError):
            return None

    def subject_from_sheet_name(name: str) -> str:
        """Extract the audited subject from a workpaper name, not a template."""

        subject = re.sub(r"^[A-Za-z]+\d+(?:[-－]\d+)*", "", str(name or ""))
        subject = re.sub(
            r"(?:审计程序表?|审定表|审核表|核实表|检查表|审查表|明细表)$",
            "",
            subject,
        )
        return subject.strip(" -－_ ").replace("帐", "账")

    def detail_label(value: Any) -> str:
        label = str(value or "").strip().replace("帐", "账")
        label = re.sub(r"^\d+[\.．、]\s*", "", label)
        return label

    def metric_from_schedule_row(value: Any) -> tuple[str, str] | None:
        """Map standard workpaper movement wording to a fact dimension."""

        text = _normalize_material_label(value)
        if "期初" in text or "年初" in text:
            return "opening", ""
        if "本期增加" in text or "本年增加" in text:
            return "current", "本期增加"
        if "摊销" in text:
            return "current", "摊销"
        if "减少" in text:
            return "current", "其他减少"
        if "期末" in text or "审定数" in text:
            return "current", ""
        return None

    def aging_bucket(value: Any) -> str:
        """Normalise a standard aging-column label for note-table matching."""

        text = _normalize_template_label(value)
        text = text.replace("以下", "以内").replace("以上", "以上")
        if not text or "年" not in text:
            return ""
        if re.fullmatch(r"(?:\d+年以内|\d+[至-]\d+年|\d+年以上)", text):
            return text
        return ""

    def put_detail(
        subject: str,
        category: str,
        *,
        current: Any = None,
        opening: Any = None,
        metric: str = "",
        source: str,
        row_number: int,
    ) -> None:
        clean_subject = _normalize_material_label(subject)
        clean_category = detail_label(category)
        if not clean_subject or not clean_category or clean_category in {"合计", "项目", "索引号"}:
            return
        key = f"{clean_subject}|{clean_category}"
        if metric:
            key = f"{key}|{metric}"
        put(key, current=current, opening=opening, source=source, row_number=row_number)

    # The four statement sheets contain the audited/approved amount in the
    # final column.  They are the canonical source for note tables; the
    # normalized trial balance does not retain the workpaper's opening/final
    # presentation adjustments.
    for sheet in sheets:
        name = str(getattr(sheet, "name", "") or "")
        if name not in {"资产表(年末数)", "资产表(年初数)", "负债及权益表(年末数)", "负债及权益表(年初数)"}:
            continue
        is_opening = name.endswith("年初数)")
        for row_number, raw_row in enumerate(getattr(sheet, "rows", []) or [], start=1):
            row = list(raw_row)
            label = str(row[0] or "").strip() if row else ""
            if len(row) <= 8 or not label or not isinstance(row[8], (int, float)):
                continue
            clean_label = _normalize_material_label(label)
            if len(clean_label) < 2 or clean_label in _MATERIAL_NOISE_LABELS:
                continue
            if is_opening:
                put(clean_label, opening=row[8], source=name, row_number=row_number)
            else:
                put(clean_label, current=row[8], source=name, row_number=row_number)

    for sheet in sheets:
        name = str(getattr(sheet, "name", "") or "")
        rows = list(getattr(sheet, "rows", []) or [])
        if name in {"利润表(本年数)", "利润表(上年数)"}:
            is_opening = name == "利润表(上年数)"
            for row_number, raw_row in enumerate(rows, start=1):
                row = list(raw_row)
                label = str(row[0] or "").strip() if row else ""
                if len(row) <= 6 or not label:
                    continue
                current = row[6]
                if not isinstance(current, (int, float)):
                    continue
                key = statement_key(label)
                if not key:
                    continue
                if is_opening:
                    put(key, opening=current, source=name, row_number=row_number)
                else:
                    put(key, current=current, source=name, row_number=row_number)
        elif name == "单独现金流量表":
            for row_number, raw_row in enumerate(rows, start=1):
                row = list(raw_row)
                label = str(row[0] or "").strip() if row else ""
                if len(row) > 2 and label and isinstance(row[2], (int, float)):
                    put(label, row[2], row[3] if len(row) > 3 else None, source=name, row_number=row_number)
        elif name == "C1-2货币资金审定":
            bank_detail = False
            for row_number, raw_row in enumerate(rows, start=1):
                row = list(raw_row)
                row_text = " ".join(str(value).strip() for value in row if value not in (None, ""))
                if "开户银行" in row_text:
                    bank_detail = True
                    continue
                label = str(row[0] or "").strip() if row else ""
                if len(row) <= 12:
                    continue
                current = row[12]
                opening = row[4] if len(row) > 4 else None
                if label == "货币资金合计":
                    put("货币资金", current=current, opening=opening, source=name, row_number=row_number)
                elif label == "现金":
                    put("库存现金", current=current, opening=opening, source=name, row_number=row_number)
                elif bank_detail and label == "合计":
                    put("银行存款", current=current, opening=opening, source=name, row_number=row_number)
        elif name == "C21-1-2固定资产、累计折旧审定表":
            for row_number, raw_row in enumerate(rows, start=1):
                row = list(raw_row)
                label = next((str(value).strip() for value in row if value not in (None, "") and not isinstance(value, (int, float))), "")
                if len(row) > 7 and isinstance(row[7], (int, float)) and isinstance(row[2], (int, float)):
                    if "固定资产原值合计" in label:
                        put("固定资产原值", current=row[7], opening=row[2], source=name, row_number=row_number)
                    elif "累计折旧合计" in label:
                        put("累计折旧", current=row[7], opening=row[2], source=name, row_number=row_number)
                    elif "固定资产净值合计" in label:
                        put("固定资产净值", current=row[7], opening=row[2], source=name, row_number=row_number)
                    elif len(row) > 1 and str(row[1] or "").strip():
                        detail = str(row[1]).strip()
                        put(detail, current=row[7], opening=row[2], source=name, row_number=row_number)
                if "累计折旧合计" in label:
                    numeric = [value for value in row if isinstance(value, (int, float))]
                    if len(numeric) >= 2:
                        put("固定资产折旧", numeric[1], source=name, row_number=row_number)
        elif name == "C29-2长期待摊费用审定表":
            for row_number, raw_row in enumerate(rows, start=1):
                row = list(raw_row)
                label = next((str(value).strip() for value in row if value not in (None, "") and not isinstance(value, (int, float))), "")
                numeric = [value for value in row if isinstance(value, (int, float))]
                if "期初余额" in label and numeric:
                    put("长期待摊费用", opening=numeric[0], source=name, row_number=row_number)
                elif "本期增加额" in label and numeric:
                    put("长期待摊费用本期增加", current=numeric[0], source=name, row_number=row_number)
                elif "本期已摊销额" in label and numeric:
                    put("长期待摊费用摊销", current=numeric[0], source=name, row_number=row_number)
                elif "期末未审数" in label and numeric:
                    put("长期待摊费用", current=numeric[0], source=name, row_number=row_number)

    # Generic detailed-workpaper adapter.  Audit schedules commonly present
    # detail in one of two orientations: a vertical account list with
    # opening/debit/credit/audited columns, or a transposed movement schedule
    # whose columns are asset categories and rows are opening/addition/
    # amortisation/closing.  Read both layouts from the source workbook and
    # retain a three-dimensional fact key (subject | category | movement).
    # This is independent of any target DOCX template; the template merely
    # asks for one of those facts by its displayed row/column labels.
    for sheet in sheets:
        name = str(getattr(sheet, "name", "") or "")
        subject = subject_from_sheet_name(name)
        rows = [list(row) for row in (getattr(sheet, "rows", []) or [])]
        if not subject or not rows:
            continue

        for header_row, header in enumerate(rows[:20]):
            header_text = [_normalize_material_label(value) for value in header]
            project_columns = [
                index
                for index, value in enumerate(header_text)
                if value in {"项目", "往来单位", "客户名称", "客户", "单位名称", "名称"}
            ]
            if not project_columns:
                continue
            project_column = project_columns[0]

            # Vertical audit schedules (for example trading financial assets)
            # keep the category in a "项目" column.  When an audited closing
            # formula has no cached value, calculate the standard roll-forward
            # only from the same row's literal opening/debit/credit/adjustment
            # evidence; never borrow another row's amount.
            columns: dict[str, int] = {}
            subheader = [
                _normalize_material_label(value)
                for value in (rows[header_row + 1] if header_row + 1 < len(rows) else [])
            ]
            for column, raw_text in enumerate(header_text):
                # A common workpaper header spans two rows: “期初余额” on
                # the first row and “原币/记账本位币” below.  Build a local
                # composite header so the RMB column is selected rather than
                # the adjacent original-currency column.
                text = raw_text
                detail = subheader[column] if column < len(subheader) else ""
                if not text and column > 0:
                    text = header_text[column - 1]
                combined = f"{text}{detail}"
                if ("期初" in combined or "年初" in combined) and "原币" in detail:
                    continue
                if ("期末" in combined or "审定" in combined) and "原币" in detail:
                    continue
                if "期初" in combined or "年初" in combined:
                    columns.setdefault("opening", column)
                elif "本期借方" in combined or "本年借方" in combined:
                    columns.setdefault("debit", column)
                elif "本期贷方" in combined or "本年贷方" in combined:
                    columns.setdefault("credit", column)
                elif "审定数" in combined:
                    columns.setdefault("current", column)
                elif "期末" in combined:
                    columns.setdefault("unreviewed", column)
                elif "调整数" in combined:
                    columns.setdefault("adjustment", column)
            if "opening" in columns and any(key in columns for key in ("current", "unreviewed", "debit", "credit")):
                for offset, row in enumerate(rows[header_row + 1:header_row + 150], start=header_row + 2):
                    if project_column >= len(row):
                        continue
                    category = detail_label(row[project_column])
                    if not category and row:
                        first_label = detail_label(row[0])
                        if first_label == "合计":
                            category = first_label
                    row_text = " ".join(str(value or "") for value in row)
                    if any(marker in row_text for marker in ("审计说明", "审计结论", "调整分录")):
                        break
                    if not category or category in {"项目", "简称", "名称"}:
                        continue
                    opening = number(row[columns["opening"]]) if columns["opening"] < len(row) else None
                    current = number(row[columns["current"]]) if "current" in columns and columns["current"] < len(row) else None
                    if current is None and "unreviewed" in columns and columns["unreviewed"] < len(row):
                        current = number(row[columns["unreviewed"]])
                    debit = number(row[columns["debit"]]) if "debit" in columns and columns["debit"] < len(row) else None
                    credit = number(row[columns["credit"]]) if "credit" in columns and columns["credit"] < len(row) else None
                    adjustment = number(row[columns["adjustment"]]) if "adjustment" in columns and columns["adjustment"] < len(row) else None
                    if current is None and any(value is not None for value in (opening, debit, credit, adjustment)):
                        current = (opening or 0.0) + (debit or 0.0) - (credit or 0.0) + (adjustment or 0.0)
                    if opening is not None or current is not None:
                        if category == "合计":
                            put(subject, current=current, opening=opening, source=name, row_number=offset)
                            # Some receivable workpapers store aging as a
                            # second header row above literal amount columns.
                            # Keep each populated aging bucket as a detail
                            # fact, so a blank repeat-row in any note template
                            # can be filled without a table-specific map.
                            for column, sublabel in enumerate(subheader):
                                bucket = aging_bucket(sublabel)
                                amount = number(row[column]) if column < len(row) else None
                                if bucket and amount is not None:
                                    put_detail(subject, bucket, current=amount, source=name, row_number=offset)
                        else:
                            put_detail(subject, category, current=current, opening=opening, source=name, row_number=offset)

            # Transposed movement schedules have category headings beside the
            # "项目" cell and standard movement labels below it.  Only create
            # a detail fact when the intersecting source cell is a literal
            # amount; an unavailable formula is left unavailable rather than
            # fabricated as zero.
            movement_rows: list[tuple[int, str, str]] = []
            for offset, row in enumerate(rows[header_row + 1:header_row + 80], start=header_row + 2):
                if project_column >= len(row):
                    continue
                movement = metric_from_schedule_row(row[project_column])
                if movement:
                    movement_rows.append((offset, *movement))
            if len(movement_rows) < 2:
                continue
            for column, category in enumerate(header):
                category_text = detail_label(category)
                normalized_category = _normalize_material_label(category_text)
                if column == project_column or not normalized_category or normalized_category in {"合计", "索引号", "项目"}:
                    continue
                if any(token in normalized_category for token in ("期初", "期末", "本期", "审定", "调整", "借方", "贷方")):
                    continue
                for row_number, period, metric in movement_rows:
                    row = rows[row_number - 1]
                    value = number(row[column]) if column < len(row) else None
                    if value is None:
                        continue
                    if metric:
                        put_detail(subject, category_text, current=value, metric=metric, source=name, row_number=row_number)
                    else:
                        put_detail(
                            subject,
                            category_text,
                            current=value if period == "current" else None,
                            opening=value if period == "opening" else None,
                            source=name,
                            row_number=row_number,
                        )
    return values


def _load_case_material_index(
    engagement_id: int,
    *,
    settings: Settings,
) -> dict[str, Any]:
    """Read the complete case workpaper again before attachment rendering.

    The case workbook is intentionally kept separate from the normalized source
    fact tables.  It is nevertheless the authoritative replay package for this
    supplied case, so template generation must read every worksheet and retain
    a compact, auditable label/value index instead of silently relying on only
    F1-2/C5-2/C1-2.
    """

    from .import_service import read_tabular_sheets
    from .workpaper_case import CASE_WORKPAPER_SOURCE_TYPE, summarize_workpaper_sheets

    with mysql_connection(settings) as connection:
        with connection.cursor() as cursor:
            cursor.execute(
                """
                SELECT source_ref, source_sha256, row_count, metadata_json
                FROM annual_import_batch
                WHERE engagement_id = %s AND source_type = %s AND status = 'completed'
                ORDER BY id DESC LIMIT 1
                """,
                (engagement_id, CASE_WORKPAPER_SOURCE_TYPE),
            )
            row = cursor.fetchone()
    if not row:
        return {
            "status": "missing",
            "source_type": CASE_WORKPAPER_SOURCE_TYPE,
            "message": "当前项目没有已完成的主底稿回放导入批次。",
        }

    source_ref = str(row.get("source_ref") or "")
    source_object_ref = source_ref.split("#", 1)[0]
    metadata = row.get("metadata_json") or {}
    if isinstance(metadata, str):
        try:
            metadata = json.loads(metadata)
        except json.JSONDecodeError:
            metadata = {}
    file_name = str((metadata or {}).get("file_name") or Path(source_object_ref).name or "主底稿.xlsx")
    raw_bytes = get_minio_service().get_object_bytes(source_object_ref)
    sheets = read_tabular_sheets(file_name, raw_bytes)
    if not sheets:
        return {
            "status": "unsupported",
            "source_type": CASE_WORKPAPER_SOURCE_TYPE,
            "file_name": file_name,
            "source_ref": source_ref,
            "source_sha256": str(row.get("source_sha256") or ""),
            "message": "主底稿文件无法按当前支持的表格格式读取。",
        }

    summary = summarize_workpaper_sheets(sheets, file_name=file_name)
    entity_profile = _extract_case_entity_profile(sheets)
    statement_values = _extract_case_statement_values(sheets)
    trial_balance = _load_trial_balance(engagement_id, settings)
    account_names = [
        str(item.get("account_name") or "")
        for item in trial_balance.get("accounts") or []
        if any(float(item.get(key) or 0) != 0 for key in ("opening_debit", "opening_credit", "period_debit", "period_credit", "closing_debit", "closing_credit"))
    ]
    label_entries: dict[str, list[dict[str, Any]]] = defaultdict(list)
    for sheet in sheets:
        for row_number, raw_row in enumerate(sheet.rows, start=1):
            values = list(raw_row)
            nonempty = [value for value in values if value not in (None, "")]
            if not nonempty:
                continue
            numeric = _numeric_material_values(values)
            text_values = [str(value).strip() for value in values if value not in (None, "") and not isinstance(value, (int, float))]
            for label in text_values[:4]:
                normalized = _normalize_material_label(label)
                if len(normalized) < 2 or normalized in {"项目", "类别", "合计", "说明", "备注"}:
                    continue
                if normalized in _MATERIAL_NOISE_LABELS:
                    continue
                if len(label_entries[normalized]) >= 4:
                    continue
                label_entries[normalized].append(
                    {
                        "label": label,
                        "sheet_name": sheet.name,
                        "row_number": row_number,
                        "values": numeric[:8],
                    }
                )

    compact_labels = {
        key: value[:4]
        for key, value in sorted(label_entries.items())
        if value
    }
    active_labels = {
        key: entries
        for key, entries in compact_labels.items()
        if any(
            any(abs(float(number)) > 1e-9 for number in (entry.get("values") or []))
            for entry in entries
            if isinstance(entry, dict)
        )
    }
    if len(compact_labels) > _CASE_MATERIAL_MAX_LABELS:
        compact_labels = dict(list(compact_labels.items())[:_CASE_MATERIAL_MAX_LABELS])
    return {
        "status": "ready",
        "source_type": CASE_WORKPAPER_SOURCE_TYPE,
        "file_name": file_name,
        "source_ref": source_ref,
        "source_sha256": str(row.get("source_sha256") or ""),
        "source_size": len(raw_bytes),
        "sheet_count": len(sheets),
        "row_count": int(row.get("row_count") or summary.get("nonempty_row_count") or 0),
        "nonempty_row_count": int(summary.get("nonempty_row_count") or 0),
        "nonempty_cell_count": int(summary.get("nonempty_cell_count") or 0),
        "formula_error_count": int(summary.get("formula_error_count") or 0),
        "sheet_names": list(summary.get("sheet_names") or []),
        "covered_categories": dict(summary.get("covered_categories") or {}),
        "program_evidence": dict(summary.get("program_evidence") or {}),
        "labels": compact_labels,
        "active_labels": active_labels,
        "trial_balance_account_names": account_names,
        "trial_balance": trial_balance,
        "entity_profile": entity_profile,
        "statement_values": statement_values,
        "read_all_sheets": True,
        "parser": "annual-spreadsheet-local-v1",
    }


def _template_text_and_structure(data: bytes, file_name: str) -> tuple[str, dict[str, Any]]:
    """Inspect a template without changing it or treating it as a result."""

    extension = Path(file_name).suffix.lower()
    if extension == ".docx":
        from docx import Document

        document = Document(BytesIO(data))
        paragraphs = [paragraph.text or "" for paragraph in document.paragraphs]
        table_cells = [
            cell.text or ""
            for table in document.tables
            for row in table.rows
            for cell in row.cells
        ]
        header_footer: list[str] = []
        for section in document.sections:
            header_footer.extend(paragraph.text or "" for paragraph in section.header.paragraphs)
            header_footer.extend(paragraph.text or "" for paragraph in section.footer.paragraphs)
        all_text = "\n".join([*paragraphs, *table_cells, *header_footer])
        table_sections = _docx_table_sections(document)
        table_plans = []
        for table_index, table in enumerate(document.tables):
            header_rows, columns = _table_column_plan(table)
            if not columns:
                continue
            table_plans.append(
                {
                    "locator": f"body.table[{table_index}]",
                    "section": table_sections[table_index] if table_index < len(table_sections) else "",
                    "header_rows": header_rows,
                    "value_columns": {
                        str(column): {"period": period, "metric": metric}
                        for column, (period, metric) in columns.items()
                    },
                    "row_count": len(table.rows),
                }
            )
        return all_text, {
            "extension": extension,
            "paragraph_count": len(document.paragraphs),
            "table_count": len(document.tables),
            "section_count": len(document.sections),
            "paragraph_text_count": sum(bool(text.strip()) for text in paragraphs),
            "table_cell_count": len(table_cells),
            # The parser is template-neutral.  A new DOCX enters the same
            # contract pipeline and contributes its own table locators rather
            # than requiring a file-name branch in the renderer.
            "template_contract": {
                "contract_version": "docx-template-contract-v1",
                "explicit_token_count": len(list(_TOKEN_RE.finditer(all_text))),
                "generic_marker_count": len(re.findall(r"一般企业模板|一般企业|XXX|20XX", all_text)),
                "data_table_count": len(table_plans),
                "data_tables": table_plans,
            },
        }
    if extension in {".xlsx", ".xlsm"}:
        from openpyxl import load_workbook

        workbook = load_workbook(BytesIO(data), read_only=True, data_only=False)
        values = [
            str(cell.value or "")
            for sheet in workbook.worksheets
            for row in sheet.iter_rows()
            for cell in row
            if cell.value not in (None, "")
        ]
        table_plans = _xlsx_workbook_table_plans(workbook)
        return "\n".join(values), {
            "extension": extension,
            "sheet_count": len(workbook.worksheets),
            "nonempty_cell_count": len(values),
            "template_contract": {
                "contract_version": "xlsx-template-contract-v1",
                "explicit_token_count": len(list(_TOKEN_RE.finditer("\n".join(values)))),
                "data_table_count": len(table_plans),
                "data_tables": [
                    {
                        "locator": f"sheet[{item['sheet']}].header_row[{item['header_row']}]",
                        "header_rows": item["header_rows"],
                        "label_column": item["label_column"],
                        "value_columns": {
                            str(column): {"period": period, "metric": metric}
                            for column, (period, metric) in item["value_columns"].items()
                        },
                    }
                    for item in table_plans
                ],
            },
        }
    return data.decode("utf-8-sig", errors="replace"), {
        "extension": extension,
        "byte_size": len(data),
    }


def _template_is_case_output_reference(file_name: str, text: str, template: dict[str, Any]) -> bool:
    lowered_name = str(file_name or "").lower()
    remark = str(template.get("files", [{}])[0].get("remark") or "").lower() if template.get("files") else ""
    return any(marker.lower() in lowered_name or marker.lower() in remark for marker in _CASE_REFERENCE_FILE_MARKERS) or (
        "北京****有限公司" in text and "会计报表附注" in text and text.count("附注") >= 2
    )


def _material_label_match(material_index: dict[str, Any], label: str) -> dict[str, Any] | None:
    labels = material_index.get("labels") or {}
    normalized = _normalize_material_label(label)
    if not normalized:
        return None
    if normalized in labels and labels[normalized]:
        return labels[normalized][0]
    for key, entries in labels.items():
        if key in normalized or normalized in key:
            return entries[0] if entries else None
    return None


def _template_field_plan(
    *,
    template_type: str,
    file_name: str,
    text: str,
    structure: dict[str, Any],
    template: dict[str, Any],
    context: dict[str, Any],
    material_index: dict[str, Any],
) -> dict[str, Any]:
    token_keys = sorted(
        {
            match.group(1) or match.group(2) or match.group(3)
            for match in _TOKEN_RE.finditer(text)
            if match.group(1) or match.group(2) or match.group(3)
        }
    )
    generic_markers = sorted(
        set(re.findall(r"一般企业模板|一般企业|XXX|20XX|披露要求：|提示：本参考格式", text))
    )
    case_reference = _template_is_case_output_reference(file_name, text, template)
    fields = [
        {
            "field": field,
            "source": source,
            "status": "mapped" if context.get(field) not in (None, "") else "missing",
            "value_preview": str(context.get(field) or "")[:120],
        }
        for field, source in _TEMPLATE_FIELD_SOURCE_LABELS.items()
    ]
    if template_type == "annual_report":
        required_fields = {"entity_name", "fiscal_year", "period_end", "audit_result"}
    elif template_type == "financial_statements":
        required_fields = {"entity_name", "fiscal_year", "period_end", "financial_statement_line_items"}
    elif template_type == "notes":
        required_fields = {"entity_name", "fiscal_year", "period_end", "notes_disclosure"}
    else:
        required_fields = {"entity_name", "fiscal_year"}
    by_field = {item["field"]: item for item in fields}
    for field in required_fields:
        if field in {"financial_statement_line_items", "notes_disclosure"}:
            by_field[field]["status"] = "mapped" if material_index.get("status") == "ready" else "missing"

    matched_labels = 0
    candidate_labels: list[str] = []
    if template_type in {"financial_statements", "notes"}:
        for label in re.findall(r"[\u4e00-\u9fffA-Za-z][\u4e00-\u9fffA-Za-z0-9（）()、，：: ]{1,24}", text):
            normalized = _normalize_material_label(label)
            if len(normalized) < 2 or normalized in {"一般企业模板", "财务报表附注", "会计报表附注"}:
                continue
            if normalized not in candidate_labels:
                candidate_labels.append(normalized)
        for label in candidate_labels[:500]:
            if _material_label_match(material_index, label):
                matched_labels += 1

    blockers: list[str] = []
    quality_issues: list[str] = []
    warnings: list[str] = []
    if material_index.get("status") != "ready":
        quality_issues.append(str(material_index.get("message") or "主底稿尚未读取完整"))
    if case_reference:
        warnings.append("模板含案例业务正文；生成时保留其结构与版式，仅在已声明字段写入本项目数据。")
    missing_required = [item["field"] for item in fields if item["field"] in required_fields and item["status"] == "missing"]
    if missing_required:
        quality_issues.append("缺少必填数据映射：" + "、".join(missing_required))
    if template_type in {"financial_statements", "notes"} and matched_labels == 0:
        quality_issues.append("模板财务项目未命中任何结构化事实")
    if int(material_index.get("formula_error_count") or 0) > 0:
        quality_issues.append(
            f"主底稿存在 {int(material_index.get('formula_error_count') or 0)} 个公式错误"
        )

    return {
        "template_type": template_type,
        "file_name": file_name,
        "template_version": template_version_ref(template),
        "structure": structure,
        "template_contract": structure.get("template_contract") or {},
        "field_schema_declared": bool(template.get("field_schema")),
        "token_keys": token_keys,
        "generic_markers": generic_markers,
        "case_output_reference_detected": case_reference,
        "fields": fields,
        "required_fields": sorted(required_fields),
        "candidate_label_count": len(candidate_labels),
        "matched_material_label_count": matched_labels,
        "composition_mode": "template_in_place_slots_only",
        "material_source": {
            "file_name": material_index.get("file_name"),
            "source_sha256": material_index.get("source_sha256"),
            "sheet_count": material_index.get("sheet_count", 0),
            "nonempty_row_count": material_index.get("nonempty_row_count", 0),
            "read_all_sheets": bool(material_index.get("read_all_sheets")),
        },
        "status": "blocked" if blockers else "ready",
        "blockers": blockers,
        "quality_issues": quality_issues,
        "warnings": warnings,
    }


def _context_identity(snapshot: dict[str, Any], report_text: str) -> dict[str, Any]:
    engagement = snapshot.get("engagement") or {}
    return {
        "entity_name": engagement.get("entity_name") or snapshot.get("entity_name") or "",
        "fiscal_year": engagement.get("fiscal_year") or snapshot.get("fiscal_year") or "",
        "period_start": engagement.get("period_start") or snapshot.get("period_start") or "",
        "period_end": engagement.get("period_end") or snapshot.get("period_end") or "",
        "audit_result": report_text,
    }


def _stable_snapshot_hash(snapshot: dict[str, Any]) -> str:
    volatile = {"evaluated_at", "generated_at", "observed_at", "heartbeat_at", "updated_at"}

    def stable(value: Any) -> Any:
        if isinstance(value, dict):
            return {key: stable(item) for key, item in value.items() if key not in volatile}
        if isinstance(value, list):
            return [stable(item) for item in value]
        return value

    encoded = json.dumps(stable(snapshot), ensure_ascii=False, sort_keys=True, default=_json_default)
    return hashlib.sha256(encoded.encode("utf-8")).hexdigest()


def plan_annual_attachment_package(
    engagement_id: int,
    *,
    requested_types: list[str] | None = None,
    settings: Settings | None = None,
) -> dict[str, Any]:
    """Freeze an existing audit result and map templates without mutating it."""

    resolved = settings or get_settings()
    selected_types = tuple(requested_types or DEFAULT_ATTACHMENT_TYPES)
    latest_report = _load_latest_report(engagement_id, resolved)
    snapshot = dict(latest_report.get("snapshot") or {})
    report_text = str(snapshot.pop("report_text", "") or "").strip()
    if not report_text:
        raise ValueError("当前项目的最新审计结果缺少报告正文，不能生成附件")
    snapshot.update(
        {
            "report_id": int(latest_report.get("id") or 0),
            "report_version": int(latest_report.get("report_version") or 0),
            "report_status": str(latest_report.get("status") or "draft"),
        }
    )
    material_index = _load_case_material_index(engagement_id, settings=resolved)
    try:
        from .attachment_material_service import load_case_material_catalog

        material_catalog = load_case_material_catalog(engagement_id, settings=resolved)
    except Exception as exc:
        material_catalog = {
            "material_context_version": "unavailable",
            "engagement_id": engagement_id,
            "catalog_sha256": "",
            "files": [],
            "warnings": [f"全项目材料目录暂不可用：{str(exc)[:240]}"],
        }
    catalog = get_active_template_catalog(settings=resolved)
    identity = _context_identity(snapshot, report_text)
    engagement_name = str(snapshot.get("engagement_name") or "")
    if not engagement_name:
        from .engagement_repository import get_engagement

        engagement_name = str((get_engagement(engagement_id, settings=resolved) or {}).get("name") or "")
    snapshot["engagement_name"] = engagement_name
    report_number_match = re.search(r"京创会审字\[\d{4}\]第\s*\d+号", engagement_name)
    identity["report_number"] = report_number_match.group(0) if report_number_match else ""
    templates: list[dict[str, Any]] = []
    quality_issues: list[str] = preflight_attachment_quality_blockers(
        snapshot=snapshot,
        material_index=material_index,
        requested_types=selected_types,
    )
    # A client attachment may not be queued merely because the template can
    # be parsed. Release, accounting and source-data defects are delivery
    # blockers, not advisory warnings. The old v6 path kept these in a
    # separate list and incorrectly advertised the plan as ready.
    global_quality_blockers = list(dict.fromkeys(quality_issues))
    blockers: list[str] = list(global_quality_blockers)
    for template_type in selected_types:
        template = catalog.get(template_type) or {}
        files = list(template.get("files") or [])
        if not files:
            blockers.append(f"{template_type} 没有已激活的真实模板文件")
            continue
        for source_file in files:
            source_name = str(source_file.get("file_name") or "template")
            source_extension = Path(source_name).suffix.lower()
            source_bytes = get_minio_service().get_object_bytes(str(source_file.get("storage_ref") or ""))
            text, structure = _template_text_and_structure(source_bytes, source_name)
            context = {**identity, "notes_disclosure": material_index.get("status") == "ready", "financial_statement_line_items": material_index.get("status") == "ready"}
            item = _template_field_plan(
                template_type=template_type,
                file_name=source_name,
                text=text,
                structure=structure,
                template=template,
                context=context,
                material_index=material_index,
            )
            item.update(
                {
                    "template_file_id": source_file.get("id"),
                    "source_template_sha256": hashlib.sha256(source_bytes).hexdigest(),
                    "source_template_size": len(source_bytes),
                    "source_storage_ref": source_file.get("storage_ref"),
                    "source_content_type": source_file.get("content_type"),
                    "source_file_ext": source_file.get("file_ext") or source_extension,
                    "template_snapshot": {
                        "template_code": template.get("template_code"),
                        "template_type": template_type,
                        "version_no": int(template.get("version_no") or 0),
                        "version_label": template_version_ref(template),
                        "content_hash": template.get("content_hash") or "",
                        "content": template.get("content") or {},
                        "field_schema": template.get("field_schema") or {},
                        "template_contract": template.get("template_contract") or {},
                    },
                }
            )
            try:
                from .attachment_blueprint_service import (
                    distill_attachment_template,
                    public_blueprint_summary,
                )
                from .attachment_document_author_service import (
                    authoring_preflight,
                    freeze_attachment_authoring_context,
                )

                blueprint = distill_attachment_template(
                    source_bytes,
                    source_name,
                    template_type,
                    field_schema=template.get("field_schema") or {},
                )
                authoring_gate = authoring_preflight(blueprint, settings=resolved)
                item["blueprint"] = blueprint
                item["blueprint_summary"] = public_blueprint_summary(blueprint)
                item["authoring"] = authoring_gate
                if source_extension == ".docx":
                    slot_contract = (
                        (blueprint.get("structure") or {}).get("slot_contract") or {}
                    )
                    if not (
                        slot_contract.get("declared")
                        and slot_contract.get("slots")
                    ):
                        message = (
                            "DOCX 模板未配置经确认的 slot_contract，"
                            "拒绝猜测正文或表格写入位置"
                        )
                        item.setdefault("blockers", []).append(message)
                        item["status"] = "blocked"
                if authoring_gate.get("required") and authoring_gate.get("ready"):
                    item["frozen_authoring_context"] = freeze_attachment_authoring_context(
                        engagement_id,
                        blueprint=blueprint,
                        snapshot=snapshot,
                        report_text=report_text,
                        material_index=material_index,
                        material_catalog=material_catalog,
                        settings=resolved,
                    )
                elif authoring_gate.get("required"):
                    message = str(authoring_gate.get("reason") or "附件 AI 完整编制不可用")
                    item.setdefault("blockers", []).append(message)
                    item["status"] = "blocked"
            except Exception as exc:
                message = f"模板蓝图/冻结证据构建失败：{str(exc)[:300]}"
                item.setdefault("blockers", []).append(message)
                item["status"] = "blocked"
            templates.append(item)
            blockers.extend(f"{source_name}：{message}" for message in item.get("blockers") or [])
            item_quality_issues = [
                f"{source_name}：{message}"
                for message in item.get("quality_issues") or []
            ]
            if item_quality_issues:
                item["status"] = "blocked"
                item["blocking_quality_issues"] = item_quality_issues
                quality_issues.extend(item_quality_issues)
                blockers.extend(item_quality_issues)
    if global_quality_blockers:
        for item in templates:
            if item.get("status") == "ready":
                item["status"] = "blocked"
            item["global_quality_blockers"] = list(global_quality_blockers)
    blockers = list(dict.fromkeys(str(item) for item in blockers if str(item).strip()))
    frozen_report = {
        "report_id": int(latest_report.get("id") or 0),
        "report_version": int(latest_report.get("report_version") or 0),
        "template_version": str(latest_report.get("template_version") or ""),
        "status": str(latest_report.get("status") or "draft"),
        "fact_snapshot": snapshot,
        "report_text": report_text,
        "snapshot_sha256": _stable_snapshot_hash(snapshot),
    }
    plan = {
        "plan_version": _TEMPLATE_PLAN_VERSION,
        "status": "blocked" if blockers else "ready",
        "engagement_id": engagement_id,
        "requested_types": list(selected_types),
        "created_at": datetime.now().isoformat(timespec="seconds"),
        "material_read": {
            key: value
            for key, value in material_index.items()
            if key not in {"labels"}
        },
        "material_index": material_index,
        "material_catalog": material_catalog,
        "frozen_report": frozen_report,
        "templates": templates,
        "blockers": blockers,
        "quality_issues": list(dict.fromkeys(quality_issues)),
        "summary": {
            "template_count": len(templates),
            "mapped_field_count": sum(sum(field.get("status") == "mapped" for field in item.get("fields") or []) for item in templates),
            "required_field_count": sum(len(item.get("required_fields") or []) for item in templates),
            "matched_material_label_count": sum(int(item.get("matched_material_label_count") or 0) for item in templates),
            "main_workpaper_sheet_count": material_index.get("sheet_count", 0),
            "main_workpaper_nonempty_rows": material_index.get("nonempty_row_count", 0),
        },
    }
    return plan


def _load_latest_report(engagement_id: int, settings: Settings) -> dict[str, Any]:
    with mysql_connection(settings) as connection:
        with connection.cursor() as cursor:
            cursor.execute(
                """
                SELECT id, report_version, template_version, status, fact_snapshot_json
                FROM audit_report
                WHERE engagement_id = %s AND report_type = 'annual_audit_draft'
                ORDER BY report_version DESC LIMIT 1
                """,
                (engagement_id,),
            )
            row = dict(cursor.fetchone() or {})
    if not row:
        raise ValueError("当前项目尚未生成审计结果，不能生成附件")
    snapshot = row.get("fact_snapshot_json")
    if isinstance(snapshot, str):
        try:
            snapshot = json.loads(snapshot)
        except json.JSONDecodeError:
            snapshot = {}
    row["snapshot"] = snapshot if isinstance(snapshot, dict) else {}
    return row


def _next_package_version(engagement_id: int, settings: Settings) -> int:
    with mysql_connection(settings) as connection:
        with connection.cursor() as cursor:
            cursor.execute(
                "SELECT COALESCE(MAX(package_version), 0) AS latest FROM annual_audit_attachment_package WHERE engagement_id = %s",
                (engagement_id,),
            )
            return int((cursor.fetchone() or {}).get("latest") or 0) + 1


def _persist_package(
    *,
    engagement_id: int,
    package_version: int,
    status: str,
    template_snapshot: dict[str, Any],
    artifacts: list[dict[str, Any]],
    created_by: str,
    settings: Settings,
) -> int:
    with mysql_connection(settings) as connection:
        with connection.cursor() as cursor:
            cursor.execute(
                """
                INSERT INTO annual_audit_attachment_package
                  (engagement_id, package_version, status, template_snapshot_json,
                   artifact_refs_json, created_by)
                VALUES (%s, %s, %s, %s, %s, %s)
                """,
                (
                    engagement_id,
                    package_version,
                    status,
                    json.dumps(template_snapshot, ensure_ascii=False, default=_json_default),
                    json.dumps(artifacts, ensure_ascii=False, default=_json_default),
                    created_by or "ai_agent",
                ),
            )
            package_id = int(cursor.lastrowid)
        connection.commit()
    return package_id


def _replace_tokens(value: str, context: dict[str, Any]) -> str:
    def replace(match: re.Match[str]) -> str:
        key = next((item for item in match.groups() if item), "")
        if key not in context:
            return match.group(0)
        value = context[key]
        return "" if value is None else str(value)

    result = _TOKEN_RE.sub(replace, value)

    def replace_single(match: re.Match[str]) -> str:
        key = match.group(1)
        if key not in context:
            return match.group(0)
        replacement = context[key]
        return "" if replacement is None else str(replacement)

    return _SINGLE_BRACKET_TOKEN_RE.sub(replace_single, result)


def _chinese_year(year: int) -> str:
    digits = "〇一二三四五六七八九"
    return "".join(digits[int(char)] for char in str(year))


def _replace_template_markers(value: str, context: dict[str, Any]) -> str:
    """Fill both explicit tokens and the literal markers used by supplied templates.

    The reference documents in ``new_docs`` predate tokenized templates: they
    use phrases such as ``一般企业模板`` and ``【被审计单位名称】``.  Treating
    those phrases as first-class markers is what makes the uploaded customer
    templates useful without mutating the original file.
    """

    result = _replace_tokens(value, context)
    entity_name = str(context.get("entity_name") or "")
    fiscal_year = int(context.get("fiscal_year") or 0)
    issue_year = int(context.get("issue_year") or (fiscal_year + 1 if fiscal_year else 0))
    period_end = str(context.get("period_end") or "")
    period_start = str(context.get("period_start") or "")
    audit_period = str(context.get("audit_period") or "")
    engagement_code = str(context.get("engagement_code") or "")
    report_number = str(context.get("report_number") or "").strip()
    engagement_partner = str(context.get("engagement_partner") or "")
    if fiscal_year:
        result = result.replace("二〇二五年度", f"{_chinese_year(fiscal_year)}年度")
        result = result.replace("二○二五年度", f"{_chinese_year(fiscal_year)}年度")
        result = result.replace("2025年度", f"{fiscal_year}年度")
        result = result.replace("2025年12月31日", period_end)
    if issue_year:
        result = result.replace("2026年度", f"{issue_year}年度")
        result = result.replace("2026年12月31日", period_end)
    # The supplied reference templates use “一般企业” as a literal
    # placeholder in both headings and body text.  Replace the complete
    # placeholder, not only the few variants that happen to be followed by
    # “模板/报表/附注”.
    result = result.replace("一般企业模板", entity_name)
    result = result.replace("一般企业", entity_name)
    result = result.replace("XXX公司", entity_name)
    # The supplied Beijing case output masks the primary entity as
    # ``北京****有限公司``.  When that exact workpaper replay is being
    # rendered, restore the engagement entity from the case profile instead
    # of carrying the masked reference text into the customer-facing draft.
    if context.get("__case_workpaper_replay__"):
        result = result.replace("北京****有限公司", entity_name)
    result = re.sub(r"一般企业(?=(?:模板|报表|附注))", entity_name, result)
    result = result.replace("【被审计单位名称】", entity_name)
    result = result.replace("[被审计单位名称]", entity_name)
    result = result.replace("【年度】", f"{fiscal_year}年度" if fiscal_year else "")
    result = result.replace("【20XX年度】", f"{fiscal_year}年度" if fiscal_year else "")
    result = result.replace("20XX年度", f"{fiscal_year}年度" if fiscal_year else "")
    result = result.replace("【日期】", period_end)
    result = result.replace("[日期]", period_end)
    result = result.replace("【资产负债日】", period_end)
    result = result.replace("【审计期间】", audit_period or f"{period_start} 至 {period_end}")
    result = result.replace("【项目编号】", engagement_code)
    result = result.replace("[项目编号]", engagement_code)
    if report_number:
        # Reference report templates commonly contain a blank number such as
        # “京创会审字[2026]第 号”.  The supplied case number is carried by
        # the engagement name and must remain visible in the rendered draft.
        result = re.sub(r"京创会审字\[\d{4}\]第\s*号", report_number, result)
        report_year_match = re.search(r"\[(\d{4})\]", report_number)
        if report_year_match:
            result = re.sub(
                r"20\d{2}年\s*月\s*日",
                "",
                result,
            )
    result = result.replace("[注册会计师签名]", "")
    result = result.replace("【注册会计师签名】", "")
    result = result.replace("[项目合伙人姓名]", engagement_partner)
    result = result.replace("【项目合伙人姓名】", engagement_partner)
    result = result.replace("[北京今创会计师事务所（普通合伙）盖章]", "")
    result = result.replace("【北京今创会计师事务所（普通合伙）盖章】", "")
    result = result.replace("[在此处报告与业务经营、税务、管理实务相关的建议和说明。]", "")
    result = result.replace("【在此处报告与业务经营、税务、管理实务相关的建议和说明。】", "")
    result = _BRACKET_PLACEHOLDER_RE.sub("", result)
    result = result.replace("[或适用]", "适用")
    result = result.replace("20XX", str(fiscal_year) if fiscal_year else "")
    # A template marker without a mapped fact must become a blank slot, never
    # an internal workflow sentence or copied sample value.
    result = re.sub(r"(?<![A-Za-z])X{2,}(?![A-Za-z])", "", result)
    # Do not turn an unresolved template marker into prose such as
    # “主底稿未提供元”.  That looks like a populated disclosure but is neither
    # a number nor an auditable conclusion.  Unsupported optional sections are
    # removed by the section plan; a marker that survives the plan is a real
    # template-contract failure and is rejected by the validation gate below.
    return result


def _replace_docx_paragraph(paragraph: Any, context: dict[str, Any]) -> bool:
    """Replace template text without flattening the paragraph's run styling.

    Word often splits a placeholder across several ``w:r`` runs so that a
    part of it can be bold, underlined, or use a different East-Asian font.
    Assigning ``paragraph.text`` (or writing the full replacement into run 0)
    recreates the paragraph as plain text and is the primary cause of broken
    report typography.  The diff is instead placed back into the original run
    carriers: unchanged text stays in its original run and a replacement uses
    the run containing the start of the replaced marker.
    """

    runs = list(paragraph.runs)
    original = "".join(run.text or "" for run in runs)
    if not original:
        return False
    replaced = _replace_template_markers(original, context)
    return _rewrite_docx_paragraph_runs(
        paragraph,
        original=original,
        replacement=replaced,
    )


def _replace_docx_paragraph_exact_replacements(
    paragraph: Any,
    replacements: list[tuple[str, str]],
) -> bool:
    """Apply administrator-approved inline substitutions without global rules.

    The scanner always consumes the source paragraph, never a previously
    substituted value.  This prevents a value that happens to contain another
    source phrase from being rewritten by a later replacement.
    """

    runs = list(paragraph.runs)
    original = "".join(run.text or "" for run in runs)
    if not original:
        return False
    if original != str(paragraph.text or ""):
        raise ValueError("DOCX 行内槽位不能包含未建模的域、制表符或换行")
    if any(
        any(control in value for control in ("\t", "\r", "\n"))
        or any(control in expected_text for control in ("\t", "\r", "\n"))
        for expected_text, value in replacements
    ):
        raise ValueError("DOCX 行内槽位 replacement 不能跨越或写入制表符、换行")
    ordered = sorted(replacements, key=lambda item: len(item[0]), reverse=True)
    position = 0
    output: list[str] = []
    while position < len(original):
        matched = False
        for expected_text, value in ordered:
            if original.startswith(expected_text, position):
                output.append(value)
                position += len(expected_text)
                matched = True
                break
        if not matched:
            output.append(original[position])
            position += 1
    replacement = "".join(output)
    if _docx_paragraph_has_direct_tabs(paragraph):
        return _rewrite_docx_paragraph_inline_replacements_preserving_tabs(
            paragraph,
            original=original,
            replacement=replacement,
            replacements=ordered,
        )
    return _rewrite_docx_paragraph_runs(
        paragraph,
        original=original,
        replacement=replacement,
    )


def _docx_paragraph_has_direct_tabs(paragraph: Any) -> bool:
    """Return whether a writer-safe direct run includes a Word tab element."""

    from docx.oxml.ns import qn

    tab_tag = qn("w:tab")
    return any(run._r.find(tab_tag) is not None for run in paragraph.runs)


def _rewrite_docx_paragraph_inline_replacements_preserving_tabs(
    paragraph: Any,
    *,
    original: str,
    replacement: str,
    replacements: list[tuple[str, str]],
) -> bool:
    """Rewrite direct text while retaining every ``w:tab`` element in place.

    ``Run.text`` presents a tab as ``\t``. Writing that value back into a run
    can leave both a real ``w:tab`` and a literal tab in ``w:t``. This path
    maps the source stream to concrete direct text nodes, treats tab elements
    as immutable separators, and writes only those text nodes.
    """

    if replacement == original:
        return False

    from docx.oxml.ns import qn

    text_tag = qn("w:t")
    tab_tag = qn("w:tab")
    text_nodes: list[Any] = []
    text_node_indexes: dict[int, int] = {}
    stream: list[tuple[int | None, str]] = []
    for run in paragraph.runs:
        for child in run._r:
            if child.tag == text_tag:
                node_index = text_node_indexes.get(id(child))
                if node_index is None:
                    node_index = len(text_nodes)
                    text_node_indexes[id(child)] = node_index
                    text_nodes.append(child)
                stream.extend(
                    (node_index, character) for character in child.text or ""
                )
            elif child.tag == tab_tag:
                stream.append((None, "\t"))

    if "".join(character for _node_index, character in stream) != original:
        raise ValueError("DOCX 行内槽位不能安全保留制表符位置")

    output_by_text_node = ["" for _node in text_nodes]
    position = 0
    while position < len(original):
        matched = False
        for expected_text, value in replacements:
            if not original.startswith(expected_text, position):
                continue
            match_end = position + len(expected_text)
            match_tokens = stream[position:match_end]
            start_node_index = stream[position][0] if position < len(stream) else None
            if (
                start_node_index is None
                or any(node_index is None for node_index, _character in match_tokens)
            ):
                raise ValueError("DOCX 行内槽位 replacement 不能跨越制表符")
            output_by_text_node[start_node_index] += value
            position = match_end
            matched = True
            break
        if matched:
            continue
        node_index, character = stream[position]
        if node_index is not None:
            output_by_text_node[node_index] += character
        position += 1

    for text_node, value in zip(text_nodes, output_by_text_node):
        if value != str(text_node.text or ""):
            _set_docx_text_node_value(text_node, value)
    if str(paragraph.text or "") != replacement:
        raise ValueError("DOCX 行内槽位不能安全保留制表符位置")
    _clear_docx_authoring_annotations(paragraph)
    return True


def _set_docx_text_node_value(text_node: Any, value: str) -> None:
    """Update one text node without rebuilding its containing run."""

    text_node.text = value
    space_attr = "{http://www.w3.org/XML/1998/namespace}space"
    if value[:1].isspace() or value[-1:].isspace():
        text_node.set(space_attr, "preserve")
    else:
        text_node.attrib.pop(space_attr, None)


def _rewrite_docx_paragraph_runs(
    paragraph: Any,
    *,
    original: str,
    replacement: str,
) -> bool:
    """Move a text diff through existing Word runs without flattening style."""

    runs = list(paragraph.runs)
    if replacement == original:
        return False
    if not runs:
        paragraph.add_run(replacement)
        return True

    boundaries: list[tuple[int, int]] = []
    position = 0
    for run in runs:
        next_position = position + len(run.text or "")
        boundaries.append((position, next_position))
        position = next_position
    output_by_run = ["" for _ in runs]

    def run_index_for(position_hint: int) -> int:
        for index, (start, end) in enumerate(boundaries):
            if start <= position_hint < end:
                return index
        for index in range(len(boundaries) - 1, -1, -1):
            if boundaries[index][1] > boundaries[index][0]:
                return index
        return 0

    def append_original(start: int, end: int) -> None:
        for index, (run_start, run_end) in enumerate(boundaries):
            left = max(start, run_start)
            right = min(end, run_end)
            if left < right:
                output_by_run[index] += original[left:right]

    matcher = SequenceMatcher(a=original, b=replacement, autojunk=False)
    for tag, old_start, old_end, new_start, new_end in matcher.get_opcodes():
        if tag == "equal":
            append_original(old_start, old_end)
            continue
        if tag in {"replace", "insert"} and new_start < new_end:
            output_by_run[run_index_for(old_start)] += replacement[new_start:new_end]

    for run, text in zip(runs, output_by_run):
        if text != (run.text or ""):
            _set_docx_run_text_in_place(run, text)
    # A highlight on a placeholder is a drafting cue, not a customer-facing
    # style.  Leave all untouched template styling alone, but clear the cue on
    # the one paragraph whose value was actually replaced.
    _clear_docx_authoring_annotations(paragraph)
    return True


def _set_docx_run_text_in_place(run: Any, value: str) -> None:
    """Change a run's text nodes without flattening fields, tabs or breaks.

    ``Run.text = ...`` reconstructs the run's contents and silently drops
    inline ``w:fldChar``, ``w:instrText``, ``w:tab`` and ``w:br`` elements.
    A financial-statement template relies heavily on those elements for its
    titles, page layout and calculated fields, so only existing ``w:t`` nodes
    are changed here.
    """

    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    text_nodes = list(run._r.findall(qn("w:t")))
    if not text_nodes:
        if not value:
            return
        text_node = OxmlElement("w:t")
        text_node.text = value
        if value[:1].isspace() or value[-1:].isspace():
            text_node.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        run._r.append(text_node)
        return

    text_nodes[0].text = value
    if value[:1].isspace() or value[-1:].isspace():
        text_nodes[0].set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    else:
        text_nodes[0].attrib.pop("{http://www.w3.org/XML/1998/namespace}space", None)
    for text_node in text_nodes[1:]:
        text_node.text = ""


def _replace_docx_container(container: Any, context: dict[str, Any]) -> int:
    changed = sum(_replace_docx_paragraph(item, context) for item in container.paragraphs)
    for table in container.tables:
        for row in table.rows:
            for cell in row.cells:
                changed += _replace_docx_container(cell, context)
    return changed


def _iter_docx_tables(container: Any):
    """Yield top-level and nested tables without yielding a nested table twice."""

    seen: set[int] = set()

    def walk(owner: Any):
        for table in getattr(owner, "tables", ()):
            marker = id(table)
            if marker in seen:
                continue
            seen.add(marker)
            yield table
            for row in table.rows:
                for cell in row.cells:
                    yield from walk(cell)

    yield from walk(container)


def _docx_table_has_outer_borders(table: Any) -> tuple[bool, bool]:
    """Return whether a table has effective direct left and right borders."""

    from docx.oxml.ns import qn

    table_properties = table._tbl.tblPr
    borders = table_properties.first_child_found_in("w:tblBorders")
    if borders is None:
        return False, False
    sides = {}
    for side in ("left", "right"):
        element = borders.find(qn(f"w:{side}"))
        sides[side] = bool(element is not None and (element.get(qn("w:val")) or "single") not in {"nil", "none"})
    return bool(sides["left"]), bool(sides["right"])


def _ensure_docx_table_outer_borders(document: Any, context: dict[str, Any]) -> int:
    """Add missing direct left/right borders while preserving the style frame.

    Supplied note templates commonly put top/bottom/internal borders in a
    table style but omit the two outer sides.  Adding only the missing direct
    sides keeps the uploaded table geometry, widths, fonts and existing border
    definitions intact while making every retained note table printable and
    visually closed.
    """

    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    touched_tables = 0
    touched_sides = 0
    for table in _iter_docx_tables(document):
        table_properties = table._tbl.tblPr
        borders = table_properties.first_child_found_in("w:tblBorders")
        if borders is None:
            borders = OxmlElement("w:tblBorders")
            table_properties.append(borders)
        table_touched = False
        for side in ("left", "right"):
            element = borders.find(qn(f"w:{side}"))
            value = element.get(qn("w:val")) if element is not None else None
            if element is not None and value not in {None, "nil", "none"}:
                continue
            if element is None:
                element = OxmlElement(f"w:{side}")
                borders.append(element)
            element.set(qn("w:val"), "single")
            element.set(qn("w:sz"), "12")
            element.set(qn("w:space"), "0")
            element.set(qn("w:color"), "auto")
            touched_sides += 1
            table_touched = True
        if table_touched:
            touched_tables += 1

    context["__notes_outer_border_table_count__"] = touched_tables
    context["__notes_outer_border_side_count__"] = touched_sides
    return touched_tables


def _docx_document_text(document: Any) -> str:
    """Collect body/table text without materialising new header/footer parts."""

    parts = [paragraph.text or "" for paragraph in document.paragraphs]
    parts.extend(
        cell.text or ""
        for table in _iter_docx_tables(document)
        for row in table.rows
        for cell in row.cells
    )
    return "\n".join(parts)


def _docx_ooxml_text_parts(data: bytes) -> dict[str, str]:
    """Read client-visible text from every relevant DOCX story part.

    python-docx's ``Document.paragraphs`` intentionally exposes only the main
    body.  Delivery validation must additionally inspect headers, footers,
    footnotes, endnotes, comments and text boxes (which live inside those XML
    parts), otherwise an internal source locator can leak outside the body.
    """

    from xml.etree import ElementTree

    namespace = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    visible_tags = {
        f"{namespace}t",
        f"{namespace}delText",
        f"{namespace}instrText",
    }
    part_name = re.compile(
        r"^word/(?:document|header\d+|footer\d+|footnotes|endnotes|comments)\.xml$"
    )
    result: dict[str, str] = {}
    with zipfile.ZipFile(BytesIO(data)) as archive:
        for name in archive.namelist():
            if not part_name.fullmatch(name):
                continue
            try:
                root = ElementTree.fromstring(archive.read(name))
            except ElementTree.ParseError:
                continue
            values = [node.text or "" for node in root.iter() if node.tag in visible_tags]
            result[name] = "".join(values)
    return result


def _set_docx_cell_text(cell: Any, value: Any) -> None:
    """Set one data cell while preserving its paragraph/cell formatting.

    The renderer only calls this for an identified value cell.  It deliberately
    leaves additional paragraphs (notes, line breaks and footnote markers) in
    place rather than rebuilding the entire ``w:tc`` node.
    """

    text = "" if value is None else str(value)
    if not cell.paragraphs:
        cell.add_paragraph()
    _set_docx_paragraph_text(cell.paragraphs[0], text)


def _materialize_docx_docvariable_fields(paragraph: Any) -> bool:
    """Turn declared DOCVARIABLE results into stable literal Word text.

    The supplied notes master uses DOCVARIABLE fields for the entity and year
    in its header.  Editing only their cached ``w:t`` result lets Word restore
    the original variable when fields are refreshed.  We materialise only
    DOCVARIABLE fields; PAGE, REF and other field types carry document
    behavior and are deliberately rejected for client-slot replacement.
    """

    from docx.oxml.ns import qn

    instr_tag = qn("w:instrText")
    field_tag = qn("w:fldChar")
    simple_field_tag = qn("w:fldSimple")
    simple_instr_attr = qn("w:instr")
    field_type_attr = qn("w:fldCharType")
    field_nodes = list(paragraph._p.iter(field_tag))
    instruction_nodes = list(paragraph._p.iter(instr_tag))
    simple_fields = list(paragraph._p.iter(simple_field_tag))
    if not field_nodes and not simple_fields:
        return False
    commands: list[str] = []
    stack: list[list[str]] = []
    for node in paragraph._p.iter():
        if node.tag == field_tag:
            field_type = str(node.get(field_type_attr) or "").lower()
            if field_type == "begin":
                stack.append([])
            elif field_type == "end" and stack:
                commands.append("".join(stack.pop()))
        elif node.tag == instr_tag and stack:
            stack[-1].append(str(node.text or ""))
        elif node.tag == simple_field_tag:
            commands.append(str(node.get(simple_instr_attr) or ""))
    if stack or not commands or any(
        not re.match(r"^\s*DOCVARIABLE(?:\s|$)", command, re.IGNORECASE)
        for command in commands
    ):
        raise ValueError("DOCX 槽位不能改写非 DOCVARIABLE 的 Word 域")

    # A simple field owns its cached result runs.  Lift those runs into the
    # paragraph before removing the field wrapper so the normal in-place text
    # writer can preserve their font and paragraph design.
    for simple_field in simple_fields:
        parent = simple_field.getparent()
        if parent is None:
            continue
        position = parent.index(simple_field)
        for child in list(simple_field):
            simple_field.remove(child)
            parent.insert(position, child)
            position += 1
        parent.remove(simple_field)
    for node in [*field_nodes, *instruction_nodes]:
        parent = node.getparent()
        if parent is not None:
            parent.remove(node)
    return True


def _set_docx_paragraph_text(paragraph: Any, value: str) -> None:
    """Replace paragraph text without replacing its paragraph/style object.

    Only the written paragraph loses direct highlight/shading used to signal a
    template author's editable slot.  Other template paragraphs remain
    byte-for-byte equivalent at the layout level.
    """

    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    text_nodes = list(paragraph._p.findall(".//" + qn("w:t")))
    if not text_nodes:
        if not value:
            return
        run = OxmlElement("w:r")
        text_node = OxmlElement("w:t")
        run.append(text_node)
        paragraph._p.append(run)
        text_nodes = [text_node]
    text_nodes[0].text = str(value)
    text_nodes[0].set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    for text_node in text_nodes[1:]:
        text_node.text = ""
    _clear_docx_authoring_annotations(paragraph)


def _normalize_template_label(value: Any) -> str:
    """Normalise a template row label without losing its total-row meaning."""

    text = str(value or "").strip().replace("　", " ")
    text = re.sub(r"^[（(]?[一二三四五六七八九十百0-9]+[）).、．]\s*", "", text)
    text = re.sub(r"^(?:其中|加|减)[:：]", "", text)
    text = re.sub(r"^[一二三四五六七八九十]+[、.．]", "", text)
    return re.sub(r"[\s　\t\r\n:：()（）\[\]【】]", "", text)


def _is_total_template_label(label: str) -> bool:
    compact = _normalize_template_label(label)
    return compact in {"合计", "小计", "总计"} or compact.endswith(("合计", "总计"))


_INDEPENDENT_STATEMENT_TOTALS = frozenset(
    {
        "资产总计",
        "负债合计",
        "所有者权益合计",
        "股东权益合计",
        "负债和所有者权益总计",
        "负债及所有者权益总计",
        "负债和股东权益总计",
        "负债及股东权益总计",
        "利润总额",
        "净利润",
        "现金及现金等价物净增加额",
    }
)


def _is_independent_statement_total(label: str) -> bool:
    return _normalize_template_label(label) in _INDEPENDENT_STATEMENT_TOTALS


def _fact_registry_from_context(context: dict[str, Any]) -> dict[str, dict[str, Any]]:
    """Return periodised facts, never a best-effort label/number grab.

    A fact is keyed by its standard accounting label and contains current and
    comparative values separately.  This keeps a table cell traceable to the
    project fact registry and prevents the old behaviour that used the first
    number found anywhere on a worksheet row.
    """

    saved = context.get("__fact_registry__")
    if isinstance(saved, dict):
        return saved
    registry: dict[str, dict[str, Any]] = {}
    values = context.get("__table_values__") or {}
    if not isinstance(values, dict):
        values = {}
    for raw_key, raw_value in values.items():
        key = str(raw_key or "")
        period = "opening" if key.endswith("_期初") else "current"
        label = key[:-3] if period == "opening" else key
        try:
            value = float(raw_value)
        except (TypeError, ValueError):
            continue
        normalized = _normalize_material_label(_normalize_template_label(label))
        if not normalized or normalized in _MATERIAL_NOISE_LABELS:
            continue
        entry = registry.setdefault(normalized, {"label": label, "current": None, "opening": None, "source": "project_fact"})
        entry[period] = value
    context["__fact_registry__"] = registry
    return registry


def _prepare_template_render_context(
    context: dict[str, Any], *, file_name: str, extension: str
) -> dict[str, Any]:
    """Initialise the format-neutral rendering contract for one template.

    A template's filename and its business attachment type must never decide
    which audit figure is written to a cell.  This small mutable run record is
    also persisted with the generated artifact, so a reviewer can distinguish
    a true template rendering from a returned source file.
    """

    context.setdefault("__source_file_name__", file_name)
    registry = _fact_registry_from_context(context)
    contract = context.get("__template_render_contract__")
    if not isinstance(contract, dict):
        contract = {}
        context["__template_render_contract__"] = contract
    contract.update(
        {
            "contract_version": _TEMPLATE_RENDER_CONTRACT_VERSION,
            "source_file_name": str(context.get("__source_file_name__") or file_name),
            "source_extension": extension,
            "template_type": str(context.get("__template_type__") or ""),
            "fact_registry_count": len(registry),
            "source_priority": "workpaper_or_statement>trial_balance>analysis_snapshot",
            "placement_policy": "discovered_template_slots_only",
            "format_preservation_policy": "edit_in_place_no_append_no_rebuild",
            "authoring_annotation_policy": "clear_written_slot_authoring_marks",
        }
    )
    return contract


def _record_template_render_result(
    context: dict[str, Any], *, renderer: str, changed: int, status: str, **details: Any
) -> None:
    """Record concise, serialisable evidence common to every renderer."""

    contract = context.get("__template_render_contract__")
    if not isinstance(contract, dict):
        contract = _prepare_template_render_context(
            context,
            file_name=str(context.get("__source_file_name__") or "template"),
            extension="",
        )
    contract.update(
        {
            "renderer": renderer,
            "changed_slot_count": int(changed),
            "render_status": status,
        }
    )
    contract.update({key: value for key, value in details.items() if value is not None})


def _template_render_contract_evidence(context: dict[str, Any]) -> dict[str, Any]:
    """Return the persisted subset of a renderer run without raw source data."""

    contract = context.get("__template_render_contract__") or {}
    if not isinstance(contract, dict):
        return {}
    allowed = {
        "contract_version",
        "source_file_name",
        "source_extension",
        "template_type",
        "fact_registry_count",
        "source_priority",
        "placement_policy",
        "format_preservation_policy",
        "authoring_annotation_policy",
        "renderer",
        "changed_slot_count",
        "render_status",
        "audit_result_mapped",
        "audit_result_mapping_reason",
        "declared_slot_semantic_coverage",
        "discovered_table_count",
        "mapped_table_count",
        "mapped_value_cell_count",
        "docx_field_schema_slot_count",
        "output_table_count",
        "sheet_count",
        "content_profile",
        "notes_outer_border_table_count",
        "notes_outer_border_side_count",
    }
    return {key: value for key, value in contract.items() if key in allowed}


def _build_periodised_fact_registry(
    table_values: dict[str, Any],
    *,
    statement_values: dict[str, Any] | None = None,
    balance_by_name: dict[str, Any] | None = None,
) -> dict[str, dict[str, Any]]:
    """Build source-prioritised facts for a template rendering run.

    Source priority is a data concern, not a template concern: audited
    statement/workpaper values override an imported trial balance, which in
    turn overrides a high-level analysis snapshot.  Current and comparative
    periods are prioritised independently so an unavailable comparative value
    never erases a verified current value.
    """

    registry: dict[str, dict[str, Any]] = {}

    def add(label: Any, current: Any, opening: Any, *, source: str, priority: int) -> None:
        normalized = _normalize_material_label(_normalize_template_label(label))
        if not normalized or normalized in _MATERIAL_NOISE_LABELS:
            return
        entry = registry.setdefault(
            normalized,
            {
                "label": str(label),
                "current": None,
                "opening": None,
                "source": source,
                "_current_priority": -1,
                "_opening_priority": -1,
            },
        )
        for period, candidate in (("current", current), ("opening", opening)):
            if not isinstance(candidate, (int, float)) or isinstance(candidate, bool):
                continue
            priority_key = f"_{period}_priority"
            if priority >= int(entry.get(priority_key) or -1):
                entry[period] = float(candidate)
                entry[priority_key] = priority
                entry["source"] = source

    for label, value in table_values.items():
        text = str(label or "")
        if text.endswith("_期初"):
            add(text[:-3], None, value, source="analysis_snapshot", priority=10)
        else:
            add(text, value, None, source="analysis_snapshot", priority=10)

    for label, account in (balance_by_name or {}).items():
        if not isinstance(account, dict):
            continue
        current = abs(float(account.get("closing_debit") or 0) - float(account.get("closing_credit") or 0))
        opening = abs(float(account.get("opening_debit") or 0) - float(account.get("opening_credit") or 0))
        add(label, current, opening, source="trial_balance", priority=60)

    for label, item in (statement_values or {}).items():
        if not isinstance(item, dict):
            continue
        add(
            label,
            item.get("current"),
            item.get("opening"),
            source=f"workpaper:{item.get('source_sheet') or 'statement'}:{item.get('source_row') or ''}",
            priority=100,
        )

    for item in registry.values():
        item.pop("_current_priority", None)
        item.pop("_opening_priority", None)
    return registry


def _resolve_template_fact(
    registry: dict[str, dict[str, Any]], label: str
) -> tuple[dict[str, Any] | None, int]:
    """Resolve an exact standard fact or declared accounting alias only."""

    normalized = _normalize_material_label(_normalize_template_label(label))
    if not normalized:
        return None, 1
    preferred = _PREFERRED_PRESENTATION_FACTS.get(normalized)
    if preferred:
        fact = registry.get(_normalize_material_label(preferred))
        if isinstance(fact, dict):
            return fact, 1
    direct = registry.get(normalized)
    if isinstance(direct, dict):
        return direct, 1
    alias_key = re.sub(r"[\"'“”‘’\-—–]", "", normalized)
    alias = _STANDARD_FACT_ALIASES.get(normalized) or _STANDARD_FACT_ALIASES.get(
        alias_key
    )
    if alias:
        source, multiplier = alias
        fact = registry.get(_normalize_material_label(source))
        if isinstance(fact, dict):
            return fact, multiplier
    return None, 1


def _remember_rendered_asset_total(
    registry: dict[str, dict[str, Any]],
    *,
    period: str,
    value: Any,
) -> None:
    """Share an asset-side total with the later liabilities/equity side."""

    if period not in {"current", "opening"} or not isinstance(value, (int, float)):
        return
    key = _normalize_material_label("资产总计")
    fact = dict(registry.get(key) or {})
    if isinstance(fact.get(period), (int, float)):
        return
    fact.setdefault("label", "资产总计")
    fact[period] = float(value)
    fact[f"{period}_source"] = "derived:rendered_asset_rows"
    registry[key] = fact


def _docx_table_sections(document: Any) -> list[str]:
    """Map each body table to the nearest heading in the template body."""

    return [item["section"] for item in _docx_table_contexts(document)]


def _docx_table_contexts(document: Any) -> list[dict[str, str]]:
    """Map every body table to its nearest heading and its parent heading."""

    from docx.text.paragraph import Paragraph

    headings: dict[int, str] = {}
    contexts: list[dict[str, str]] = []
    for child in list(document._element.body):
        if str(child.tag).endswith("}p"):
            paragraph = Paragraph(child, document)
            level = _docx_heading_level(paragraph)
            if level is not None and (paragraph.text or "").strip():
                headings[level] = (paragraph.text or "").strip()
                for deeper in [item for item in headings if item > level]:
                    headings.pop(deeper, None)
        elif str(child.tag).endswith("}tbl"):
            section = headings.get(3) or headings.get(2) or headings.get(1) or ""
            parent = headings.get(2) if headings.get(3) else headings.get(1, "")
            contexts.append({"section": section, "parent": parent or ""})
    if len(contexts) < len(document.tables):
        contexts.extend({"section": "", "parent": ""} for _ in range(len(document.tables) - len(contexts)))
    return contexts[: len(document.tables)]


def _table_column_plan(table: Any) -> tuple[int, dict[int, tuple[str, str]]]:
    """Identify a table's data start and period/metric columns from headers."""

    if not table.rows:
        return 0, {}
    max_columns = max((len(row.cells) for row in table.rows), default=0)
    def header_score(value: str) -> int:
        text = value.strip().replace(" ", "")
        if not text:
            return 0
        if text in {"项目", "类别", "名称", "账龄", "税种", "补充资料", "金额", "比例（%）", "计提比例（%）"}:
            return 1
        if text in {"期末", "期初", "本期", "上期", "本年", "上年", "期末余额", "期初余额", "本期数", "上期数", "本期金额", "上期金额"}:
            return 1
        if re.match(r"^(?:期末|期初|本期|上期|本年|上年)", text) and any(
            suffix in text for suffix in ("余额", "金额", "数")
        ):
            return 1
        if any(token in text for token in ("账面余额", "账面价值", "坏账准备", "计提比例")):
            return 1
        if any(token in text for token in ("第一阶段", "第二阶段", "第三阶段", "未来12个月", "整个存续期")):
            return 1
        return 0

    header_depth = 0
    for row_index, row in enumerate(table.rows[:3]):
        scores = [header_score(cell.text or "") for cell in row.cells]
        # The first row can be a compact two-column heading; subsequent rows
        # must contain at least two recognised header cells.  This avoids
        # treating a sample value such as “案例原始期末值” as a second header.
        if (row_index == 0 and sum(scores) >= 2) or (row_index > 0 and sum(scores) >= 2):
            header_depth = row_index + 1
            continue
        break
    if header_depth == 0:
        return 0, {}
    columns: dict[int, tuple[str, str]] = {}
    for column in range(max_columns):
        header = " ".join(
            (table.rows[row_index].cells[column].text or "").strip()
            for row_index in range(header_depth)
            if column < len(table.rows[row_index].cells)
        )
        period = ""
        if any(token in header for token in ("期末", "本期", "本年")):
            period = "current"
        elif any(token in header for token in ("期初", "上期", "上年")):
            period = "opening"
        if not period:
            continue
        if "收入" in header:
            metric = "收入"
        elif "成本" in header:
            metric = "成本"
        elif "本期增加" in header:
            metric = "increase"
        elif "摊销" in header:
            metric = "amortization"
        elif "减少" in header:
            metric = "decrease"
        elif "计提比例" in header:
            metric = "provision_ratio"
        elif "比例" in header:
            metric = "ratio"
        elif "坏账准备" in header:
            metric = "provision"
        elif "账面价值" in header:
            metric = "net"
        elif "账面余额" in header or "金额" in header:
            metric = "amount"
        else:
            metric = ""
        columns[column] = (period, metric)
    return header_depth, columns


def _write_fact_to_cell(cell: Any, value: Any) -> bool:
    if not isinstance(value, (int, float)):
        return False
    _set_docx_cell_text(cell, _money(value))
    return True


def _detail_fact_candidates(registry: dict[str, dict[str, Any]], section: str) -> list[str]:
    """Return the source-backed detail labels available for one note section."""

    prefix = f"{_normalize_material_label(_normalize_template_label(section))}|"
    if not prefix or prefix == "|":
        return []
    candidates: list[str] = []
    for key, fact in registry.items():
        if not key.startswith(prefix):
            continue
        suffix = key[len(prefix):]
        # A base detail fact has exactly ``subject|category``.  Movement facts
        # (``subject|category|本期增加``) support its columns but must not
        # become a separate template row.
        if "|" in suffix or not isinstance(fact, dict):
            continue
        if any(isinstance(fact.get(period), (int, float)) for period in ("current", "opening")):
            candidates.append(suffix)
    return candidates


def _table_cell_source_label(
    row_label: str,
    section: str,
    metric: str,
    registry: dict[str, dict[str, Any]],
) -> tuple[str, str]:
    """Resolve a data-grid cell using declared accounting-table semantics.

    The returned discriminator identifies a ratio cell, which is derived from
    the same table's populated components rather than copied from an amount
    fact.
    This supports common receivable impairment grids without a template-name
    branch or a substring search across unrelated account names.
    """

    label = _normalize_template_label(row_label)
    detail_metric = {
        "increase": "本期增加",
        "amortization": "摊销",
        "decrease": "其他减少",
    }.get(metric, "")
    detail_source = f"{section}|{label}" if section and label else ""
    if detail_metric:
        detail_source = f"{detail_source}|{detail_metric}" if detail_source else ""
    if detail_source and _resolve_template_fact(registry, detail_source)[0]:
        if metric == "ratio":
            return detail_source, "balance"
        if metric == "provision_ratio":
            return f"{detail_source}坏账准备", "provision"
        if metric == "provision":
            return f"{detail_source}坏账准备", ""
        return detail_source, ""
    if metric in {"收入", "成本"}:
        return _COMPOSITE_TABLE_FACTS.get(label, {}).get(metric, ""), ""
    if metric == "increase":
        return f"{section}本期增加", ""
    if metric == "amortization":
        return f"{section}摊销", ""
    if metric == "decrease":
        return f"{section}其他减少", ""
    if label == "其他" and _resolve_template_fact(registry, section)[0]:
        return section, ""

    # Standard impairment schedules describe an account in a longer row name
    # (for example “按信用风险特征组合计提坏账准备的应收账款”).  The row is
    # not a different account.  Limit this containment rule to the two
    # standard receivable roots, never to arbitrary labels in a template.
    root = ""
    for candidate in ("应收账款", "其他应收款"):
        # A broad account balance can support the credit-risk-combination
        # line only.  It must never be copied into a "按单项计提" row merely
        # because the account name appears in that row's prose.
        if section == candidate and any(token in label for token in ("按信用", "账龄组合")):
            root = candidate
            break
    base = root or label
    if metric == "provision":
        return f"{base}坏账准备", ""
    if metric == "provision_ratio":
        return f"{base}坏账准备", "provision"
    if metric == "ratio":
        return base, "balance"
    return base, ""


def _table_fact_period_for_row(row_label: str, table_period: str) -> str | None:
    """Apply a row's explicit period semantics to a year-comparative table."""

    label = _normalize_template_label(row_label)
    # In a cash-flow supplement, the "本期金额" for "现金的期初余额"
    # is the opening balance of the cash fact, rather than its closing balance.
    # The prior-year opening is deliberately left blank unless independently
    # sourced; copying the same number across both columns is misleading.
    if "现金" in label and any(token in label for token in ("期初余额", "年初余额")):
        return "opening" if table_period == "current" else None
    return table_period


def _fill_docx_tables(document: Any, context: dict[str, Any]) -> int:
    """Fill planned value cells while preserving the template's table layout.

    Unlike the earlier substring-based routine, this renderer never selects a
    random account merely because its name appears somewhere in a table.  It
    maps period columns from the header, resolves each row through the audited
    fact registry, and calculates a total only from its section fact or from
    values actually written to that same table.
    """

    registry = _fact_registry_from_context(context)
    table_contexts = _docx_table_contexts(document)
    composition_manifest = context.get("__attachment_composition_manifest__") or {}
    prototype_bindings = {
        int(item.get("output_table_index")): str(item.get("fact_section_key") or "")
        for item in (composition_manifest.get("prototype_table_bindings") or [])
        if isinstance(item, dict) and str(item.get("output_table_index", "")).isdigit()
    }
    changed = 0
    table_stats: list[dict[str, Any]] = []
    for table_index, table in enumerate(document.tables):
        header_depth, columns = _table_column_plan(table)
        table_context = table_contexts[table_index] if table_index < len(table_contexts) else {"section": "", "parent": ""}
        section = str(table_context.get("section") or "")
        parent_section = str(table_context.get("parent") or "")
        prototype_fact_section = prototype_bindings.get(table_index, "")
        fact_section = next(
            (
                candidate
                for candidate in (prototype_fact_section, section, parent_section)
                if candidate and (
                    _resolve_template_fact(registry, candidate)[0]
                    or _detail_fact_candidates(registry, candidate)
                )
            ),
            section,
        )
        stats: dict[str, Any] = {
            "table_index": table_index,
            "section": section,
            "fact_section": fact_section,
            "header_rows": header_depth,
            "period_columns": {str(column): period for column, (period, _) in columns.items()},
            "filled_cells": 0,
            "mapped_rows": [],
        }
        if not columns or header_depth >= len(table.rows):
            table_stats.append(stats)
            continue

        # A template is a layout contract, not a source of financial facts.
        # Clear literal example/default amounts in recognised amount columns
        # before applying this engagement's fact registry.  Otherwise a stale
        # template ``0.00`` can survive as an invented audit conclusion.
        default_numeric = re.compile(r"^[-+]?\d{1,3}(?:,\d{3})*(?:\.\d+)?(?:%|％)?$")
        for row in table.rows[header_depth:]:
            for column in columns:
                if column >= len(row.cells):
                    continue
                if default_numeric.fullmatch((row.cells[column].text or "").strip()):
                    _set_docx_cell_text(row.cells[column], "")
                    changed += 1
                    stats["cleared_template_default_cells"] = int(stats.get("cleared_template_default_cells") or 0) + 1

        row_written: dict[int, dict[tuple[int, str], float]] = {}
        computed_totals: dict[int, dict[tuple[int, str], float]] = {}
        total_rows: list[tuple[int, Any, str]] = []
        used_detail_labels: set[str] = set()
        claimed_detail_rows: set[str] = set()
        detail_candidates = _detail_fact_candidates(registry, fact_section)
        for row in table.rows[header_depth:]:
            if row.cells:
                label = _normalize_template_label(row.cells[0].text or "")
                if label in detail_candidates:
                    used_detail_labels.add(label)

        # Some standard disclosure templates leave a repeat-row label blank.
        # When the workpaper offers an unambiguous, source-backed category,
        # place that category into the template's existing first-column cell
        # and let the normal row/column resolver populate its amounts.  This
        # is a structural table rule, not a template-name special case.
        for row_index, row in enumerate(table.rows[header_depth:], start=header_depth):
            if not row.cells or (row.cells[0].text or "").strip():
                continue
            candidate = next((item for item in detail_candidates if item not in used_detail_labels), "")
            if not candidate:
                break
            _set_docx_cell_text(row.cells[0], candidate)
            used_detail_labels.add(candidate)
            stats.setdefault("inserted_detail_rows", []).append({"row": row_index, "label": candidate})

        for row_index, row in enumerate(table.rows[header_depth:], start=header_depth):
            if not row.cells:
                continue
            label = str(row.cells[0].text or "").strip()
            if not label:
                continue
            if _is_total_template_label(label):
                total_rows.append((row_index, row, label))
                continue

            clean_label = _normalize_template_label(label)
            detail_base = f"{fact_section}|{clean_label}" if fact_section and clean_label else ""
            is_detail_row = bool(detail_base and _resolve_template_fact(registry, detail_base)[0])
            if is_detail_row and detail_base in claimed_detail_rows:
                # Identical labels can occur in separate optional groups of a
                # disclosure table.  With one source fact and no additional
                # group axis, filling both would duplicate an audited amount.
                continue
            row_used_detail_fact = False
            for column, (period, metric) in columns.items():
                if column >= len(row.cells):
                    continue
                source_label, ratio_kind = _table_cell_source_label(
                    clean_label, fact_section, metric, registry
                )
                fact, multiplier = _resolve_template_fact(registry, source_label)
                if not fact:
                    continue
                fact_period = _table_fact_period_for_row(label, period)
                if fact_period is None:
                    continue
                value = fact.get(fact_period)
                if not isinstance(value, (int, float)):
                    continue
                numeric_value: float | None = float(value) * multiplier
                if ratio_kind == "balance":
                    denominator, _ = _resolve_template_fact(registry, fact_section)
                    denominator_value = denominator.get(fact_period) if denominator else None
                    if not isinstance(denominator_value, (int, float)) or abs(float(denominator_value)) <= 1e-9:
                        continue
                    numeric_value = float(value) * multiplier / float(denominator_value) * 100.0
                elif ratio_kind == "provision":
                    amount_source, _ = _table_cell_source_label(clean_label, fact_section, "amount", registry)
                    amount_fact, amount_multiplier = _resolve_template_fact(registry, amount_source)
                    amount_value = amount_fact.get(fact_period) if amount_fact else None
                    if not isinstance(amount_value, (int, float)) or abs(float(amount_value)) <= 1e-9:
                        continue
                    numeric_value = float(value) * multiplier / (float(amount_value) * amount_multiplier) * 100.0
                if numeric_value is None:
                    continue
                if _write_fact_to_cell(row.cells[column], numeric_value):
                    changed += 1
                    stats["filled_cells"] += 1
                    stats["mapped_rows"].append({"row": row_index, "label": label, "fact": fact.get("label") or source_label})
                    row_written.setdefault(row_index, {})[(column, period)] = numeric_value
                    if is_detail_row and (source_label.startswith(f"{detail_base}|") or source_label == detail_base):
                        row_used_detail_fact = True
            if row_used_detail_fact:
                claimed_detail_rows.add(detail_base)

        for total_row_index, row, total_label in total_rows:
            # A total without one mapped line is a template artefact, not an
            # evidence-backed disclosure.  Leaving it in place would create
            # the misleading pattern “blank details + total amount”.
            has_non_total_row_slot = any(
                row.cells and not _is_total_template_label(row.cells[0].text or "")
                for row in table.rows[header_depth:]
            )
            if has_non_total_row_slot and not stats["mapped_rows"]:
                continue
            for column, (period, metric) in columns.items():
                if column >= len(row.cells):
                    continue
                if metric == "ratio":
                    total_fact, _ = _resolve_template_fact(registry, fact_section)
                    total_value = total_fact.get(period) if total_fact else None
                    if isinstance(total_value, (int, float)) and abs(float(total_value)) > 1e-9:
                        if _write_fact_to_cell(row.cells[column], 100.0):
                            changed += 1
                            stats["filled_cells"] += 1
                            stats["mapped_rows"].append({"row": "total", "label": total_label, "fact": "derived:total_amount_ratio"})
                    continue
                if metric == "provision_ratio":
                    provision_fact, provision_multiplier = _resolve_template_fact(registry, f"{fact_section}坏账准备")
                    amount_fact, amount_multiplier = _resolve_template_fact(registry, fact_section)
                    provision_value = provision_fact.get(period) if provision_fact else None
                    amount_value = amount_fact.get(period) if amount_fact else None
                    if isinstance(provision_value, (int, float)) and isinstance(amount_value, (int, float)) and abs(float(amount_value)) > 1e-9:
                        ratio = float(provision_value) * provision_multiplier / (float(amount_value) * amount_multiplier) * 100.0
                        if _write_fact_to_cell(row.cells[column], ratio):
                            changed += 1
                            stats["filled_cells"] += 1
                            stats["mapped_rows"].append({"row": "total", "label": total_label, "fact": "derived:total_provision_ratio"})
                    continue
                source_label = ""
                if metric in {"收入", "成本"}:
                    source_label = _COMPOSITE_TABLE_FACTS.get(_normalize_template_label(total_label), {}).get(metric, "")
                elif metric == "provision":
                    source_label = f"{fact_section}坏账准备"
                elif metric in {"increase", "amortization", "decrease"}:
                    source_label, _ = _table_cell_source_label(total_label, fact_section, metric, registry)
                if not source_label:
                    source_label = _normalize_template_label(total_label)
                fact, multiplier = _resolve_template_fact(registry, source_label)
                if not fact and source_label != fact_section:
                    source_label = fact_section
                    fact, multiplier = _resolve_template_fact(registry, source_label)
                value = (
                    fact.get(period) * multiplier
                    if fact and isinstance(fact.get(period), (int, float))
                    else None
                )
                # Prefer a direct statement fact, but keep a total local to
                # the detail rows in its own section.  The old global sum
                # could make a later grand total equal to every preceding
                # group total plus every detail row, or make one group total
                # absorb an unrelated schedule in the same table.
                previous_total_indices = [
                    index for index in computed_totals if index < total_row_index
                ]
                previous_total_index = max(previous_total_indices, default=header_depth - 1)
                detail_sum = sum(
                    values.get((column, period), 0.0)
                    for index, values in row_written.items()
                    if previous_total_index < index < total_row_index
                )
                prior_total_sum = sum(
                    computed_totals[index].get((column, period), 0.0)
                    for index in previous_total_indices
                )
                if "总计" in _normalize_template_label(total_label) and prior_total_sum:
                    summed_value = prior_total_sum + detail_sum
                else:
                    summed_value = detail_sum if detail_sum else prior_total_sum
                if value is None:
                    value = summed_value if summed_value else None
                elif (
                    isinstance(summed_value, (int, float))
                    and abs(float(value)) <= 1e-9
                    and abs(float(summed_value)) > 1e-9
                ):
                    # A detail schedule with a non-zero sourced sum is more
                    # reliable than a zero carried over from an unavailable
                    # comparative/formula cell in a high-level statement.
                    value = summed_value
                    stats.setdefault("total_fallbacks", []).append(
                        {"label": total_label, "period": period, "source": "sum_of_mapped_detail_rows"}
                    )
                elif (
                    stats["mapped_rows"]
                    and isinstance(summed_value, (int, float))
                    and abs(float(summed_value)) > 1e-9
                    and abs(float(value) - float(summed_value)) > 0.01
                ):
                    conflict = {
                        "label": total_label,
                        "period": period,
                        "statement_value": value,
                        "detail_sum": summed_value,
                    }
                    stats.setdefault("total_conflicts", []).append(conflict)
                    if not _is_independent_statement_total(total_label):
                        stats.setdefault("auto_reconciled_totals", []).append(conflict)
                        value = summed_value
                if (
                    not has_non_total_row_slot
                    and not stats["mapped_rows"]
                    and isinstance(value, (int, float))
                    and abs(float(value)) <= 1e-9
                ):
                    # A standalone all-zero total is an optional template
                    # schedule, not a disclosure requiring publication.
                    continue
                if _write_fact_to_cell(row.cells[column], value):
                    changed += 1
                    stats["filled_cells"] += 1
                    stats["mapped_rows"].append({"row": "total", "label": total_label, "fact": source_label or "sum_of_written_rows"})
                    computed_totals.setdefault(total_row_index, {})[(column, period)] = float(value)
        table_stats.append(stats)
    context["__docx_table_render_stats__"] = table_stats
    return changed


def _downgrade_formal_report_opinion(document: Any) -> int:
    """Prevent a reference unmodified-opinion template from becoming a false opinion."""

    changed = 0
    for paragraph in document.paragraphs:
        text = paragraph.text or ""
        if "审计意见" in text and text.strip() in {"一、审计意见", "审计意见"}:
            _set_docx_paragraph_text(paragraph, "一、审计意见（待项目组复核）")
            changed += 1
        elif "我们认为，后附的财务报表在所有重大方面" in text:
            _set_docx_paragraph_text(
                paragraph,
                "本文件为基于当前已导入资料形成的审计报告格式草稿，尚未完成全部审计程序，未形成正式审计意见，未经项目组复核和批准，不得作为正式审计报告对外使用。",
            )
            changed += 1
    return changed


_NOTES_OPTIONAL_SECTION_HINTS: dict[str, tuple[str, ...]] = {
    "合营安排分类及共同经营会计处理方法": ("合营", "联营", "C17", "长期股权投资"),
    "外币业务和外币报表折算": ("外币", "外币报表", "汇率"),
    "合同资产": ("合同资产", "合同负债"),
    "持有待售和终止经营": ("持有待售", "终止经营"),
    "长期股权投资": ("长期股权投资", "C17", "联营", "合营"),
    "投资性房地产": ("投资性房地产", "C18"),
    "在建工程": ("在建工程", "C22"),
    "借款费用": ("借款", "借款费用", "短期借款", "长期借款"),
    "生物资产": ("生物资产",),
    "油气资产": ("油气资产",),
    "无形资产": ("无形资产",),
    "开发支出": ("开发支出",),
    "商誉": ("商誉",),
    "长期待摊费用": ("长期待摊费用",),
    "职工薪酬": ("职工薪酬", "工资", "D8"),
    "政府补助": ("政府补助",),
    "租赁": ("租赁", "使用权资产", "租赁负债"),
    "股份支付": ("股份支付", "股权激励"),
    "递延所得税资产和递延所得税负债": ("递延所得税",),
    "分部报告": ("分部", "分部报告"),
}

# These are the sections that are structurally expected in the generic note
# template.  The template contains the disclosure text for many unrelated
# accounting situations; the generated note must select the sections that
# are supported by this engagement's evidence instead of emitting the whole
# reference manual.
_NOTES_MAJOR_KEEP = {
    "公司基本情况",
    "财务报表的编制基础",
    "重要会计政策及会计估计",
    "税项",
    "财务报表项目注释",
    "关联方及关联交易",
    "承诺及或有事项",
    "资产负债表日后事项",
    "其他重要事项",
}
_NOTES_POLICY_KEEP = {
    "遵循企业会计准则的声明",
    "会计期间",
    "营业周期",
    "记账本位币",
    "现金及现金等价物的确定标准",
    "金融工具",
    "应收款项",
    "固定资产",
    "长期待摊费用",
    "职工薪酬",
    "收入",
    "所得税",
    "重要会计政策和会计估计变更",
}
_NOTES_CASE_FINANCIAL_KEEP = {
    "货币资金",
    "交易性金融资产",
    "应收账款",
    "其他应收款",
    "其他流动资产",
    "固定资产",
    "长期待摊费用",
    "应付账款",
    "预收款项",
    "应付职工薪酬",
    "应交税费",
    "其他应付款",
    "实收资本（或股本）",
    "未分配利润",
    "营业收入和营业成本",
    "税金及附加",
    "销售费用",
    "管理费用",
    "财务费用",
    "其他收益",
    "投资收益",
    "营业外收入",
    "营业外支出",
    "所得税费用",
    "现金流量表项目",
    "现金流量表补充资料",
}
_NOTES_CASE_RELATED_KEEP = {
    "本公司的母公司情况",
    "关联交易情况",
    "关联方应收应付等未结算项目",
}
_NOTES_CASE_COMMITMENT_KEEP = {"重要承诺事项", "或有事项"}


def _remove_docx_body_range(document: Any, start: Any, stop: Any) -> None:
    body = document._element.body
    children = list(body)
    try:
        start_index = children.index(start)
    except ValueError:
        return
    stop_index = children.index(stop) if stop in children else len(children)
    for child in children[start_index:stop_index]:
        body.remove(child)


def _docx_heading_level(paragraph: Any) -> int | None:
    style_name = str(paragraph.style.name if paragraph.style else "")
    lowered = style_name.lower()
    # This template family uses both “附注标题 [n级]” and “附注正文 [n级]”.
    # The old check treated every body paragraph with a level suffix as a
    # heading and then deleted its text during section pruning.  A hierarchy
    # is a document-contract concept: require an explicit title style (or a
    # native Word Heading style), never the level suffix alone.
    is_heading = "标题" in style_name or lowered.startswith("heading ")
    if not is_heading:
        return None
    if "[1级]" in style_name or lowered == "heading 1":
        return 1
    if "[2级]" in style_name or lowered == "heading 2":
        return 2
    if "[3级]" in style_name or lowered == "heading 3":
        return 3
    return None


def _notes_material_corpus(context: dict[str, Any]) -> str:
    material_index = context.get("__case_material_index__") or {}
    # ``active_labels`` also contains headings from procedure/template sheets
    # in the 296-sheet workpaper pack.  Those headings are not evidence that
    # the entity has the corresponding balance.  Prefer the normalized trial
    # balance account names and a few high-confidence snapshot facts.
    parts = [str(item) for item in (material_index.get("trial_balance_account_names") or [])]
    if context.get("net_revenue"):
        parts.append("营业收入")
    if context.get("receivables_balance"):
        parts.append("应收账款")
    if context.get("total_inflow") or context.get("total_outflow"):
        parts.append("货币资金")
    profile = material_index.get("entity_profile") or {}
    if isinstance(profile, dict):
        # Business-scope wording is evidence for choosing an industry-specific
        # accounting-policy branch only when no more specific actual activity
        # is available.  A broad registered business scope commonly lists
        # trading/construction permissions that are not this year's operation.
        # It is deliberately not inferred from a template's example text.
        actual_activity = str(profile.get("main_activity") or "").strip()
        parts.append(actual_activity or str(profile.get("business_scope") or ""))
    return " ".join(parts)


def _notes_supported_financial_sections(context: dict[str, Any]) -> set[str]:
    """Return only note sections backed by a populated project fact.

    The old case-replay switch kept every section in a generic 200-page note
    template merely because the master workbook was present.  Presence of a
    workbook is not evidence that every accounting disclosure applies.  A
    section is kept only when the periodised fact registry contains a value for
    its standard financial-statement line or a declared group of line items.
    """

    registry = _fact_registry_from_context(context)

    def has_value(label: str) -> bool:
        fact, _ = _resolve_template_fact(registry, label)
        return bool(
            fact
            and any(
                isinstance(fact.get(period), (int, float)) and abs(float(fact.get(period) or 0)) > 1e-9
                for period in ("current", "opening")
            )
        )

    grouped: dict[str, tuple[str, ...]] = {
        "营业收入和营业成本": ("营业收入", "营业成本", "主营业务收入", "主营业务成本"),
        "现金流量表项目": ("经营活动产生的现金流量净额", "货币资金"),
        "现金流量表补充资料": ("净利润", "固定资产折旧", "长期待摊费用摊销", "货币资金"),
    }
    supported = {
        heading
        for heading in _NOTES_CASE_FINANCIAL_KEEP
        if any(has_value(item) for item in grouped.get(heading, (heading,)))
    }
    context["__notes_supported_sections__"] = sorted(supported)
    return supported


def _prune_generic_notes_sections(document: Any, context: dict[str, Any]) -> int:
    """Remove optional reference-template sections before filling values.

    A generic enterprise-note template contains hundreds of optional policy and
    disclosure pages.  Leaving them in the customer document creates a huge
    200-page reference manual full of blank tables.  Section selection is based
    on the complete workpaper material index; no case-output document is used as
    the content source.
    """

    corpus = _notes_material_corpus(context)
    body = document._element.body
    from docx.text.paragraph import Paragraph

    headings: list[tuple[Any, int, str]] = []
    for child in list(body):
        if not str(child.tag).endswith("}p"):
            continue
        paragraph = Paragraph(child, document)
        level = _docx_heading_level(paragraph)
        heading = (paragraph.text or "").strip()
        if level is not None and heading:
            headings.append((child, level, heading))

    current_major = ""
    current_subsection = ""
    heading_major: dict[int, str] = {}
    heading_subsection: dict[int, str] = {}
    for element, level, heading in headings:
        if level == 1:
            current_major = heading
            current_subsection = ""
        elif level == 2:
            current_subsection = heading
        heading_major[id(element)] = current_major
        heading_subsection[id(element)] = current_subsection

    case_replay = bool(context.get("__case_workpaper_replay__"))
    financial_keep = _notes_supported_financial_sections(context) if case_replay else set()
    related_keep: set[str] = set()
    commitment_keep = set(_NOTES_CASE_COMMITMENT_KEEP) if case_replay else set()
    removable: list[tuple[Any, str]] = []

    def should_remove(level: int, heading: str, major: str, subsection: str) -> bool:
        if level == 1:
            if major == "关联方及关联交易" and not related_keep:
                return True
            return major not in _NOTES_MAJOR_KEEP
        if major == "重要会计政策及会计估计":
            # Retain the policy's internal headings and prose once its
            # level-2 policy has been selected.  They are substantive
            # template content, not optional disclosure schedules.
            if level >= 3 and subsection in _NOTES_POLICY_KEEP:
                return False
            return heading not in _NOTES_POLICY_KEEP
        if major == "财务报表项目注释":
            if case_replay:
                if level >= 3 and subsection in financial_keep:
                    return False
                return heading not in financial_keep and heading != "其他说明"
            hints = _NOTES_OPTIONAL_SECTION_HINTS.get(heading)
            return bool(hints and not any(hint in corpus for hint in hints))
        if major == "关联方及关联交易":
            if level >= 3 and subsection in related_keep:
                return False
            return heading not in related_keep
        if major == "承诺及或有事项" and commitment_keep:
            return heading not in commitment_keep
        if case_replay and major in {"资产负债表日后事项", "其他重要事项"}:
            return True
        # A caller may render an extracted template fragment (for example a
        # single note table in a review preview).  Without a level-1 parent it
        # is not safe to classify that fragment as an optional disclosure.
        if not major:
            return False
        if heading in _NOTES_OPTIONAL_SECTION_HINTS:
            hints = _NOTES_OPTIONAL_SECTION_HINTS[heading]
            return not any(hint in corpus for hint in hints)
        return False

    # Work backwards so removing a parent section cannot invalidate the
    # location of an unprocessed preceding section.
    for element, level, heading in reversed(headings):
        if element.getparent() is None or not should_remove(
            level,
            heading,
            heading_major[id(element)],
            heading_subsection[id(element)],
        ):
            continue
        children = list(body)
        try:
            start_index = children.index(element)
        except ValueError:
            continue
        stop = None
        for next_element, next_level, _ in headings:
            if next_element is element or next_element not in children:
                continue
            try:
                next_index = children.index(next_element)
            except ValueError:
                continue
            if next_index > start_index and next_level <= level:
                stop = next_element
                break
        removable.append((element, heading))
        _remove_docx_body_range(document, element, stop)
    if removable:
        context["__notes_pruned_sections__"] = [heading for _, heading in removable]
    return len(removable)


def _add_docx_paragraph_after(paragraph: Any, text: str, style_name: str = "") -> None:
    from docx.oxml import OxmlElement
    from docx.text.paragraph import Paragraph

    element = OxmlElement("w:p")
    paragraph._p.addnext(element)
    inserted = Paragraph(element, paragraph._parent)
    if style_name:
        try:
            inserted.style = style_name
        except Exception:
            pass
    inserted.add_run(text)


def _fill_case_empty_disclosure_sections(document: Any, context: dict[str, Any]) -> int:
    """Replace optional blank reference subtrees with a traceable statement."""

    if not context.get("__case_workpaper_replay__"):
        return 0
    period_end = str(context.get("period_end") or "")
    statements = {
        "承诺及或有事项": f"截至{period_end}，主底稿索引中未识别到需单独披露的重大承诺及或有事项；仍需结合管理层声明和项目组复核记录确认。",
        "资产负债表日后事项": f"截至{period_end}，主底稿索引中未识别到需单独披露的重大资产负债表日后事项；仍需结合期后检查记录确认。",
        "其他重要事项": f"截至{period_end}，主底稿索引中未识别到需单独披露的其他重大事项。",
    }
    body = document._element.body
    from docx.text.paragraph import Paragraph

    headings = []
    for child in list(body):
        if not str(child.tag).endswith("}p"):
            continue
        paragraph = Paragraph(child, document)
        if _docx_heading_level(paragraph) == 1 and (paragraph.text or "").strip() in statements:
            headings.append(paragraph)

    changed = 0
    for heading in reversed(headings):
        children = list(body)
        try:
            start_index = children.index(heading._p)
        except ValueError:
            continue
        stop = None
        for child in children[start_index + 1:]:
            if not str(child.tag).endswith("}p"):
                continue
            candidate = Paragraph(child, document)
            if _docx_heading_level(candidate) == 1:
                stop = child
                break
        stop_index = children.index(stop) if stop is not None else len(children)
        for child in children[start_index + 1:stop_index]:
            body.remove(child)
        _add_docx_paragraph_after(
            heading,
            statements[(heading.text or "").strip()],
            "立信附注正文 [2级]",
        )
        changed += 1
    return changed


def _remove_generic_instruction_paragraphs(document: Any) -> int:
    removed = 0
    body = document._element.body
    for paragraph in list(document.paragraphs):
        text = (paragraph.text or "").strip()
        if not text:
            continue
        if any(marker in text for marker in _GENERIC_NOTE_INSTRUCTION_MARKERS):
            body.remove(paragraph._p)
            removed += 1
    # A bracketed option in a table label is a drafting instruction, not a
    # client fact.  Keep the table and replace it with the neutral selected
    # wording rather than publishing the template's choice prompt.
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if "[或适用]" in (cell.text or ""):
                    _set_docx_cell_text(cell, (cell.text or "").replace("[或适用]", "适用"))
                    removed += 1
    return removed


def _clear_generic_instruction_text(document: Any) -> int:
    """Clear note-template drafting instructions without deleting body nodes."""

    changed = 0
    for paragraph in document.paragraphs:
        text = (paragraph.text or "").strip()
        if text and any(marker in text for marker in _GENERIC_NOTE_INSTRUCTION_MARKERS):
            _set_docx_paragraph_text(paragraph, "")
            changed += 1
    for table in _iter_docx_tables(document):
        for row in table.rows:
            for cell in row.cells:
                text = cell.text or ""
                if any(marker in text for marker in _GENERIC_NOTE_INSTRUCTION_MARKERS):
                    _set_docx_cell_text(cell, "")
                    changed += 1
                elif "[或适用]" in text:
                    _set_docx_cell_text(cell, text.replace("[或适用]", "适用"))
                    changed += 1
    return changed


def _docx_run_authoring_annotations(run: Any) -> list[Any]:
    """Return direct run highlight/shading elements used as template cues."""

    from docx.oxml.ns import qn

    properties = run._r.rPr
    if properties is None:
        return []
    annotations: list[Any] = []
    for element in list(properties.findall(qn("w:highlight"))):
        value = str(element.get(qn("w:val")) or "").lower()
        if value not in {"", "none", "auto"}:
            annotations.append(element)
    for element in list(properties.findall(qn("w:shd"))):
        value = str(element.get(qn("w:fill")) or "").lower()
        if value not in {"", "auto", "ffffff", "none"}:
            annotations.append(element)
    return annotations


def _paragraph_has_docx_authoring_annotation(paragraph: Any) -> bool:
    return any(_docx_run_authoring_annotations(run) for run in paragraph.runs)


def _optional_business_policy_branch_name(text: str) -> str:
    """Return a compact highlighted industry-choice label's evidence key."""

    label = _normalize_template_label(text)
    if 2 < len(label) <= 18 and label.endswith("业务"):
        return label[:-2]
    # Templates also use labels such as “建筑施工业” and “制造业”.  These
    # are choices of an industry policy branch, not the generic word “营业”.
    if 3 < len(label) <= 18 and label.endswith("业"):
        return label[:-1]
    return ""


def _clear_docx_authoring_annotations(paragraph: Any) -> int:
    """Clear template authoring colour while retaining paragraph/run formatting."""

    changed = 0
    for run in paragraph.runs:
        for element in _docx_run_authoring_annotations(run):
            parent = element.getparent()
            if parent is not None:
                parent.remove(element)
                changed += 1
    return changed


def _clear_document_authoring_annotations(document: Any) -> int:
    """Clear authoring highlight/shading from every retained document slot."""

    changed = sum(_clear_docx_authoring_annotations(paragraph) for paragraph in document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                changed += sum(_clear_docx_authoring_annotations(paragraph) for paragraph in cell.paragraphs)
    return changed


def _normalise_generic_note_authoring_annotations(document: Any, context: dict[str, Any]) -> int:
    """Remove unsupported highlighted policy choices and clear authoring ink.

    Generic note templates commonly use yellow/grey run highlighting to flag
    selectable drafting alternatives.  Highlight colour is not audit evidence.
    An industry branch is kept only when the project business scope mentions
    the branch; all remaining retained text is returned to the template's
    normal typography rather than exposing its drafting instructions.
    """

    corpus = _normalize_template_label(_notes_material_corpus(context))
    body = document._element.body
    removed = 0
    remove_following_highlighted = False
    for paragraph in list(document.paragraphs):
        text = (paragraph.text or "").strip()
        highlighted = _paragraph_has_docx_authoring_annotation(paragraph)
        if highlighted and (branch := _optional_business_policy_branch_name(text)):
            if branch and branch not in corpus:
                body.remove(paragraph._p)
                removed += 1
                remove_following_highlighted = True
                continue
            remove_following_highlighted = False
        if remove_following_highlighted and highlighted:
            body.remove(paragraph._p)
            removed += 1
            continue
        # A non-highlighted body paragraph marks the end of the selected
        # branch.  Keep it available for the normal policy-selection rules.
        if not highlighted:
            remove_following_highlighted = False

    cleared = _clear_document_authoring_annotations(document)
    context["__notes_removed_authoring_annotation_count__"] = removed
    context["__notes_cleared_authoring_annotation_count__"] = cleared
    return removed + cleared


def _docx_authoring_annotation_count(document: Any) -> int:
    """Count residual direct highlight/shading in template body content."""

    count = sum(
        len(_docx_run_authoring_annotations(run))
        for paragraph in document.paragraphs
        for run in paragraph.runs
    )
    count += sum(
        len(_docx_run_authoring_annotations(run))
        for table in document.tables
        for row in table.rows
        for cell in row.cells
        for paragraph in cell.paragraphs
        for run in paragraph.runs
    )
    return count


def _fill_notes_entity_profile(document: Any, context: dict[str, Any]) -> int:
    """Fill the generic note's entity section from labelled workpaper facts."""

    material_index = context.get("__case_material_index__") or {}
    profile = material_index.get("entity_profile") or {}
    if not isinstance(profile, dict):
        return 0
    entity = str(context.get("entity_name") or profile.get("entity_masked") or "待补充被审计单位")
    address = str(profile.get("registered_address") or "")
    nature = str(profile.get("entity_nature") or "")
    capital = str(profile.get("registered_capital") or "")
    shareholder = str(profile.get("shareholder") or "")
    scope = str(profile.get("business_scope") or "")
    activity = str(profile.get("main_activity") or "").rstrip("。；; ")
    changed = 0
    for paragraph in list(document.paragraphs):
        text = (paragraph.text or "").strip()
        if not text:
            continue
        replacement = ""
        if "本公司实际从事的主要经营活动为" in text and activity:
            replacement = f"本公司实际从事的主要经营活动为：{activity}。"
        elif "本财务报表业经" in text:
            replacement = "本财务报表为基于已导入主底稿形成的审计工作草稿，批准报出日期及批准人待管理层确认。"
        elif "系于" in text and "营业执照" in text:
            details = [item for item in (nature, address, capital, shareholder) if item]
            suffix = "；".join(details)
            replacement = f"{entity}（以下简称“本公司”）的主体信息已从主底稿客户基本情况表提取。"
            if suffix:
                replacement += f"{suffix}。"
        elif "主要经营范围包括" in text and scope:
            replacement = f"公司的主要经营范围包括：{scope}"
        if replacement and replacement != text:
            _set_docx_paragraph_text(paragraph, replacement)
            changed += 1
    return changed


def _remove_empty_generic_tables(document: Any) -> int:
    removed = 0
    body = document._element.body
    for table in list(document.tables):
        if len(table.rows) <= 1:
            continue
        meaningful = False
        for row in table.rows[1:]:
            data_cells = row.cells[1:] if len(row.cells) > 1 else row.cells
            if any(
                (cell.text or "").strip()
                and not re.fullmatch(r"[Xx＊*…\.。\s]+", (cell.text or "").strip())
                for cell in data_cells
            ):
                meaningful = True
                break
        if not meaningful:
            body.remove(table._tbl)
            removed += 1
    return removed


def _remove_unmapped_note_tables(document: Any, context: dict[str, Any]) -> int:
    """Remove optional disclosure grids that have no mapped project facts.

    A generic note template is a catalogue of possible disclosures.  Keeping a
    period/rate table solely because it existed in that catalogue produces an
    empty customer-facing schedule.  This is deliberately limited to note
    data grids; policy tables with pre-approved accounting-policy content are
    left untouched.
    """

    stats = {
        int(item.get("table_index")): item
        for item in (context.get("__docx_table_render_stats__") or [])
        if isinstance(item, dict)
    }
    removed = 0
    for index, table in reversed(list(enumerate(document.tables))):
        stat = stats.get(index) or {}
        if int(stat.get("filled_cells") or 0) > 0:
            continue
        header = " ".join(
            cell.text or ""
            for row in table.rows[:3]
            for cell in row.cells
        )
        data_rows = table.rows[min(int(stat.get("header_rows") or 1), len(table.rows)):]
        # A genuine policy matrix already carries its approved content in the
        # template body.  A disclosure grid has labels in the first column but
        # no populated data cell after the header.  Remove only the latter.
        # Descriptive samples such as “例如：收到发票后 30 天” are part of
        # a generic template, not evidence for this engagement.  Retain an
        # otherwise-unmapped grid only when it has an actual numeric amount
        # that was not a template default cleared by the renderer.
        amount_pattern = re.compile(r"[-+]?\d{1,3}(?:,\d{3})*(?:\.\d+)?")
        has_existing_data = any(
            "例如" not in (cell.text or "") and amount_pattern.search(cell.text or "")
            for row in data_rows
            for cell in row.cells[1:]
        )
        is_data_grid = any(
            token in header
            for token in (
                "期末", "期初", "本期", "上期", "本年", "上年", "税率",
                "单位名称", "借款单位", "项目", "类别", "账龄", "名称",
            )
        )
        if not is_data_grid:
            continue
        # A tax-rate grid is incomplete until rates are sourced, even though
        # its description column contains generic explanatory prose.
        if "税率" not in header and has_existing_data:
            continue
        parent = table._tbl.getparent()
        if parent is not None:
            parent.remove(table._tbl)
            removed += 1
    context["__notes_removed_unmapped_table_count__"] = removed
    return removed


def _note_range_has_audited_content(document: Any, elements: list[Any]) -> bool:
    """Return whether a note subtree still carries an audited value.

    Tables in a note template often retain labels after the data grid has been
    removed.  A heading followed only by those labels is not a disclosure.  We
    deliberately use a narrow test here: a populated amount / percentage in a
    remaining table, or a numeric fact in a narrative paragraph.  Headings
    alone, generic instructions and placeholders never keep a subtree alive.
    """

    from docx.table import Table
    from docx.text.paragraph import Paragraph

    # Paragraph labels such as “（1）” are not evidence.  Rendered monetary
    # facts carry either a thousands separator or a decimal part, while a
    # bare section-number does not.
    numeric = re.compile(r"[-+]?(?:\d{1,3}(?:,\d{3})+(?:\.\d+)?|\d+\.\d+)(?:%|％)?")
    for element in elements:
        if str(element.tag).endswith("}tbl"):
            table = Table(element, document)
            for row in table.rows[1:]:
                cells = row.cells[1:] if len(row.cells) > 1 else row.cells
                if any(numeric.search(cell.text or "") for cell in cells):
                    return True
        elif str(element.tag).endswith("}p"):
            paragraph = Paragraph(element, document)
            if _docx_heading_level(paragraph) is None and numeric.search(paragraph.text or ""):
                return True
    return False


def _remove_empty_financial_note_subsections(document: Any, context: dict[str, Any]) -> int:
    """Remove orphaned note headings after unmapped disclosure grids are cut.

    This is intentionally driven by the template hierarchy, not by a list of
    case-specific headings.  It makes a new generic notes template compact
    whenever one of its optional schedules has no evidence for this audit,
    while retaining every subtree that contains a rendered fact.
    """

    body = document._element.body
    from docx.text.paragraph import Paragraph

    headings: list[tuple[Any, int, str]] = []
    major = ""
    for child in list(body):
        if not str(child.tag).endswith("}p"):
            continue
        paragraph = Paragraph(child, document)
        level = _docx_heading_level(paragraph)
        text = (paragraph.text or "").strip()
        if level is None or not text:
            continue
        if level == 1:
            major = text
        if major == "财务报表项目注释" and level >= 2:
            headings.append((child, level, text))

    removed = 0
    # Process deepest/later headings first; this makes an empty level-2 note
    # disappear after all of its empty level-3 schedules have disappeared.
    for element, level, _ in reversed(headings):
        children = list(body)
        if element not in children:
            continue
        start = children.index(element)
        stop_index = len(children)
        for candidate, candidate_level, _ in headings:
            if candidate not in children:
                continue
            candidate_index = children.index(candidate)
            if candidate_index > start and candidate_level <= level:
                stop_index = candidate_index
                break
        contents = children[start + 1:stop_index]
        if _note_range_has_audited_content(document, contents):
            continue
        for child in contents:
            body.remove(child)
        body.remove(element)
        removed += 1
    context["__notes_removed_empty_subsection_count__"] = removed
    return removed


def _remove_unresolved_generic_note_components(document: Any, context: dict[str, Any]) -> int:
    """Remove optional generic-note components that still contain placeholders.

    A surviving ``XX`` in a notes template is a conditional disclosure (tax
    incentives, segment reporting, contract-performance obligations, etc.)
    for which the project supplied no evidence.  It must not become a fake
    sentence such as “主底稿未提供元”; remove the optional component and retain
    the removal count in the package manifest.  Validation still blocks any
    marker that escapes this controlled clean-up.
    """

    marker = re.compile(r"X{2,}|x{2,}|20XX|\{\{[^}]+\}\}|\[\[[^\]]+\]\]")
    removed = 0
    body = document._element.body
    for paragraph in list(document.paragraphs):
        if marker.search(paragraph.text or ""):
            body.remove(paragraph._p)
            removed += 1
    for table in list(document.tables):
        table_text = "\n".join(cell.text or "" for row in table.rows for cell in row.cells)
        if marker.search(table_text):
            parent = table._tbl.getparent()
            if parent is not None:
                parent.remove(table._tbl)
                removed += 1
    context["__notes_removed_placeholder_component_count__"] = removed
    return removed


def _render_docx(data: bytes, context: dict[str, Any]) -> tuple[bytes, str]:
    from docx import Document
    from .attachment_document_composer import compose_attachment_from_context

    strict_client_delivery = bool(context.get("__strict_client_delivery__"))
    if strict_client_delivery:
        # Bind the contract against this exact source package before any
        # python-docx save operation.  A template revision cannot silently
        # reuse an old paragraph address or source-text fingerprint.
        from .attachment_blueprint_service import bind_docx_slot_contract

        existing_contract = context.get("__attachment_slot_contract__")
        if isinstance(existing_contract, dict) and existing_contract.get("declared"):
            schema_for_contract: dict[str, Any] = {
                "docx": {"slot_contract": existing_contract}
            }
        else:
            schema_for_contract = context.get("__field_schema__") or {}
        contract = bind_docx_slot_contract(
            data,
            field_schema=schema_for_contract,
            file_name=str(context.get("__source_file_name__") or ""),
        )
        if not contract.get("declared") or not contract.get("slots"):
            raise ValueError("DOCX 正式交付必须配置经确认的 slot_contract")
        context["__attachment_slot_contract__"] = contract
        context["__source_template_sha256__"] = hashlib.sha256(data).hexdigest()

    _prepare_template_render_context(
        context,
        file_name=str(context.get("__source_file_name__") or "template.docx"),
        extension=".docx",
    )
    composed_data = compose_attachment_from_context(
        data,
        str(context.get("__source_file_name__") or "template.docx"),
        context,
    )
    composition_manifest = context.get("__attachment_composition_manifest__") or {}
    ai_authored_changes = (
        int(composition_manifest.get("authored_block_count") or 0)
        + int(bool(composition_manifest.get("title_written")))
    )
    document = Document(BytesIO(composed_data))
    original_document_text = _docx_document_text(document)
    has_result_token = _template_contains_result_token(original_document_text)
    source_name = str(context.get("__source_file_name__") or "")
    # Removing generic-note sections rewrites the client's body and table
    # sequence, so it is no longer part of the normal customer-facing path.
    # It remains opt-in only for one-off legacy migration jobs that explicitly
    # acknowledge that they are not template-fidelity renders.
    is_generic_notes = bool(context.get("__legacy_generic_note_pruning__"))
    if strict_client_delivery and is_generic_notes:
        raise ValueError("DOCX 正式交付不允许启用遗留附注裁剪")
    if is_generic_notes:
        context["__template_render_contract__"]["content_profile"] = "evidence_selected_generic_note"
        _prune_generic_notes_sections(document, context)
        _fill_case_empty_disclosure_sections(document, context)
        _remove_generic_instruction_paragraphs(document)
    schema_changed = _apply_docx_field_schema(document, context)
    declared_semantic_coverage = (
        _declared_docx_semantic_coverage(
            context,
            template_type=str(context.get("__template_type__") or ""),
        )
        if strict_client_delivery
        else {"mapped": False, "reason": "not_declared_contract", "structured_value_count": 0, "ai_value_count": 0, "slot_count": 0}
    )
    legacy_marker_changes = (
        0 if strict_client_delivery else _replace_docx_container(document, context)
    )
    changed = ai_authored_changes + schema_changed + legacy_marker_changes
    if context.get("__clear_generic_note_instruction_slots__") and not strict_client_delivery:
        changed += _clear_generic_instruction_text(document)
    if is_generic_notes:
        changed += _fill_notes_entity_profile(document, context)
        changed += _normalise_generic_note_authoring_annotations(document, context)
    # A complete case-replay notes reference is itself the authoritative
    # 26-page case output.  Re-running the generic trial-balance table mapper
    # can alter values and pagination (for example, turning 26 pages into 27),
    # so preserve its tables and only apply explicit metadata replacements.
    if not strict_client_delivery:
        changed += _fill_docx_tables(document, context)
    if is_generic_notes:
        # An optional data grid with no fact mapping is not a reviewed blank
        # field; it is a generic-template disclosure that does not apply to
        # this engagement.  Remove it as a whole, preserving the layout of
        # every table that did receive auditable project values.
        changed += _remove_unmapped_note_tables(document, context)
        changed += _remove_unresolved_generic_note_components(document, context)
        changed += _remove_empty_financial_note_subsections(document, context)
    if context.get("__template_type__") == "notes" and is_generic_notes:
        # Legacy generic-note cleanup can still use this narrowly scoped
        # presentation repair.  Slot-based templates retain their original
        # table borders exactly as supplied.
        changed += _ensure_docx_table_outer_borders(document, context)
    if context.get("__template_type__") == "annual_report" and not strict_client_delivery:
        opinion_changed = _downgrade_formal_report_opinion(document)
        changed += opinion_changed
    else:
        opinion_changed = 0
    # Headers and footers are part of the template frame.  Accessing an
    # inherited header/footer through python-docx can materialise a new part,
    # so they deliberately remain read-only in a fidelity render.  Likewise,
    # untouched body highlights and shading stay intact; replacement helpers
    # clear authoring ink only on slots they actually write.
    residual_authoring_annotations = 0
    if is_generic_notes:
        residual_authoring_annotations = _clear_document_authoring_annotations(document)
        changed += residual_authoring_annotations
    total_authoring_annotations_cleared = (
        int(context.get("__notes_cleared_authoring_annotation_count__") or 0)
        + residual_authoring_annotations
    )
    context["__cleared_template_authoring_annotation_count__"] = total_authoring_annotations_cleared
    # Authored text is written only into declared template slots.  No generic
    # chat appendix, evidence trace or reconstructed document body is added.
    output = BytesIO()
    document.save(output)
    status = "filled" if changed else "copied_no_matching_placeholders"
    table_stats = context.get("__docx_table_render_stats__") or []
    mapped_value_cell_count = sum(
        int(item.get("filled_cells") or 0) for item in table_stats if isinstance(item, dict)
    )
    result_token_replaced = has_result_token and not _template_contains_result_token(_docx_document_text(document))
    has_schema_result_mapping = _schema_contains_result_mapping(context.get("__field_schema__"))
    declared_contract_result_mapping = bool(
        strict_client_delivery
        and context.get("__attachment_slot_contract__")
        and declared_semantic_coverage.get("mapped")
    )
    audit_result_mapped = bool(
        ai_authored_changes
        or result_token_replaced
        or opinion_changed
        or mapped_value_cell_count
        or declared_contract_result_mapping
        or has_schema_result_mapping and schema_changed
    )
    contract = context.get("__template_render_contract__") or {}
    contract["audit_result_mapped"] = audit_result_mapped
    contract["declared_slot_semantic_coverage"] = declared_semantic_coverage
    contract["audit_result_mapping_reason"] = (
        str(declared_semantic_coverage.get("reason") or "declared_slot_contract")
        if declared_contract_result_mapping else
        "template_slot_authoring" if ai_authored_changes else
        "explicit_result_token" if result_token_replaced else
        "formal_opinion_downgrade" if opinion_changed else
        "field_schema" if has_schema_result_mapping and schema_changed else
        "structured_table_values" if mapped_value_cell_count else
        "none"
    )
    if context.get("__enforce_result_mapping__") and not audit_result_mapped:
        status = "metadata_filled_no_result_mapping" if changed else "copied_no_matching_placeholders"
    _record_template_render_result(
        context,
        renderer="docx",
        changed=changed,
        status=status,
        discovered_table_count=len(table_stats),
        output_table_count=len(document.tables),
        mapped_table_count=sum(
            1 for item in table_stats if isinstance(item, dict) and int(item.get("filled_cells") or 0) > 0
        ),
        mapped_value_cell_count=mapped_value_cell_count,
        docx_field_schema_slot_count=schema_changed,
        ai_authored_block_count=max(
            ai_authored_changes,
            int(declared_semantic_coverage.get("ai_value_count") or 0),
        ),
        source_business_body_reused=composition_manifest.get("source_business_body_reused"),
        audit_result_mapped=audit_result_mapped,
        audit_result_mapping_reason=contract.get("audit_result_mapping_reason"),
        notes_outer_border_table_count=int(context.get("__notes_outer_border_table_count__") or 0),
        notes_outer_border_side_count=int(context.get("__notes_outer_border_side_count__") or 0),
    )
    return output.getvalue(), status


def _xlsx_header_period_metric(value: Any) -> tuple[str, str]:
    """Classify an uploaded workbook header using the DOCX table semantics."""

    header = str(value or "").strip().replace(" ", "")
    if not header:
        return "", ""
    if any(token in header for token in ("期末", "本期", "本年")):
        period = "current"
    elif any(token in header for token in ("期初", "上期", "上年")):
        period = "opening"
    else:
        return "", ""
    if "收入" in header:
        metric = "收入"
    elif "成本" in header:
        metric = "成本"
    elif "本期增加" in header:
        metric = "increase"
    elif "摊销" in header:
        metric = "amortization"
    elif "减少" in header:
        metric = "decrease"
    elif "计提比例" in header:
        metric = "provision_ratio"
    elif "比例" in header:
        metric = "ratio"
    elif "坏账准备" in header:
        metric = "provision"
    elif "账面价值" in header:
        metric = "net"
    elif "账面余额" in header or "金额" in header or "余额" in header:
        metric = "amount"
    else:
        metric = ""
    return period, metric


def _xlsx_header_score(value: Any) -> int:
    """Return whether a cell looks like an uploaded financial-table header."""

    text = str(value or "").strip().replace(" ", "")
    if not text:
        return 0
    if text in {"项目", "类别", "名称", "账龄", "税种", "补充资料", "金额", "比例（%）", "计提比例（%）", "行次"}:
        return 1
    if _xlsx_header_period_metric(text)[0]:
        return 1
    if any(token in text for token in ("账面余额", "账面价值", "坏账准备", "计提比例")):
        return 1
    if any(token in text for token in ("第一阶段", "第二阶段", "第三阶段", "未来12个月", "整个存续期")):
        return 1
    return 0


def _xlsx_header_text(sheet: Any, row: int, column: int, depth: int) -> str:
    """Join a vertical multi-row header without modifying merged cells."""

    values: list[str] = []
    for item_row in range(row, min(sheet.max_row or 0, row + depth - 1) + 1):
        value = sheet.cell(row=item_row, column=column).value
        if value not in (None, ""):
            values.append(str(value).strip())
    return " ".join(values)


def _xlsx_workbook_table_plans(workbook: Any) -> list[dict[str, Any]]:
    """Discover data grids from headers rather than sheet or file names.

    The plan intentionally keeps the same three concepts as the DOCX mapper:
    a label column, current/comparative amount columns, and existing row
    slots.  A sheet may contain multiple tables; each receives its own plan.
    """

    plans: list[dict[str, Any]] = []
    label_headers = ("项目", "科目", "名称", "类别", "账龄", "行次", "资产", "负债", "所有者权益")
    for sheet in workbook.worksheets:
        max_row = int(sheet.max_row or 0)
        max_column = int(sheet.max_column or 0)
        if max_row < 2 or max_column < 2:
            continue
        row = 1
        while row <= max_row:
            first_row_scores = [
                _xlsx_header_score(sheet.cell(row=row, column=column).value)
                for column in range(1, max_column + 1)
            ]
            if sum(first_row_scores) < 2:
                row += 1
                continue
            header_depth = 1
            for candidate_row in range(row + 1, min(max_row, row + 2) + 1):
                scores = [
                    _xlsx_header_score(sheet.cell(row=candidate_row, column=column).value)
                    for column in range(1, max_column + 1)
                ]
                if sum(scores) >= 2:
                    header_depth += 1
                else:
                    break
            columns: dict[int, tuple[str, str]] = {}
            label_column: int | None = None
            for column in range(1, max_column + 1):
                header = _xlsx_header_text(sheet, row, column, header_depth)
                period, metric = _xlsx_header_period_metric(header)
                if period:
                    columns[column] = (period, metric)
                if label_column is None and any(marker in header.replace(" ", "") for marker in label_headers):
                    label_column = column
            if columns and label_column is not None:
                plans.append(
                    {
                        "sheet": sheet.title,
                        "header_row": row,
                        "header_rows": header_depth,
                        "label_column": label_column,
                        "value_columns": columns,
                    }
                )
                row += header_depth
            else:
                row += 1
    return plans


def _xlsx_value_source_label(row_label: str, metric: str) -> str:
    """Resolve a workbook row by accounting semantics, never a substring."""

    label = _normalize_template_label(row_label)
    if metric == "increase":
        return f"{label}本期增加"
    if metric == "amortization":
        return f"{label}摊销"
    if metric == "decrease":
        return f"{label}其他减少"
    if metric == "provision":
        return f"{label}坏账准备"
    # Ratio grids need a table-level denominator.  A workbook template must
    # declare that denominator through a field schema; guessing one from an
    # arbitrary sheet would produce a plausible but false audit percentage.
    if metric in {"ratio", "provision_ratio"}:
        return ""
    return label


def _is_literal_template_value(value: Any) -> bool:
    """Identify a sample numeric value that must not survive a new render."""

    if isinstance(value, bool) or value is None:
        return False
    if isinstance(value, (int, float)):
        return True
    text = str(value).strip().replace(",", "").replace("，", "")
    text = text.rstrip("%％")
    if not text:
        return False
    try:
        float(text)
    except (TypeError, ValueError):
        return False
    return True


def _prepare_financial_semantic_authoring(
    data: bytes,
    extension: str,
    context: dict[str, Any],
) -> dict[str, Any]:
    """Ask AI to bind cell semantics to fact IDs without sending amounts."""

    from openpyxl import load_workbook
    from .attachment_financial_author_service import (
        author_financial_semantic_plan,
        fact_id_for_label,
    )

    workbook = load_workbook(
        BytesIO(data),
        read_only=False,
        data_only=False,
        keep_vba=extension == ".xlsm",
    )
    registry = _fact_registry_from_context(context)
    fact_catalog: list[dict[str, Any]] = []
    fact_key_by_id: dict[str, str] = {}
    fact_id_by_key: dict[str, str] = {}
    fact_key_by_object_id = {id(fact): key for key, fact in registry.items()}
    for key, fact in registry.items():
        if not isinstance(fact, dict):
            continue
        fact_id = fact_id_for_label(key)
        periods = [
            period
            for period in ("current", "opening")
            if isinstance(fact.get(period), (int, float)) and not isinstance(fact.get(period), bool)
        ]
        if not periods:
            continue
        fact_catalog.append(
            {
                "fact_id": fact_id,
                "label": str(fact.get("label") or key),
                "available_periods": periods,
                "source_kind": str(fact.get("source") or "frozen_fact_registry"),
            }
        )
        fact_key_by_id[fact_id] = key
        fact_id_by_key[key] = fact_id

    candidate_cells: list[dict[str, Any]] = []
    candidate_multipliers: dict[str, int] = {}
    for plan in _xlsx_workbook_table_plans(workbook):
        sheet = workbook[plan["sheet"]]
        header_end = int(plan["header_row"]) + int(plan["header_rows"]) - 1
        empty_streak = 0
        for row in range(header_end + 1, int(sheet.max_row or 0) + 1):
            label = str(sheet.cell(row=row, column=int(plan["label_column"])).value or "").strip()
            if not label:
                empty_streak += 1
                if empty_streak >= 3:
                    break
                continue
            empty_streak = 0
            if _is_total_template_label(label):
                continue
            for column, (period, metric) in plan["value_columns"].items():
                target = sheet.cell(row=row, column=column)
                if target.__class__.__name__ == "MergedCell" or (
                    isinstance(target.value, str) and target.value.startswith("=")
                ):
                    continue
                row_label = _normalize_template_label(label)
                source_label = _xlsx_value_source_label(row_label, metric)
                resolved_fact, multiplier = _resolve_template_fact(
                    registry,
                    source_label,
                )
                resolved_key = (
                    fact_key_by_object_id.get(id(resolved_fact))
                    if resolved_fact is not None
                    else ""
                )
                allowed_fact_id = fact_id_by_key.get(str(resolved_key or ""), "")
                cell_id = f"{sheet.title}!{target.coordinate}|{period}|{metric}"
                candidate_cells.append(
                    {
                        "cell_id": cell_id,
                        "sheet": sheet.title,
                        "coordinate": target.coordinate,
                        "row_label": row_label,
                        "period": period,
                        "metric": metric or "balance",
                        "allowed_fact_ids": [allowed_fact_id] if allowed_fact_id else [],
                        "multiplier": int(multiplier),
                    }
                )
                if allowed_fact_id:
                    candidate_multipliers[cell_id] = int(multiplier)
    workbook.close()
    plan = author_financial_semantic_plan(
        candidate_cells=candidate_cells,
        fact_catalog=fact_catalog,
    )
    context["__financial_semantic_candidate_ids__"] = [item["cell_id"] for item in candidate_cells]
    context["__financial_semantic_bindings__"] = dict(plan.get("bindings") or {})
    context["__financial_fact_key_by_id__"] = fact_key_by_id
    context["__financial_semantic_multipliers__"] = {
        cell_id: candidate_multipliers.get(cell_id, 1)
        for cell_id in (plan.get("bindings") or {})
    }
    context["__financial_semantic_manifest__"] = {
        key: value
        for key, value in plan.items()
        if key not in {"bindings"}
    }
    return plan


def _fill_xlsx_tables(workbook: Any, context: dict[str, Any]) -> int:
    """Fill discovered XLSX value cells from the shared fact registry.

    Unlike the retired implementation, this routine does not scan arbitrary
    labels for substrings and does not create a new result sheet.  It writes
    only into source-template rows whose labels resolve to a project fact.
    Existing number formats, merged cells, formulas, print setup and sheet
    order remain intact.
    """

    registry = _fact_registry_from_context(context)
    semantic_candidate_ids = set(context.get("__financial_semantic_candidate_ids__") or [])
    semantic_bindings = context.get("__financial_semantic_bindings__") or {}
    fact_key_by_id = context.get("__financial_fact_key_by_id__") or {}
    semantic_multipliers = context.get("__financial_semantic_multipliers__") or {}
    plans = _xlsx_workbook_table_plans(workbook)
    changed = 0
    stats: list[dict[str, Any]] = []
    for plan in plans:
        sheet = workbook[plan["sheet"]]
        header_end = int(plan["header_row"]) + int(plan["header_rows"]) - 1
        mapped_rows = 0
        filled_cells = 0
        mapped_value_cells = 0
        cleared_template_default_cells = 0
        row_written: dict[int, dict[tuple[int, str], float]] = {}
        total_rows: list[tuple[int, str]] = []
        computed_totals: dict[int, dict[tuple[int, str], float]] = {}
        total_fallbacks: list[dict[str, Any]] = []
        total_conflicts: list[dict[str, Any]] = []
        auto_reconciled_totals: list[dict[str, Any]] = []
        data_rows: list[int] = []
        empty_streak = 0
        for row in range(header_end + 1, int(sheet.max_row or 0) + 1):
            label = str(sheet.cell(row=row, column=int(plan["label_column"])).value or "").strip()
            if not label:
                empty_streak += 1
                if empty_streak >= 3:
                    break
                continue
            empty_streak = 0
            data_rows.append(row)

        # Literal amounts in a supplied template are examples, not audit
        # facts.  Clear them in recognised value columns before mapping the
        # current engagement; formulas and merged cells remain untouched.
        for row in data_rows:
            for column in plan["value_columns"]:
                target = sheet.cell(row=row, column=column)
                if target.__class__.__name__ == "MergedCell":
                    continue
                if isinstance(target.value, str) and target.value.startswith("="):
                    continue
                if _is_literal_template_value(target.value):
                    target.value = None
                    changed += 1
                    cleared_template_default_cells += 1
        for row in data_rows:
            label_cell = sheet.cell(row=row, column=int(plan["label_column"]))
            label = str(label_cell.value or "").strip()
            clean_label = _normalize_template_label(label)
            if clean_label in {"项目", "科目", "名称", "类别", "账龄", "行次"}:
                continue
            if _is_total_template_label(label):
                total_rows.append((row, label))
                continue
            row_filled = False
            for column, (period, metric) in plan["value_columns"].items():
                source_label = _xlsx_value_source_label(clean_label, metric)
                target = sheet.cell(row=row, column=column)
                cell_id = f"{sheet.title}!{target.coordinate}|{period}|{metric}"
                bound_fact_id = str(semantic_bindings.get(cell_id) or "")
                if bound_fact_id:
                    fact = registry.get(str(fact_key_by_id.get(bound_fact_id) or ""))
                    multiplier = int(semantic_multipliers.get(cell_id, 1) or 1)
                elif cell_id in semantic_candidate_ids:
                    fact, multiplier = None, 1
                else:
                    fact, multiplier = _resolve_template_fact(registry, source_label)
                if not fact:
                    continue
                fact_period = _table_fact_period_for_row(label, period)
                value = fact.get(fact_period) if fact_period else None
                if not isinstance(value, (int, float)):
                    continue
                if target.__class__.__name__ == "MergedCell" or (
                    isinstance(target.value, str) and target.value.startswith("=")
                ):
                    continue
                rendered_value = float(value) * multiplier
                mapped_value_cells += 1
                if target.value != rendered_value:
                    target.value = rendered_value
                    changed += 1
                    filled_cells += 1
                    row_filled = True
                row_written.setdefault(row, {})[(column, period)] = rendered_value
            if row_filled:
                mapped_rows += 1

        # Totals are resolved from a direct audited fact first.  If that fact
        # is unavailable, calculate only from rows since the previous total;
        # a grand total may additionally include the already calculated group
        # totals.  This keeps independent schedules in one sheet isolated.
        for total_row, total_label in total_rows:
            clean_label = _normalize_template_label(total_label)
            row_filled = False
            for column, (period, metric) in plan["value_columns"].items():
                source_label = _xlsx_value_source_label(clean_label, metric)
                fact, multiplier = _resolve_template_fact(registry, source_label)
                fact_period = _table_fact_period_for_row(total_label, period)
                value = fact.get(fact_period) * multiplier if fact and fact_period and isinstance(fact.get(fact_period), (int, float)) else None

                previous_total_indices = [index for index in computed_totals if index < total_row]
                previous_total_index = max(previous_total_indices, default=header_end)
                detail_sum = sum(
                    values.get((column, period), 0.0)
                    for index, values in row_written.items()
                    if previous_total_index < index < total_row
                )
                prior_total_sum = sum(
                    computed_totals[index].get((column, period), 0.0)
                    for index in previous_total_indices
                )
                summed_value = (
                    prior_total_sum + detail_sum
                    if "总计" in clean_label and prior_total_sum
                    else detail_sum if detail_sum else prior_total_sum
                )
                if value is None:
                    value = summed_value if summed_value else None
                    if value is not None:
                        stats_entry = {"label": total_label, "period": period, "source": "sum_of_mapped_detail_rows"}
                    else:
                        stats_entry = None
                elif summed_value and abs(float(value)) <= 1e-9 and abs(float(summed_value)) > 1e-9:
                    value = summed_value
                    stats_entry = {"label": total_label, "period": period, "source": "sum_of_mapped_detail_rows"}
                elif summed_value and abs(float(value) - float(summed_value)) > 0.01:
                    conflict = {
                        "label": total_label,
                        "period": period,
                        "statement_value": value,
                        "detail_sum": summed_value,
                    }
                    total_conflicts.append(conflict)
                    if not _is_independent_statement_total(total_label):
                        auto_reconciled_totals.append(conflict)
                        value = summed_value
                    stats_entry = None
                else:
                    stats_entry = None
                if stats_entry:
                    total_fallbacks.append(stats_entry)
                if not isinstance(value, (int, float)):
                    continue
                target = sheet.cell(row=total_row, column=column)
                if target.__class__.__name__ == "MergedCell" or (
                    isinstance(target.value, str) and target.value.startswith("=")
                ):
                    continue
                mapped_value_cells += 1
                if target.value != value:
                    target.value = float(value)
                    changed += 1
                    filled_cells += 1
                    row_filled = True
                computed_totals.setdefault(total_row, {})[(column, period)] = float(value)
            if row_filled:
                mapped_rows += 1
        stats.append(
            {
                "sheet": plan["sheet"],
                "header_row": plan["header_row"],
                "mapped_rows": mapped_rows,
                "filled_cells": filled_cells,
                "mapped_value_cells": mapped_value_cells,
                "cleared_template_default_cells": cleared_template_default_cells,
                "total_fallbacks": total_fallbacks,
                "total_conflicts": total_conflicts,
                "auto_reconciled_totals": auto_reconciled_totals,
            }
        )
    context["__xlsx_table_render_stats__"] = stats
    return changed


def _render_xlsx(data: bytes, extension: str, context: dict[str, Any]) -> tuple[bytes, str]:
    """Render an XLSX/XLSM template through the shared structure contract."""

    from openpyxl import load_workbook

    _prepare_template_render_context(
        context,
        file_name=str(context.get("__source_file_name__") or f"template{extension}"),
        extension=extension,
    )
    workbook = load_workbook(BytesIO(data), keep_vba=extension == ".xlsm")
    changed = 0
    template_text = "\n".join(
        str(cell.value or "")
        for sheet in workbook.worksheets
        for row in sheet.iter_rows()
        for cell in row
        if isinstance(cell.value, str)
    )
    has_result_token = _template_contains_result_token(template_text)
    changed += _apply_xlsx_field_schema(workbook, context)
    for sheet in workbook.worksheets:
        for row in sheet.iter_rows():
            for cell in row:
                if isinstance(cell.value, str) and not cell.value.startswith("="):
                    replaced = _replace_template_markers(cell.value, context)
                    if replaced != cell.value:
                        cell.value = replaced
                        changed += 1
    changed += _fill_xlsx_tables(workbook, context)

    # Keep the workbook's original sheets and layout.  A result sheet made by
    # the renderer would be a second, unrelated report rather than a filled
    # version of the supplied template.
    output = BytesIO()
    workbook.save(output)
    semantic_manifest = context.get("__financial_semantic_manifest__") or {}
    semantic_planned = bool(
        semantic_manifest and int(semantic_manifest.get("bound_count") or 0) > 0
    )
    status = "filled" if changed or semantic_planned else "copied_no_matching_placeholders"
    has_schema_result_mapping = _schema_contains_result_mapping(context.get("__field_schema__"))
    table_stats = context.get("__xlsx_table_render_stats__") or []
    mapped_value_cell_count = sum(
        int(item.get("mapped_value_cells") or item.get("filled_cells") or 0)
        for item in table_stats
        if isinstance(item, dict)
    )
    result_token_replaced = has_result_token and not _template_contains_result_token(
        "\n".join(
            str(cell.value or "")
            for sheet in workbook.worksheets
            for row in sheet.iter_rows()
            for cell in row
            if isinstance(cell.value, str)
        )
    )
    audit_result_mapped = bool(
        semantic_planned
        or result_token_replaced
        or has_schema_result_mapping and changed
        or mapped_value_cell_count
    )
    contract = context.get("__template_render_contract__") or {}
    contract["audit_result_mapped"] = audit_result_mapped
    contract["audit_result_mapping_reason"] = (
        "ai_semantic_fact_plan" if semantic_planned else
        "explicit_result_token" if result_token_replaced else
        "field_schema" if has_schema_result_mapping and changed else
        "structured_table_values" if mapped_value_cell_count else
        "none"
    )
    if context.get("__enforce_result_mapping__") and not audit_result_mapped:
        status = "metadata_filled_no_result_mapping" if changed else "copied_no_matching_placeholders"
    _record_template_render_result(
        context,
        renderer="xlsx",
        changed=changed,
        status=status,
        discovered_table_count=len(table_stats),
        mapped_table_count=sum(
            1 for item in table_stats if isinstance(item, dict) and int(item.get("filled_cells") or 0) > 0
        ),
        mapped_value_cell_count=mapped_value_cell_count,
        financial_semantic_plan_sha256=semantic_manifest.get("plan_sha256"),
        financial_semantic_candidate_count=semantic_manifest.get("candidate_count"),
        financial_semantic_bound_count=semantic_manifest.get("bound_count"),
        audit_result_mapped=audit_result_mapped,
        audit_result_mapping_reason=contract.get("audit_result_mapping_reason"),
        sheet_count=len(workbook.worksheets),
    )
    return output.getvalue(), status


def _render_xls(data: bytes, context: dict[str, Any]) -> tuple[bytes, str]:
    """Fill BIFF8 through Excel while keeping the original ``.xls`` format.

    ``xlrd`` is read-only and rebuilding the workbook with ``openpyxl`` turns
    BIFF8 into OOXML, which changes both the extension and the workbook's
    structure.  On Windows deployments with Excel automation available, use
    Excel's own SaveAs(FileFormat=56) path so formulas, styles, merged cells,
    print settings and sheet order remain under the original file format. If
    automation is unavailable, return the original bytes unchanged rather than
    silently delivering a falsely converted workbook.
    """

    try:
        import pythoncom
        import win32com.client
    except ImportError:
        return data, "copied_legacy_unmodified"

    changed = 0
    result_token_replaced = False
    excel = None
    workbook = None
    pythoncom.CoInitialize()
    try:
        with tempfile.TemporaryDirectory(prefix="annual-audit-xls-") as temp_dir:
            source_path = Path(temp_dir) / "source.xls"
            output_path = Path(temp_dir) / "rendered.xls"
            source_path.write_bytes(data)

            excel = win32com.client.DispatchEx("Excel.Application")
            excel.Visible = False
            excel.DisplayAlerts = False
            excel.ScreenUpdating = False
            try:
                # 3 = msoAutomationSecurityForceDisable; templates must never
                # execute embedded macros during server-side rendering.
                excel.AutomationSecurity = 3
            except Exception:
                pass
            workbook = excel.Workbooks.Open(
                str(source_path),
                UpdateLinks=0,
                ReadOnly=False,
                AddToMru=False,
            )
            for worksheet in workbook.Worksheets:
                used_range = worksheet.UsedRange
                row_count = int(used_range.Rows.Count)
                column_count = int(used_range.Columns.Count)
                for row_index in range(1, row_count + 1):
                    for column_index in range(1, column_count + 1):
                        cell = used_range.Cells(row_index, column_index)
                        value = cell.Value
                        if not isinstance(value, str):
                            continue
                        has_result_token = _template_contains_result_token(value)
                        replaced = _replace_template_markers(value, context)
                        if replaced != value:
                            cell.Value = replaced
                            changed += 1
                            if has_result_token and not _template_contains_result_token(replaced):
                                result_token_replaced = True

            # 56 is Excel's native Excel 97-2003 Workbook format (.xls).
            workbook.SaveAs(str(output_path), FileFormat=56)
            workbook.Close(SaveChanges=False)
            workbook = None
            excel.Quit()
            excel = None
            output_bytes = output_path.read_bytes()
            audit_result_mapped = bool(result_token_replaced)
            status = "filled" if changed else "copied_no_matching_placeholders"
            contract = context.get("__template_render_contract__") or {}
            contract["audit_result_mapped"] = audit_result_mapped
            contract["audit_result_mapping_reason"] = "explicit_result_token" if result_token_replaced else "none"
            if context.get("__enforce_result_mapping__") and not audit_result_mapped:
                status = "metadata_filled_no_result_mapping" if changed else "copied_no_matching_placeholders"
            _record_template_render_result(
                context,
                renderer="excel_legacy",
                changed=changed,
                status=status,
                audit_result_mapped=audit_result_mapped,
                audit_result_mapping_reason=contract.get("audit_result_mapping_reason"),
            )
            return output_bytes, status
    except Exception:
        # A headless/service account may not be allowed to automate Excel. Keep
        # the source format intact and let package validation report the
        # limitation; never fall back to an .xlsx reconstruction.
        return data, "copied_legacy_unmodified"
    finally:
        if workbook is not None:
            try:
                workbook.Close(SaveChanges=False)
            except Exception:
                pass
        if excel is not None:
            try:
                excel.Quit()
            except Exception:
                pass
        pythoncom.CoUninitialize()


def _legacy_word_replacement_pairs(context: dict[str, Any]) -> dict[str, str]:
    """Return literal pairs safe to apply through Word's Find/Replace API."""

    entity_name = str(context.get("entity_name") or "待补充被审计单位")
    fiscal_year = int(context.get("fiscal_year") or 0)
    issue_year = int(context.get("issue_year") or (fiscal_year + 1 if fiscal_year else 0))
    period_end = str(context.get("period_end") or "待补充日期")
    pairs: dict[str, str] = {}
    for key, value in context.items():
        if key.startswith("__") or isinstance(value, (dict, list, tuple)):
            continue
        replacement = "" if value is None else str(value)
        pairs.update(
            {
                f"{{{{{key}}}}}": replacement,
                f"[[{key}]]": replacement,
                f"${{{key}}}": replacement,
            }
        )
    if fiscal_year:
        pairs.update(
            {
                "二〇二五年度": f"{_chinese_year(fiscal_year)}年度",
                "二○二五年度": f"{_chinese_year(fiscal_year)}年度",
                "2025年度": f"{fiscal_year}年度",
                "2025年12月31日": period_end,
            }
        )
    if issue_year:
        pairs["2026年度"] = f"{issue_year}年度"
        pairs["2026年12月31日"] = period_end
    pairs.update(
        {
            "一般企业模板": entity_name,
            "一般企业报表": entity_name,
            "一般企业附注": entity_name,
            "XXX公司": entity_name,
            "【被审计单位名称】": entity_name,
            "[被审计单位名称]": entity_name,
            "【年度】": f"{fiscal_year}年度" if fiscal_year else "待补充年度",
            "【20XX年度】": f"{fiscal_year}年度" if fiscal_year else "待补充年度",
            "20XX年度": f"{fiscal_year}年度" if fiscal_year else "待补充年度",
            "【日期】": period_end,
            "[日期]": period_end,
        }
    )
    return {needle: replacement for needle, replacement in pairs.items() if needle and needle != replacement}


def _render_doc(data: bytes, context: dict[str, Any]) -> tuple[bytes, str]:
    """Fill legacy binary Word documents without converting them to DOCX."""

    try:
        import pythoncom
        import win32com.client
    except ImportError:
        return data, "copied_legacy_unmodified"

    word = None
    document = None
    pythoncom.CoInitialize()
    try:
        with tempfile.TemporaryDirectory(prefix="annual-audit-doc-") as temp_dir:
            source_path = Path(temp_dir) / "source.doc"
            output_path = Path(temp_dir) / "rendered.doc"
            source_path.write_bytes(data)
            word = win32com.client.DispatchEx("Word.Application")
            word.Visible = False
            word.DisplayAlerts = 0
            document = word.Documents.Open(
                str(source_path),
                ConfirmConversions=False,
                ReadOnly=False,
                AddToRecentFiles=False,
            )
            changed = 0
            has_result_token = _template_contains_result_token(str(document.Content.Text or ""))
            for needle, replacement in _legacy_word_replacement_pairs(context).items():
                for story in document.StoryRanges:
                    current = story
                    while current is not None:
                        find = current.Find
                        find.ClearFormatting()
                        find.Replacement.ClearFormatting()
                        find.Text = needle
                        find.Replacement.Text = replacement
                        find.Forward = True
                        find.Wrap = 1  # wdFindContinue
                        find.Format = False
                        find.MatchCase = False
                        find.MatchWholeWord = False
                        if find.Execute(Replace=2):  # wdReplaceAll
                            changed += 1
                        try:
                            current = current.NextStoryRange
                        except Exception:
                            current = None
            result_text = str(document.Content.Text or "")
            document.SaveAs(str(output_path), FileFormat=0)  # wdFormatDocument
            output_bytes = output_path.read_bytes()
            document.Close(SaveChanges=False)
            document = None
            word.Quit()
            word = None
            status = "filled" if changed else "copied_no_matching_placeholders"
            # A changed legacy file without an explicit result token is
            # metadata-only.
            result_token_replaced = has_result_token and not _template_contains_result_token(
                result_text
            )
            audit_result_mapped = bool(result_token_replaced)
            contract = context.get("__template_render_contract__") or {}
            contract["audit_result_mapped"] = audit_result_mapped
            contract["audit_result_mapping_reason"] = "explicit_result_token" if result_token_replaced else "none"
            if context.get("__enforce_result_mapping__") and not audit_result_mapped:
                status = "metadata_filled_no_result_mapping" if changed else "copied_no_matching_placeholders"
            _record_template_render_result(
                context,
                renderer="word_legacy",
                changed=changed,
                status=status,
                audit_result_mapped=audit_result_mapped,
                audit_result_mapping_reason=contract.get("audit_result_mapping_reason"),
            )
            return output_bytes, status
    except Exception:
        return data, "copied_legacy_unmodified"
    finally:
        if document is not None:
            try:
                document.Close(SaveChanges=False)
            except Exception:
                pass
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass
        pythoncom.CoUninitialize()


def _render_pdf(data: bytes, context: dict[str, Any]) -> tuple[bytes, str]:
    """Fill only fields that are actually defined by the supplied PDF form."""

    try:
        from pypdf import PdfReader, PdfWriter

        reader = PdfReader(BytesIO(data))
        fields = reader.get_fields() or {}
        updates: dict[str, str] = {}
        for field_name in fields:
            key = str(field_name).strip()
            if key in context:
                updates[key] = "" if context[key] is None else str(context[key])
        base = data
        if updates:
            writer = PdfWriter()
            writer.clone_document_from_reader(reader)
            for page in writer.pages:
                writer.update_page_form_field_values(page, updates, auto_regenerate=False)
            output = BytesIO()
            writer.write(output)
            base = output.getvalue()
        if not updates:
            _record_template_render_result(
                context,
                renderer="pdf",
                changed=0,
                status="copied_no_matching_form_fields",
                audit_result_mapped=False,
                audit_result_mapping_reason="none",
            )
            return data, "copied_no_matching_form_fields"
        audit_result_mapped = any(key in _RESULT_TOKEN_KEYS for key in updates)
        status = "filled_acroform" if audit_result_mapped else "metadata_filled_no_result_mapping"
        _record_template_render_result(
            context,
            renderer="pdf",
            changed=len(updates),
            status=status,
            audit_result_mapped=audit_result_mapped,
            audit_result_mapping_reason="explicit_result_token" if audit_result_mapped else "none",
        )
        return base, status
    except Exception:
        return data, "copied_pdf"


def _iter_findings(snapshot: dict[str, Any]) -> list[dict[str, Any]]:
    """Return de-duplicated deterministic findings in report order."""

    candidates: list[dict[str, Any]] = []
    sales = snapshot.get("sales_receivables") or {}
    candidates.extend(item for item in sales.get("findings") or [] if isinstance(item, dict))
    candidates.extend(item for item in (sales.get("revenue") or {}).get("findings") or [] if isinstance(item, dict))
    candidates.extend(item for item in (sales.get("receivables") or {}).get("findings") or [] if isinstance(item, dict))
    candidates.extend(item for item in (snapshot.get("cash_and_bank") or {}).get("findings") or [] if isinstance(item, dict))

    result: list[dict[str, Any]] = []
    seen: set[tuple[str, str, str]] = set()
    for item in candidates:
        key = (
            str(item.get("finding_type") or item.get("title") or ""),
            str(item.get("description") or ""),
            str(item.get("amount") or ""),
        )
        if key in seen:
            continue
        seen.add(key)
        result.append(item)
    return result


def _evidence_location_text(reference: dict[str, Any]) -> str:
    locator = reference.get("source_locator") or reference.get("source_locator_json") or {}
    if not isinstance(locator, dict):
        return "来源定位待补充"
    parts = [
        str(locator.get(key) or "").strip()
        for key in ("file_name", "sheet_name", "page_number", "row_number", "cell_address", "voucher_no")
    ]
    parts = [part for part in parts if part]
    return " / ".join(parts) if parts else "来源定位待补充"


def _template_result_text(
    template_type: str,
    report_text: str,
    findings_summary: str,
    snapshot: dict[str, Any],
) -> str:
    """Return a bounded result value for an explicit template result field.

    Explicit result fields are still semantic fields, not an invitation to
    paste the entire conversational report into the attachment.  Formal
    report and statement templates receive only their controlled status text;
    a management-letter template may receive the bounded finding summary.
    """

    gate = snapshot.get("release_gate") or {}
    gate_ready = gate.get("gate_status") == "ready_for_signature"
    if template_type == "management_letter":
        return findings_summary
    if template_type == "annual_report":
        return (
            "当前项目组签发门禁已通过，审计意见待项目组按正式签发流程确认。"
            if gate_ready
            else "本文件为基于已导入资料形成的审计报告格式草稿，当前未形成正式审计意见。"
        )
    if template_type == "financial_statements":
        return "财务报表金额已按模板对应报表项目填列；审计意见及证据结论以审计报告和项目组复核记录为准。"
    if template_type == "notes":
        return "财务报表附注仅填列模板对应的主体信息、期间信息和报表项目数据，不嵌入审计工作底稿全文。"
    return "审计结果已按模板定义的结果字段填列。"


def _build_audit_result_appendix(
    snapshot: dict[str, Any],
    report_text: str,
    *,
    template_type: str,
    template: dict[str, Any],
) -> str:
    raise RuntimeError("通用审计结果附录已停用：结果必须写入模板定义的位置")

    """Build a source-traceable result appendix shared by every file format.

    This is deliberately derived from the frozen report snapshot.  It does not
    invent an audit opinion when the deterministic run is incomplete, and it
    keeps the evidence limitations visible in the delivered attachment.
    """

    engagement = snapshot.get("engagement") or {}
    readiness = snapshot.get("readiness") or {}
    counts = readiness.get("counts") or {}
    citation = snapshot.get("citation_plan_summary") or {}
    gate = snapshot.get("release_gate") or {}
    lines = [
        "审计结果及证据复核说明（自动化审计工作草稿）",
        f"附件用途：{_ATTACHMENT_NAME_LABELS.get(template_type, template_type)}",
        f"模板版本：{template_version_ref(template)}",
        f"被审计单位：{engagement.get('entity_name') or snapshot.get('entity_name') or '待补充'}",
        f"审计期间：{engagement.get('period_start') or '-'} 至 {engagement.get('period_end') or snapshot.get('period_end') or '-'}",
        "",
        "一、结论性质",
        "本附件只反映当前已落库资料和确定性审计规则的执行结果。规则命中不等同于错报或舞弊；在证据闭环、抽样、管理层沟通、期后事项和项目组复核完成前，不形成正式审计意见。",
        "",
        "二、已执行数据范围",
        f"- 科目余额行数：{int(counts.get('account_balance_rows') or 0)}",
        f"- 序时账/凭证明细行数：{int(counts.get('journal_entry_rows') or 0)}",
        f"- 应收账款明细行数：{int(counts.get('receivable_rows') or 0)}",
        f"- 银行流水行数：{int(counts.get('bank_transaction_rows') or 0)}",
    ]

    available = readiness.get("available_data") or []
    if available:
        lines.append("- 已识别资料：" + "；".join(
            f"{item.get('name') or item.get('code') or '未命名'} {int(item.get('row_count') or 0)} 行"
            for item in available[:20]
            if isinstance(item, dict)
        ))
    missing = [str(item) for item in (
        *(readiness.get("missing_required_data") or []),
        *(readiness.get("supplemental_required_data") or []),
    ) if str(item).strip()]
    lines.append("- 待补充资料：" + ("、".join(dict.fromkeys(missing)) if missing else "当前演示范围内未记录结构化资料缺口"))

    findings = _iter_findings(snapshot)
    lines.extend(["", "三、确定性规则结果"])
    if not findings:
        lines.append("- 当前已执行规则未形成命中事项；若对应源数据未导入或为 0 行，该结果不应解释为不存在异常。")
    for index, finding in enumerate(findings, start=1):
        level = {
            "urgent": "紧急",
            "high": "高",
            "medium": "中",
            "low": "低",
        }.get(str(finding.get("risk_level") or "").lower(), "待定")
        amount = finding.get("amount")
        amount_text = f"；涉及金额：{_money(amount)} 元" if amount not in (None, "") else ""
        refs = [item for item in finding.get("evidence_refs") or [] if isinstance(item, dict)]
        locations = "、".join(_evidence_location_text(item) for item in refs[:3])
        evidence_text = f"；证据定位：{locations}" if locations else "；证据定位：待项目组补充或核验"
        lines.append(
            f"- {index}. [{level}风险] {finding.get('title') or '待复核事项'}："
            f"{finding.get('description') or '缺少规则说明'}{amount_text}{evidence_text}"
        )

    lines.extend([
        "",
        "四、引用与签发门禁",
        f"- 已绑定引用：{int(citation.get('cited_claims') or 0)} / {int(citation.get('total_claims') or 0)} 项主张；引用投影数：{int(citation.get('projected_count') or 0)}。",
        f"- 当前签发门禁：{'已具备待签发条件' if gate.get('gate_status') == 'ready_for_signature' else '阻断，仍需完成项目组程序和复核'}。",
        "- 交付状态：工作草稿；不得将本附件直接作为已签发审计报告、已审计财务报表或已完成管理层沟通的证明。",
    ])
    if report_text.strip():
        lines.extend(["", "五、审计结果正文（确定性审计草稿）", report_text[:40000]])
    return "\n".join(lines).strip()


def _append_pdf_audit_result(data: bytes, context: dict[str, Any]) -> bytes:
    """Append a selectable, paginated result section while keeping PDF pages."""

    raise RuntimeError("PDF 通用结果附页已停用：PDF 必须使用模板已有字段")

    from pypdf import PdfReader, PdfWriter
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.pdfgen import canvas

    appendix = str(context.get("__audit_result_appendix__") or "").strip()
    if not appendix:
        raise ValueError("审计结果为空，拒绝生成无结果附件")

    font_name = "Helvetica"
    font_candidates = [
        Path("C:/Windows/Fonts/msyh.ttf"),
        Path("C:/Windows/Fonts/simhei.ttf"),
        Path("/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc"),
        Path("/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttf"),
    ]
    for font_path in font_candidates:
        if not font_path.exists():
            continue
        try:
            pdfmetrics.registerFont(TTFont("AnnualAuditCJK", str(font_path)))
            font_name = "AnnualAuditCJK"
            break
        except Exception:
            continue
    if font_name == "Helvetica":
        try:
            from reportlab.pdfbase.cidfonts import UnicodeCIDFont

            pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
            font_name = "STSong-Light"
        except Exception:
            pass

    appendix_pdf = BytesIO()
    page_width, page_height = A4
    pdf = canvas.Canvas(appendix_pdf, pagesize=A4)
    pdf.setTitle("年度审计结果及证据复核说明")
    pdf.setFont(font_name, 10)
    left = 42
    top = page_height - 48
    y = top
    max_chars = 58 if font_name != "Helvetica" else 92

    def new_page() -> None:
        nonlocal y
        pdf.showPage()
        pdf.setFont(font_name, 10)
        y = top

    for raw_line in appendix.splitlines():
        line = raw_line or " "
        while line:
            chunk = line[:max_chars]
            line = line[max_chars:]
            if y < 48:
                new_page()
            pdf.drawString(left, y, chunk)
            y -= 15
        if y < 48:
            new_page()
    pdf.save()
    appendix_reader = PdfReader(BytesIO(appendix_pdf.getvalue()))
    source_reader = PdfReader(BytesIO(data))
    writer = PdfWriter()
    for page in source_reader.pages:
        writer.add_page(page)
    for page in appendix_reader.pages:
        writer.add_page(page)
    output = BytesIO()
    writer.write(output)
    return output.getvalue()


def _template_contains_result_token(text: str) -> bool:
    for match in _TOKEN_RE.finditer(text):
        if any(key in _RESULT_TOKEN_KEYS for key in match.groups() if key):
            return True
    for match in _SINGLE_BRACKET_TOKEN_RE.finditer(text):
        if match.group(1) in _RESULT_TOKEN_KEYS:
            return True
    return False


def _schema_contains_result_mapping(schema: Any) -> bool:
    """Return whether an explicit field schema maps a result-like value."""

    if isinstance(schema, dict):
        for key, value in schema.items():
            text = str(key)
            if any(marker in text for marker in _RESULT_TOKEN_KEYS):
                return True
            if isinstance(value, str) and any(marker in value for marker in _RESULT_TOKEN_KEYS):
                return True
            if _schema_contains_result_mapping(value):
                return True
    elif isinstance(schema, list):
        return any(_schema_contains_result_mapping(item) for item in schema)
    return False


def _render_text_template(data: bytes, extension: str, context: dict[str, Any]) -> tuple[bytes, str]:
    """Render text templates without changing encoding or the file suffix."""

    from .attachment_document_composer import compose_attachment_from_context

    data = compose_attachment_from_context(
        data,
        str(context.get("__source_file_name__") or f"template{extension}"),
        context,
    )
    composition_manifest = context.get("__attachment_composition_manifest__") or {}
    ai_authored_changes = int(composition_manifest.get("authored_block_count") or 0)
    has_bom = data.startswith(b"\xef\xbb\xbf")
    text = data.decode("utf-8-sig")
    replaced = _replace_template_markers(text, context)
    has_result_token = _template_contains_result_token(text)
    result_token_replaced = has_result_token and not _template_contains_result_token(replaced)
    audit_result_mapped = bool(ai_authored_changes or result_token_replaced)
    status = "filled" if ai_authored_changes or replaced != text else "copied_no_matching_placeholders"
    if context.get("__enforce_result_mapping__") and not audit_result_mapped:
        status = "metadata_filled_no_result_mapping" if replaced != text else "copied_no_matching_placeholders"
    contract = context.get("__template_render_contract__") or {}
    contract["audit_result_mapped"] = audit_result_mapped
    contract["audit_result_mapping_reason"] = (
        "template_slot_authoring" if ai_authored_changes else
        "explicit_result_token" if result_token_replaced else
        "none"
    )
    _record_template_render_result(
        context,
        renderer="text",
        changed=ai_authored_changes + int(replaced != text),
        ai_authored_block_count=ai_authored_changes,
        source_business_body_reused=composition_manifest.get("source_business_body_reused"),
        status=status,
        audit_result_mapped=audit_result_mapped,
        audit_result_mapping_reason=contract.get("audit_result_mapping_reason"),
    )
    encoded = replaced.encode("utf-8")
    return (b"\xef\xbb\xbf" + encoded if has_bom else encoded), status


def _append_docx_audit_result(document: Any, context: dict[str, Any]) -> None:
    """Append the result only when the template has no result field."""

    raise RuntimeError("DOCX 通用结果附页已停用：DOCX 必须使用模板已有段落或表格")

    document_text = [p.text or "" for p in document.paragraphs]
    for table in document.tables:
        document_text.extend(cell.text or "" for row in table.rows for cell in row.cells)
    if _template_contains_result_token("\n".join(document_text)):
        return
    appendix = str(context.get("__audit_result_appendix__") or "").strip()
    if not appendix:
        raise ValueError("审计结果为空，拒绝生成无结果附件")
    document.add_page_break()
    lines = appendix.splitlines()
    title = lines.pop(0) if lines else "审计结果及证据复核说明"
    try:
        document.add_heading(title, level=1)
    except Exception:
        document.add_paragraph(title)
    for line in lines:
        text = line.strip()
        if not text:
            document.add_paragraph("")
        elif text.startswith("## "):
            document.add_heading(text[3:].strip(), level=2)
        elif text.startswith("### "):
            document.add_heading(text[4:].strip(), level=3)
        else:
            document.add_paragraph(text)


def _append_xlsx_audit_result(workbook: Any, context: dict[str, Any]) -> None:
    raise RuntimeError("XLSX 通用结果工作表已停用：XLSX 必须使用模板已有工作表")

    from openpyxl.styles import Alignment, Font, PatternFill

    appendix = str(context.get("__audit_result_appendix__") or "").strip()
    if not appendix:
        raise ValueError("审计结果为空，拒绝生成无结果附件")
    base_name = "审计结果"
    name = base_name
    index = 2
    while name in workbook.sheetnames:
        name = f"{base_name}{index}"
        index += 1
    sheet = workbook.create_sheet(name)
    sheet.column_dimensions["A"].width = 110
    sheet["A1"] = appendix.splitlines()[0] if appendix.splitlines() else base_name
    sheet["A1"].font = Font(bold=True, size=14, color="FFFFFF")
    sheet["A1"].fill = PatternFill("solid", fgColor="1F4E78")
    sheet["A1"].alignment = Alignment(wrap_text=True, vertical="top")
    for row_index, line in enumerate(appendix.splitlines()[1:], start=2):
        sheet.cell(row=row_index, column=1, value=line)
        sheet.cell(row=row_index, column=1).alignment = Alignment(wrap_text=True, vertical="top")
    sheet.freeze_panes = "A2"


def _field_schema_value(schema: Any, context: dict[str, Any], key: str) -> Any:
    value: Any = context.get(key, "")
    if isinstance(schema, dict):
        source = schema.get(key)
        if isinstance(source, str):
            value = context.get(source, value)
        elif isinstance(source, dict):
            source_key = source.get("source") or source.get("key")
            if source_key:
                value = context.get(str(source_key), value)
    return value


def _format_docx_contract_value(value: Any, source: dict[str, Any]) -> Any:
    if value in (None, ""):
        return ""
    if isinstance(value, (dict, list, tuple, set, frozenset)):
        raise ValueError("DOCX 槽位 source 只能解析为客户可见的标量值")
    presentation = str(source.get("format") or "").lower()
    rendered: Any = value
    if presentation == "money":
        rendered = _money(value)
    elif presentation == "integer":
        try:
            rendered = str(int(Decimal(str(value))))
        except (InvalidOperation, TypeError, ValueError):
            rendered = value
    elif presentation == "percent":
        try:
            numeric = Decimal(str(value))
        except (InvalidOperation, TypeError, ValueError):
            rendered = value
        else:
            if abs(numeric) <= 1:
                numeric *= Decimal("100")
            rendered = f"{numeric.quantize(Decimal('0.01'))}%"
    elif presentation == "chinese_year":
        try:
            rendered = _chinese_year(int(str(value)))
        except (TypeError, ValueError):
            raise ValueError("DOCX 槽位 chinese_year source 必须是四位财务年度") from None
    prefix = str(source.get("prefix") or "")
    suffix = str(source.get("suffix") or "")
    if prefix or suffix:
        return f"{prefix}{rendered}{suffix}"
    return rendered


def _docx_contract_source_kind(source: Any) -> str:
    if isinstance(source, str):
        candidate = source.strip()
        if candidate.startswith("fact:"):
            return "fact"
        if candidate.startswith("table_value:"):
            return "table_value"
        return "context"
    if isinstance(source, dict):
        return str(source.get("kind") or "context").lower()
    return ""


def _docx_contract_source_value(source: Any, context: dict[str, Any]) -> Any:
    """Resolve one allowlisted deterministic field for a declared DOCX slot."""

    if isinstance(source, str):
        key = source.strip()
        if key.startswith("fact:") or key.startswith("table_value:"):
            kind, _, remainder = key.partition(":")
            label, separator, period = remainder.rpartition(":")
            if not separator:
                label, period = remainder, "current"
            source = {"kind": kind, "label": label, "period": period}
        else:
            return _format_docx_contract_value(context.get(key, ""), {})
    if not isinstance(source, dict):
        return ""
    kind = str(source.get("kind") or "").lower()
    if kind == "context":
        return _format_docx_contract_value(context.get(str(source.get("key") or ""), ""), source)
    label = str(source.get("label") or "")
    period = str(source.get("period") or "current").lower()
    if kind == "table_value":
        values = context.get("__table_values__") or {}
        if not isinstance(values, dict):
            return ""
        lookup = f"{label}_期初" if period == "opening" else label
        return _format_docx_contract_value(values.get(lookup, ""), source)
    if kind == "fact":
        registry = context.get("__fact_registry__") or {}
        if not isinstance(registry, dict):
            return ""
        key = _normalize_material_label(_normalize_template_label(label))
        item = registry.get(key) or registry.get(label) or {}
        if not isinstance(item, dict):
            return ""
        return _format_docx_contract_value(item.get(period, ""), source)
    return ""


def _declared_docx_semantic_coverage(
    context: dict[str, Any],
    *,
    template_type: str,
) -> dict[str, Any]:
    """Classify what a declared DOCX contract actually placed in the file."""

    raw_stats = context.get("__docx_slot_contract_stats__") or []
    stats = [item for item in raw_stats if isinstance(item, dict)]
    structured_value_count = sum(
        int(item.get("structured_value_count") or 0) for item in stats
    )
    ai_value_count = sum(int(item.get("ai_value_count") or 0) for item in stats)
    if template_type == "financial_statements":
        mapped = structured_value_count > 0
        reason = "declared_structured_fact_slots"
    elif template_type == "notes":
        mapped = structured_value_count > 0 and ai_value_count > 0
        reason = "declared_note_facts_and_ai_slots"
    else:
        mapped = ai_value_count > 0
        reason = "declared_ai_narrative_slots"
    return {
        "mapped": mapped,
        "reason": reason if mapped else "declared_slots_only_metadata",
        "structured_value_count": structured_value_count,
        "ai_value_count": ai_value_count,
        "slot_count": len(stats),
    }


def _apply_xlsx_field_schema(workbook: Any, context: dict[str, Any]) -> int:
    """Apply optional explicit cell mappings saved with a template version."""

    schema = context.get("__field_schema__") or {}
    mappings: dict[str, Any] = {}
    if isinstance(schema, dict):
        for key in ("cells", "field_to_cell", "fields"):
            candidate = schema.get(key)
            if isinstance(candidate, dict):
                mappings.update(candidate)
    elif isinstance(schema, list):
        for item in schema:
            if not isinstance(item, dict):
                continue
            key = item.get("key") or item.get("field") or item.get("name")
            address = item.get("cell") or item.get("address")
            if key and address:
                mappings[str(address)] = str(key)
    changed = 0
    for address, key_spec in mappings.items():
        if "!" in str(address):
            sheet_name, cell_address = str(address).split("!", 1)
            sheet = workbook[sheet_name.strip("'")] if sheet_name.strip("'") in workbook.sheetnames else None
        else:
            sheet = workbook.active
            cell_address = str(address)
        if sheet is None:
            continue
        key = str(key_spec.get("source") or key_spec.get("key") if isinstance(key_spec, dict) else key_spec)
        if not key:
            continue
        cell = sheet[cell_address]
        if cell.__class__.__name__ == "MergedCell":
            continue
        if isinstance(cell.value, str) and cell.value.startswith("="):
            continue
        value = _field_schema_value(schema, context, key)
        if value in (None, ""):
            continue
        if cell.value != value:
            cell.value = value
            changed += 1
    return changed


def _apply_docx_field_schema(document: Any, context: dict[str, Any]) -> int:
    """Fill approved DOCX paragraph/cell slots recorded with a template.

    The schema is intentionally address-based and opt-in.  It gives template
    administrators a deterministic path for dynamic prose without allowing a
    model to invent, delete or reorder Word paragraphs.
    """

    schema = context.get("__field_schema__") or {}
    declared_contract = context.get("__attachment_slot_contract__") or {}
    if not isinstance(declared_contract, dict):
        declared_contract = {}
    if declared_contract.get("declared"):
        expected_template_sha = str(declared_contract.get("template_sha256") or "")
        actual_template_sha = str(context.get("__source_template_sha256__") or "")
        if expected_template_sha and actual_template_sha and expected_template_sha != actual_template_sha:
            raise ValueError("DOCX 槽位契约与当前模板 SHA-256 不一致")
        slot_values = context.get("__attachment_slot_values__") or {}
        if not isinstance(slot_values, dict):
            raise ValueError("DOCX AI 槽位值必须是对象")
        changed = 0
        written_slot_ids: list[str] = []
        materialized_docvariable_slot_ids: list[str] = []
        slot_contract_stats: list[dict[str, Any]] = []
        for raw_slot in declared_contract.get("slots") or []:
            if not isinstance(raw_slot, dict):
                raise ValueError("DOCX 槽位契约包含无效槽位")
            slot_id = str(raw_slot.get("slot_id") or "").strip()
            target = str(raw_slot.get("target") or "").strip()
            mode = str(raw_slot.get("mode") or "replace").strip().lower()
            if not slot_id or not target:
                raise ValueError("DOCX 槽位契约缺少 slot_id 或 target")
            paragraph_match = re.fullmatch(
                r"paragraph:(?P<index>\d+)", target, re.IGNORECASE
            )
            paragraph: Any | None = None
            cell: Any | None = None
            current_text = ""
            if paragraph_match:
                index = int(paragraph_match.group("index"))
                if not 0 <= index < len(document.paragraphs):
                    raise ValueError(f"DOCX 槽位段落不存在：{slot_id}")
                paragraph = document.paragraphs[index]
                current_text = paragraph.text or ""
            else:
                frame_match = re.fullmatch(
                    r"(?P<frame>header|footer):(?P<section>\d+),paragraph:(?P<paragraph>\d+)",
                    target,
                    re.IGNORECASE,
                )
                if frame_match:
                    section_index = int(frame_match.group("section"))
                    paragraph_index = int(frame_match.group("paragraph"))
                    if not 0 <= section_index < len(document.sections):
                        raise ValueError(f"DOCX 槽位页眉页脚分节不存在：{slot_id}")
                    frame = getattr(
                        document.sections[section_index],
                        frame_match.group("frame").lower(),
                    )
                    if frame.is_linked_to_previous:
                        raise ValueError(f"DOCX 槽位不能直接写入继承的页眉页脚：{slot_id}")
                    if not 0 <= paragraph_index < len(frame.paragraphs):
                        raise ValueError(f"DOCX 槽位页眉页脚段落不存在：{slot_id}")
                    paragraph = frame.paragraphs[paragraph_index]
                    current_text = paragraph.text or ""
                else:
                    cell_match = re.fullmatch(
                        r"table:(?P<table>\d+),row:(?P<row>\d+),cell:(?P<cell>\d+)",
                        target,
                        re.IGNORECASE,
                    )
                    if not cell_match:
                        raise ValueError(f"DOCX 槽位目标无效：{slot_id}")
                    table_index = int(cell_match.group("table"))
                    row_index = int(cell_match.group("row"))
                    cell_index = int(cell_match.group("cell"))
                    if not 0 <= table_index < len(document.tables):
                        raise ValueError(f"DOCX 槽位表格不存在：{slot_id}")
                    table = document.tables[table_index]
                    if not 0 <= row_index < len(table.rows) or not 0 <= cell_index < len(table.rows[row_index].cells):
                        raise ValueError(f"DOCX 槽位单元格不存在：{slot_id}")
                    cell = table.rows[row_index].cells[cell_index]
                    if len(cell.paragraphs) != 1:
                        raise ValueError(f"DOCX 槽位单元格不是单段落字段：{slot_id}")
                    current_text = cell.paragraphs[0].text or ""
            expected_text_sha = str(raw_slot.get("source_text_sha256") or "")
            actual_text_sha = hashlib.sha256(current_text.encode("utf-8")).hexdigest()
            if expected_text_sha and expected_text_sha != actual_text_sha:
                raise ValueError(f"DOCX 槽位原文已变化：{slot_id}")

            target_paragraph = paragraph if paragraph is not None else cell.paragraphs[0]
            field_materialized = (
                _materialize_docx_docvariable_fields(target_paragraph)
                if mode != "preserve"
                else False
            )
            if field_materialized:
                materialized_docvariable_slot_ids.append(slot_id)

            stats = {
                "slot_id": slot_id,
                "target": target,
                "mode": mode,
                "changed": False,
                "source_kinds": [],
                "structured_value_count": 0,
                "ai_value_count": 0,
            }
            if mode == "preserve":
                slot_contract_stats.append(stats)
                continue
            if mode == "replace_inline":
                replacements: list[tuple[str, str]] = []
                for raw_replacement in raw_slot.get("replacements") or []:
                    if not isinstance(raw_replacement, dict):
                        raise ValueError(f"DOCX 行内槽位 replacement 无效：{slot_id}")
                    expected_text = str(raw_replacement.get("expected_text") or "")
                    source = raw_replacement.get("source")
                    if not expected_text or not source:
                        raise ValueError(f"DOCX 行内槽位 replacement 不完整：{slot_id}")
                    occurrences = int(raw_replacement.get("occurrences") or 0)
                    if occurrences <= 0 or current_text.count(expected_text) != occurrences:
                        raise ValueError(f"DOCX 行内槽位原文已变化：{slot_id}")
                    value = _docx_contract_source_value(source, context)
                    if value in (None, "") and bool(raw_replacement.get("required", True)):
                        raise ValueError(f"DOCX 行内槽位没有可写入值：{slot_id}")
                    value_text = "" if value is None else str(value)
                    if len(value_text) > int(raw_slot.get("max_chars") or 1200):
                        raise ValueError(f"DOCX 行内槽位文本超过声明长度：{slot_id}")
                    source_kind = _docx_contract_source_kind(source)
                    stats["source_kinds"].append(source_kind)
                    if source_kind in {"fact", "table_value"} and value_text:
                        stats["structured_value_count"] += 1
                    replacements.append((expected_text, value_text))
                wrote_value = (
                    _replace_docx_paragraph_exact_replacements(
                        target_paragraph,
                        replacements,
                    )
                    or field_materialized
                )
                if wrote_value:
                    changed += 1
                stats["changed"] = wrote_value
                written_slot_ids.append(slot_id)
                slot_contract_stats.append(stats)
                continue

            source_kinds: list[str] = []
            if mode == "suppress_instruction":
                value = ""
            elif mode == "replace":
                source = raw_slot.get("source")
                has_source = bool(source) if not isinstance(source, str) else bool(source.strip())
                if has_source:
                    value = _docx_contract_source_value(source, context)
                    source_kinds.append(_docx_contract_source_kind(source))
                else:
                    if slot_id not in slot_values:
                        if bool(raw_slot.get("required", True)):
                            raise ValueError(f"DOCX 必填 AI 槽位尚未编制：{slot_id}")
                        slot_contract_stats.append(stats)
                        continue
                    value = slot_values[slot_id]
                    stats["ai_value_count"] = 1
                if value in (None, "") and bool(raw_slot.get("required", True)):
                    raise ValueError(f"DOCX 必填槽位没有可写入值：{slot_id}")
            else:
                raise ValueError(f"DOCX 槽位 mode 无效：{slot_id}")
            value_text = "" if value is None else str(value)
            if bool(raw_slot.get("preserves_direct_layout_markers")) and any(
                control in value_text for control in ("\t", "\r", "\n")
            ):
                raise ValueError(
                    f"DOCX 带布局标记的 DOCVARIABLE 槽位不能写入制表符或换行：{slot_id}"
                )
            if len(value_text) > int(raw_slot.get("max_chars") or 1200):
                raise ValueError(f"DOCX 槽位文本超过声明长度：{slot_id}")
            text_changed = value_text != current_text
            if text_changed:
                if paragraph is not None:
                    _set_docx_paragraph_text(paragraph, value_text)
                elif cell is not None:
                    _set_docx_cell_text(cell, value_text)
            wrote_value = text_changed or field_materialized
            if wrote_value:
                changed += 1
            stats["changed"] = wrote_value
            stats["source_kinds"] = source_kinds
            if any(kind in {"fact", "table_value"} for kind in source_kinds) and value_text:
                stats["structured_value_count"] = 1
            written_slot_ids.append(slot_id)
            slot_contract_stats.append(stats)
        context["__docx_field_schema_slot_count__"] = changed
        context["__docx_written_slot_ids__"] = written_slot_ids
        context["__docx_materialized_docvariable_slot_ids__"] = (
            materialized_docvariable_slot_ids
        )
        context["__docx_slot_contract_stats__"] = slot_contract_stats
        return changed
    if context.get("__strict_client_delivery__"):
        raise ValueError("DOCX 正式交付必须使用经确认的 slot_contract")
    if not isinstance(schema, dict):
        return 0
    docx_schema = schema.get("docx") if isinstance(schema.get("docx"), dict) else schema
    raw_mappings: dict[str, Any] = {}
    for key in ("paragraphs", "field_to_paragraph", "cells", "field_to_cell", "slots"):
        candidate = docx_schema.get(key)
        if isinstance(candidate, dict):
            raw_mappings.update(candidate)
        elif isinstance(candidate, list):
            for item in candidate:
                if not isinstance(item, dict):
                    continue
                target = item.get("target") or item.get("slot") or item.get("address")
                field = item.get("field") or item.get("key") or item.get("source")
                if target and field:
                    raw_mappings[str(target)] = field

    def is_target(value: Any) -> bool:
        return bool(
            re.fullmatch(r"(?:paragraph|p)[:\[]\d+\]?", str(value or ""), re.IGNORECASE)
            or re.fullmatch(
                r"(?:table|tbl)[:\[]\d+\]?(?:[,./:]?(?:row|r)[:\[]\d+\]?)?(?:[,./:]?(?:cell|col|c)[:\[]\d+\]?)?",
                str(value or ""),
                re.IGNORECASE,
            )
        )

    changed = 0
    for raw_target, key_spec in raw_mappings.items():
        target = str(raw_target)
        field_spec = key_spec
        if not is_target(target) and isinstance(key_spec, str) and is_target(key_spec):
            target, field_spec = key_spec, raw_target
        if not is_target(target):
            continue
        field = str(
            field_spec.get("source") or field_spec.get("key")
            if isinstance(field_spec, dict)
            else field_spec
        ).strip()
        if not field:
            continue
        value = _field_schema_value(schema, context, field)
        if value in (None, ""):
            continue
        paragraph_match = re.fullmatch(r"(?:paragraph|p)[:\[](?P<index>\d+)\]?", target, re.IGNORECASE)
        if paragraph_match:
            index = int(paragraph_match.group("index"))
            if 0 <= index < len(document.paragraphs):
                _set_docx_paragraph_text(document.paragraphs[index], str(value))
                changed += 1
            continue
        numbers = [int(item) for item in re.findall(r"\d+", target)]
        if len(numbers) != 3:
            continue
        table_index, row_index, column_index = numbers
        if (
            0 <= table_index < len(document.tables)
            and 0 <= row_index < len(document.tables[table_index].rows)
            and 0 <= column_index < len(document.tables[table_index].rows[row_index].cells)
        ):
            _set_docx_cell_text(
                document.tables[table_index].rows[row_index].cells[column_index], value
            )
            changed += 1
    context["__docx_field_schema_slot_count__"] = changed
    return changed


def render_template_file(data: bytes, file_name: str, context: dict[str, Any]) -> tuple[bytes, str]:
    """Render a template without changing its original file format."""

    extension = Path(file_name).suffix.lower()
    _prepare_template_render_context(context, file_name=file_name, extension=extension)
    if extension == ".docx":
        # Direct renderer callers (tests, template previews and future
        # business lines) do not always build their context through
        # ``_context``.  The source file name is part of the template
        # contract, so retain it here rather than silently bypassing generic
        # note section planning.
        context.setdefault("__source_file_name__", file_name)
        return _render_docx(data, context)
    if extension in {".xlsx", ".xlsm"}:
        return _render_xlsx(data, extension, context)
    if extension == ".xls":
        return _render_xls(data, context)
    if extension == ".doc":
        if context.get("__strict_client_delivery__"):
            raise ValueError(
                "年度审计客户附件不支持旧版 .doc 模板，请使用经确认槽位契约的 DOCX 模板"
            )
        return _render_doc(data, context)
    if extension in {".md", ".markdown", ".txt", ".csv"}:
        return _render_text_template(data, extension, context)
    if extension == ".pdf":
        return _render_pdf(data, context)
    # Never publish an unknown format as if it had been rendered.
    return data, "copied_legacy_or_unknown_format"


def _money(value: Any) -> str:
    """将冻结事实金额格式化为财务附件展示值。"""

    from decimal import Decimal, InvalidOperation

    try:
        amount = Decimal(str(value or 0)).quantize(Decimal("0.01"))
    except (InvalidOperation, ValueError):
        amount = Decimal("0.00")
    return f"{amount:,.2f}"


def _context(snapshot: dict[str, Any], report_text: str, package_version: int, template: dict[str, Any], workpaper: dict[str, Any] | None = None, *, template_type: str = "", engagement_id: int = 0, settings: Settings | None = None, material_index: dict[str, Any] | None = None, source_file_name: str = "") -> dict[str, Any]:
    """构建冻结附件编制上下文。

    提取冻结审计快照中的可用事实，使报告、财务报表、附注和管理
    建议书从同一确定性证据上下文中编制；模板只定义内容范围和设计原型。
    """

    engagement = snapshot.get("engagement") or {}
    fiscal_year = engagement.get("fiscal_year") or snapshot.get("fiscal_year") or ""
    period_end = engagement.get("period_end") or snapshot.get("period_end") or ""
    period_start = engagement.get("period_start") or snapshot.get("period_start") or ""
    if not period_end and fiscal_year:
        period_end = f"{fiscal_year}-12-31"
    if not period_start and fiscal_year:
        period_start = f"{fiscal_year}-01-01"
    period_end_display = str(period_end)
    if re.fullmatch(r"\d{4}-\d{2}-\d{2}", period_end_display):
        period_end_display = f"{period_end_display[0:4]}年{int(period_end_display[5:7])}月{int(period_end_display[8:10])}日"
    period_start_display = str(period_start)
    if re.fullmatch(r"\d{4}-\d{2}-\d{2}", period_start_display):
        period_start_display = f"{period_start_display[0:4]}年{int(period_start_display[5:7])}月{int(period_start_display[8:10])}日"
    issue_year = int(fiscal_year or 0) + 1 if str(fiscal_year).isdigit() else ""
    issue_date = (
        engagement.get("issue_date")
        or engagement.get("report_date")
        or snapshot.get("issue_date")
        or snapshot.get("report_date")
        or ""
    )
    issue_date_display = str(issue_date)
    issue_date_match = re.match(r"^(\d{4})-(\d{2})-(\d{2})", issue_date_display)
    if issue_date_match:
        issue_date_display = (
            f"{issue_date_match.group(1)}年{int(issue_date_match.group(2))}月"
            f"{int(issue_date_match.group(3))}日"
        )
    engagement_name = str(
        engagement.get("name")
        or engagement.get("case_name")
        or snapshot.get("engagement_name")
        or ""
    )
    if not engagement_name and engagement_id and settings:
        with mysql_connection(settings) as connection:
            with connection.cursor() as cursor:
                cursor.execute("SELECT name FROM audit_engagement WHERE id = %s", (engagement_id,))
                engagement_name = str((cursor.fetchone() or {}).get("name") or "")
    report_number_match = re.search(r"京创会审字\[\d{4}\]第\s*\d+号", engagement_name)
    report_number = report_number_match.group(0) if report_number_match else ""
    sales = snapshot.get("sales_receivables") or {}
    revenue = sales.get("revenue") or {}
    receivables = sales.get("receivables") or {}
    cash = snapshot.get("cash_and_bank") or {}
    readiness = snapshot.get("readiness") or {}
    counts = readiness.get("counts") or {}

    # 营业收入明细
    net_revenue = revenue.get("net_revenue")
    gross_revenue = revenue.get("gross_revenue")
    monthly_revenue = revenue.get("monthly_revenue") or []

    # 应收账款明细
    receivables_balance = receivables.get("total_balance")
    aging_buckets = receivables.get("aging_buckets") or {}

    # 货币资金/银行流水明细
    total_inflow = cash.get("total_inflow")
    total_outflow = cash.get("total_outflow")

    # 审计发现汇总
    sales_findings = list(sales.get("findings") or [])
    cash_findings = list(cash.get("findings") or [])
    revenue_findings = list(revenue.get("findings") or [])
    receivables_findings = list(receivables.get("findings") or [])
    all_findings = [*sales_findings, *cash_findings]

    # 构建审计发现摘要文本
    finding_lines: list[str] = []
    risk_names = {"high": "高", "medium": "中", "low": "低"}
    for finding in all_findings:
        risk = risk_names.get(str(finding.get("risk_level") or "").lower(), "待定")
        title = finding.get("title") or "待核查事项"
        desc = finding.get("description") or ""
        amount = finding.get("amount")
        suffix = f"；涉及金额 {_money(amount)} 元" if amount not in (None, "") else ""
        finding_lines.append(f"【{risk}风险】{title}：{desc}{suffix}")
    findings_summary = "\n".join(finding_lines) if finding_lines else "当前已导入数据未触发已配置的确定性风险规则。"
    bounded_result_text = _template_result_text(
        template_type,
        report_text,
        findings_summary,
        snapshot,
    )

    # 结构化表格事实：覆盖模板设计中常见的报表项目。
    table_values = {
        "应收账款": receivables_balance,
        "营业收入": net_revenue,
        "营业收入净额": net_revenue,
        "营业收入总额": gross_revenue,
        "银行流入": total_inflow,
        "银行流出": total_outflow,
    }

    material_index = material_index or {}
    statement_values = filter_unavailable_statement_sources(
        material_index.get("statement_values") or {}
    )

    # 从数据库加载科目余额表，以实际科目余额重建财务报表业务值。
    # 科目余额表是最权威的数据源，优先于快照中的汇总值。
    trial_balance: dict[str, Any] = dict(material_index.get("trial_balance") or {})
    if not trial_balance and engagement_id and settings:
        trial_balance = _load_trial_balance(engagement_id, settings)
    balance_by_name = trial_balance.get("by_name") or {}

    # 将科目余额表中的每个科目名称和余额加入 table_values
    # 科目余额表是最权威的数据源，直接覆盖快照中的汇总值
    for acct_name, acct_data in balance_by_name.items():
        if not acct_name:
            continue
        closing_debit = acct_data.get("closing_debit") or 0
        closing_credit = acct_data.get("closing_credit") or 0
        opening_debit = acct_data.get("opening_debit") or 0
        opening_credit = acct_data.get("opening_credit") or 0
        # 跳过余额全为零的科目
        if closing_debit == 0 and closing_credit == 0 and opening_debit == 0 and opening_credit == 0:
            continue
        # 同一科目可能同时存在借、贷余额列；以净额绝对值呈现，不能按“借方非零
        # 就忽略贷方”的方式制造虚假余额。
        balance_value = abs(closing_debit - closing_credit)
        # 科目余额表数据优先覆盖
        table_values[acct_name] = balance_value
        # 同时提供期初余额
        opening_value = abs(opening_debit - opening_credit)
        table_values[f"{acct_name}_期初"] = opening_value

    # Presentation aliases are standard accounting labels, not a per-template
    # hack.  Do not substitute the cash total for bank deposits: if a source
    # lacks a bank-deposit fact the disclosure remains unsupported instead of
    # silently receiving an incorrect number.
    aliases = {
        "预收款项": "预收账款",
        "实收资本（或股本）": "股本",
    }
    for alias, source_name in aliases.items():
        if alias in table_values or source_name not in balance_by_name:
            continue
        source = balance_by_name[source_name]
        debit = source.get("closing_debit") or 0
        credit = source.get("closing_credit") or 0
        table_values[alias] = abs(debit - credit)
        opening_debit = source.get("opening_debit") or 0
        opening_credit = source.get("opening_credit") or 0
        table_values[f"{alias}_期初"] = abs(opening_debit - opening_credit)

    # Profit/cash-flow/detail workpapers provide line-item values that are not
    # represented in the normalized balance table. They override broad
    # account-label matches (for example, investment income must not inherit
    # the trading-financial-asset balance).
    for label, item in statement_values.items():
        if not isinstance(item, dict) or not any(
            isinstance(item.get(period), (int, float)) for period in ("current", "opening")
        ):
            continue
        if template_type == "notes" and not any(
            abs(float(item.get(key) or 0)) > 1e-9 for key in ("current", "opening")
        ):
            # In note disclosures, an unsupported/zero optional line is kept
            # blank like the reference report; filling every optional line
            # with 0.00 would turn the generic template into a false claim
            # that the line was separately tested and disclosed.
            continue
        if isinstance(item.get("current"), (int, float)):
            table_values[label] = item["current"]
        if isinstance(item.get("opening"), (int, float)):
            table_values[f"{label}_期初"] = item["opening"]

    # Presentation labels in the generic note template differ from the
    # audited statement sheets.  These are deterministic aliases over the
    # same extracted facts, never newly inferred amounts.
    statement_index = statement_values
    for alias, source_name in {
        "银行存款": "银行存款",
        "其他应收款项": "其他应收款",
        "实收资本（或股本）": "股本",
    }.items():
        source = statement_index.get(source_name)
        if not isinstance(source, dict) or not any(
            isinstance(source.get(period), (int, float)) for period in ("current", "opening")
        ):
            continue
        if template_type == "notes" and not any(
            abs(float(source.get(key) or 0)) > 1e-9 for key in ("current", "opening")
        ):
            continue
        if isinstance(source.get("current"), (int, float)):
            table_values[alias] = source["current"]
        if isinstance(source.get("opening"), (int, float)):
            table_values[f"{alias}_期初"] = source["opening"]

    case_summary = snapshot.get("case_workpaper") or (
        snapshot.get("readiness") or {}
    ).get("case_workpaper") or {}
    case_replay_complete = bool(
        snapshot.get("case_pack_complete")
        or (snapshot.get("readiness") or {}).get("case_pack_complete")
        or case_summary.get("is_complete_case")
    )

    fact_statement_values = statement_values
    if template_type == "notes":
        # A 0.00 carried by a generic/formula-broken statement cell is not a
        # basis for retaining an optional disclosure schedule.  Non-zero
        # current or comparative facts remain available, as do normal trial
        # balance facts with a real opening/closing balance.
        fact_statement_values = {
            label: item
            for label, item in fact_statement_values.items()
            if isinstance(item, dict)
            and any(
                isinstance(item.get(period), (int, float))
                and abs(float(item.get(period) or 0)) > 1e-9
                for period in ("current", "opening")
            )
        }

    values: dict[str, Any] = {
        # 被审计单位信息
        "entity_name": engagement.get("entity_name") or snapshot.get("entity_name") or "",
        "entity_uscc": engagement.get("entity_uscc") or snapshot.get("entity_uscc") or "",
        "被审计单位名称": engagement.get("entity_name") or snapshot.get("entity_name") or "",
        "被审计单位": engagement.get("entity_name") or snapshot.get("entity_name") or "",
        # 审计期间
        "audit_period": f"{fiscal_year}年度" if fiscal_year else "",
        "fiscal_year": fiscal_year,
        "period_end": period_end_display,
        "period_start": period_start_display,
        "审计期间": f"{period_start_display} 至 {period_end_display}" if period_start_display and period_end_display else "",
        "资产负债日": period_end_display,
        "issue_year": issue_year,
        "issue_date": issue_date_display,
        "审计报告日期": issue_date_display,
        # 项目信息
        "engagement_code": engagement.get("engagement_code") or snapshot.get("engagement_code") or "",
        "项目编号": engagement.get("engagement_code") or snapshot.get("engagement_code") or "",
        "report_number": report_number,
        "审计报告文号": report_number,
        "engagement_partner": engagement.get("engagement_partner") or "",
        "项目合伙人": engagement.get("engagement_partner") or "",
        # 审计结果
        # Explicit template fields receive a bounded, usage-specific value.
        # The conversational report remains available to the chat layer but
        # is never copied wholesale into a customer-facing attachment.
        "report_text": bounded_result_text,
        "audit_result": bounded_result_text,
        "审计结果": bounded_result_text,
        "审计发现摘要": findings_summary,
        "findings_summary": findings_summary,
        # 财务数据
        "net_revenue": _money(net_revenue) if net_revenue is not None else "",
        "营业收入净额": _money(net_revenue) if net_revenue is not None else "",
        "gross_revenue": _money(gross_revenue) if gross_revenue is not None else "",
        "营业收入总额": _money(gross_revenue) if gross_revenue is not None else "",
        "receivables_balance": _money(receivables_balance) if receivables_balance is not None else "",
        "应收账款余额": _money(receivables_balance) if receivables_balance is not None else "",
        "total_inflow": _money(total_inflow) if total_inflow is not None else "",
        "银行流入合计": _money(total_inflow) if total_inflow is not None else "",
        "total_outflow": _money(total_outflow) if total_outflow is not None else "",
        "银行流出合计": _money(total_outflow) if total_outflow is not None else "",
        # 资料就绪度
        "account_balance_rows": int(counts.get("account_balance_rows") or 0),
        "journal_entry_rows": int(counts.get("journal_entry_rows") or 0),
        "receivable_rows": int(counts.get("receivable_rows") or 0),
        "bank_transaction_rows": int(counts.get("bank_transaction_rows") or 0),
        # 元数据
        "report_version": snapshot.get("report_version") or "",
        "template_version": template_version_ref(template),
        "package_version": package_version,
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "__template_type__": template_type,
        "__table_values__": table_values,
        "__fact_registry__": _build_periodised_fact_registry(
            table_values,
            statement_values=fact_statement_values,
            balance_by_name=balance_by_name,
        ),
        "__case_workpaper_replay__": case_replay_complete,
        "__case_material_index__": material_index,
        "__source_file_name__": source_file_name,
    }
    if workpaper:
        values.update(
            {
                "workpaper_code": workpaper.get("code") or "",
                "workpaper_name": workpaper.get("name") or "",
                "工作底稿编号": workpaper.get("code") or "",
                "工作底稿名称": workpaper.get("name") or "",
            }
        )
    values["__field_schema__"] = template.get("field_schema") or {}
    return values


def _safe_filename_part(value: Any, fallback: str) -> str:
    """Return a readable filename segment that is safe on Windows and POSIX."""

    text = _INVALID_FILENAME_CHARS_RE.sub("", str(value or "")).strip().strip(".")
    return text or fallback


def _content_type_for_file_name(file_name: str, fallback: str = "") -> str:
    """Derive content type from the output extension, not the upload hint."""

    suffix = Path(file_name).suffix.lower()
    return _CONTENT_TYPES_BY_SUFFIX.get(
        suffix,
        mimetypes.guess_type(file_name)[0] or fallback or "application/octet-stream",
    )


def _validate_rendered_file(data: bytes, file_name: str) -> dict[str, Any]:
    """Validate the container signature before an artifact is published."""

    extension = Path(file_name).suffix.lower()
    if not data:
        raise ValueError("渲染结果为空")
    if extension in {".docx", ".xlsx", ".xlsm"}:
        if not zipfile.is_zipfile(BytesIO(data)):
            raise ValueError(f"{extension} 输出不是有效的 OOXML 文件")
        with zipfile.ZipFile(BytesIO(data)) as archive:
            names = set(archive.namelist())
            required = "word/document.xml" if extension == ".docx" else "xl/workbook.xml"
            if required not in names or "[Content_Types].xml" not in names:
                raise ValueError(f"{extension} 输出缺少必要的 OOXML 部件")
    elif extension in {".doc", ".xls"}:
        if not data.startswith(b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"):
            raise ValueError(f"{extension} 输出不是有效的 OLE/BIFF 文件")
    elif extension == ".pdf":
        if not data.startswith(b"%PDF"):
            raise ValueError("PDF 输出缺少 %PDF 文件签名")
    elif extension in {".md", ".markdown", ".txt", ".csv"}:
        data.decode("utf-8-sig")
    else:
        raise ValueError(f"不支持校验的附件格式：{extension or file_name}")
    return {
        "extension": extension,
        "sha256": hashlib.sha256(data).hexdigest(),
        "size": len(data),
        "signature_valid": True,
    }


def _materialized_docvariable_target_paths(
    document: Any,
    render_context: dict[str, Any] | None,
    materialized_slot_ids: set[str],
) -> dict[str, set[str]]:
    """Locate only the source paragraphs whose DOCVARIABLEs were approved.

    The fidelity comparison must not ignore every DOCVARIABLE in a document
    merely because one header identity slot was materialised.  These paths are
    computed from the frozen contract and exact source document; a resolution
    failure stays visible as a failed format check.
    """

    if not materialized_slot_ids:
        return {}
    contract = (render_context or {}).get("__attachment_slot_contract__") or {}
    if not isinstance(contract, dict):
        return {}
    result: dict[str, set[str]] = {}
    for raw_slot in contract.get("slots") or []:
        if not isinstance(raw_slot, dict):
            continue
        if str(raw_slot.get("slot_id") or "") not in materialized_slot_ids:
            continue
        if not bool(raw_slot.get("materializes_docvariable_fields")):
            continue
        target = str(raw_slot.get("target") or "")
        paragraph: Any | None = None
        part_name = "word/document.xml"
        body_match = re.fullmatch(r"paragraph:(?P<index>\d+)", target)
        frame_match = re.fullmatch(
            r"(?P<frame>header|footer):(?P<section>\d+),paragraph:(?P<paragraph>\d+)",
            target,
            re.IGNORECASE,
        )
        table_match = re.fullmatch(
            r"table:(?P<table>\d+),row:(?P<row>\d+),cell:(?P<cell>\d+)",
            target,
            re.IGNORECASE,
        )
        try:
            if body_match:
                paragraph = document.paragraphs[int(body_match.group("index"))]
            elif frame_match:
                section = document.sections[int(frame_match.group("section"))]
                frame = getattr(section, frame_match.group("frame").lower())
                paragraph = frame.paragraphs[int(frame_match.group("paragraph"))]
                part_name = str(frame.part.partname).lstrip("/")
            elif table_match:
                table = document.tables[int(table_match.group("table"))]
                cell = table.rows[int(table_match.group("row"))].cells[
                    int(table_match.group("cell"))
                ]
                if len(cell.paragraphs) == 1:
                    paragraph = cell.paragraphs[0]
        except (AttributeError, IndexError, KeyError, TypeError, ValueError):
            paragraph = None
        if paragraph is None:
            continue
        result.setdefault(part_name, set()).add(
            paragraph._p.getroottree().getpath(paragraph._p)
        )
    return result


def _strip_docx_docvariable_field_geometry(
    root: Any,
    *,
    target_paragraph_paths: set[str],
) -> None:
    """Remove only DOCVARIABLE field mechanics from an in-memory XML tree.

    A declared header slot may convert a DOCVARIABLE result into literal text
    so that Word refreshes cannot restore a sample entity or period.  That
    conversion is a content change, not a visual-layout change.  The caller
    supplies only exact bound paragraph paths; PAGE, REF, every unbound
    DOCVARIABLE and all other fields remain in the fingerprint.
    """

    namespace = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    field_tag = f"{{{namespace}}}fldChar"
    instruction_tag = f"{{{namespace}}}instrText"
    simple_field_tag = f"{{{namespace}}}fldSimple"
    field_type_attr = f"{{{namespace}}}fldCharType"
    simple_instruction_attr = f"{{{namespace}}}instr"
    paragraph_tag = f"{{{namespace}}}p"
    tree = root.getroottree()

    def is_target_paragraph(node: Any) -> bool:
        current = node
        while current is not None and current.tag != paragraph_tag:
            current = current.getparent()
        return current is not None and tree.getpath(current) in target_paragraph_paths

    # Normalize simple fields to their cached result runs before handling
    # complex fields.  This mirrors the materialisation performed by the
    # renderer without changing the source package itself.
    for simple_field in list(root.iter(simple_field_tag)):
        if not is_target_paragraph(simple_field):
            continue
        instruction = str(simple_field.get(simple_instruction_attr) or "")
        if not re.match(r"^\s*DOCVARIABLE(?:\s|$)", instruction, re.IGNORECASE):
            continue
        parent = simple_field.getparent()
        if parent is None:
            continue
        position = parent.index(simple_field)
        for child in list(simple_field):
            simple_field.remove(child)
            parent.insert(position, child)
            position += 1
        parent.remove(simple_field)

    removable: list[Any] = []
    stack: list[dict[str, Any]] = []
    for node in root.iter():
        if node.tag == field_tag:
            field_type = str(node.get(field_type_attr) or "").lower()
            if field_type == "begin":
                stack.append(
                    {
                        "nodes": [node],
                        "instructions": [],
                        "allowed": is_target_paragraph(node),
                    }
                )
            elif stack:
                stack[-1]["nodes"].append(node)
                if field_type == "end":
                    field = stack.pop()
                    command = "".join(
                        str(item.text or "") for item in field["instructions"]
                    )
                    if field["allowed"] and re.match(
                        r"^\s*DOCVARIABLE(?:\s|$)", command, re.IGNORECASE
                    ):
                        removable.extend(field["nodes"])
                        removable.extend(field["instructions"])
        elif node.tag == instruction_tag and stack:
            stack[-1]["instructions"].append(node)
    for node in removable:
        parent = node.getparent()
        if parent is not None:
            parent.remove(node)


def _normalized_docx_layout_xml(
    value: bytes,
    *,
    materialized_docvariable_target_paths: set[str] | None = None,
) -> bytes:
    """Return Word layout XML with text and authoring ink removed.

    This is deliberately narrower than a file hash: client values may change,
    but a template-fidelity render must leave body order, sections, table
    geometry, paragraph properties, run properties, drawings and headers in
    the same structural form.
    """

    from lxml import etree

    word_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    root = etree.fromstring(value)
    if materialized_docvariable_target_paths:
        _strip_docx_docvariable_field_geometry(
            root,
            target_paragraph_paths=materialized_docvariable_target_paths,
        )
    for element in root.iter():
        local_name = etree.QName(element).localname
        # Field instructions are document behavior (PAGE, REF, formulas), not
        # customer text.  Retain them in the layout fingerprint so an
        # undeclared field cannot be changed or deleted behind a text-only
        # fidelity comparison.  Approved DOCVARIABLE instructions are removed
        # earlier and only at their exact bound paragraph paths.
        if local_name in {"t", "delText"}:
            element.text = ""
            element.attrib.pop("{http://www.w3.org/XML/1998/namespace}space", None)
        for attribute in list(element.attrib):
            if etree.QName(attribute).localname.lower().startswith("rsid"):
                del element.attrib[attribute]
    namespaces = {"w": word_ns}
    for run_properties in root.xpath(".//w:rPr", namespaces=namespaces):
        for marker in run_properties.xpath("./w:highlight | ./w:shd", namespaces=namespaces):
            run_properties.remove(marker)
    for marker in root.xpath(
        ".//w:commentRangeStart | .//w:commentRangeEnd | .//w:commentReference",
        namespaces=namespaces,
    ):
        parent = marker.getparent()
        if parent is not None:
            parent.remove(marker)
    # A value written into a previously empty table cell necessarily adds a
    # text-only run.  It inherits the paragraph's template typography, so it
    # is data rather than layout geometry.  Ignore such carriers while still
    # retaining every run that has explicit formatting or field/drawing/tab
    # structure of its own.
    run_tag = f"{{{word_ns}}}r"
    text_tag = f"{{{word_ns}}}t"
    for run in list(root.iter(run_tag)):
        if all(child.tag == text_tag for child in run):
            parent = run.getparent()
            if parent is not None:
                parent.remove(run)
    return etree.tostring(root, method="c14n")


def _docx_layout_part_hashes(
    data: bytes,
    *,
    materialized_docvariable_target_paths: dict[str, set[str]] | None = None,
) -> dict[str, str]:
    """Fingerprint all layout-bearing Word parts, ignoring visible text."""

    with zipfile.ZipFile(BytesIO(data)) as archive:
        parts = {
            name: _normalized_docx_layout_xml(
                archive.read(name),
                materialized_docvariable_target_paths=(
                    materialized_docvariable_target_paths or {}
                ).get(name),
            )
            for name in archive.namelist()
            if name == "word/document.xml"
            or name.startswith("word/header") and name.endswith(".xml")
            or name.startswith("word/footer") and name.endswith(".xml")
        }
    return {name: hashlib.sha256(value).hexdigest() for name, value in parts.items()}


def _docx_preserve_only_part_hashes(data: bytes) -> dict[str, str]:
    """Fingerprint design-system parts without serializer-only OOXML noise.

    ``python-docx`` rewrites XML namespace/attribute order and relationship
    serialisation even when a part's visible design is untouched.  XML parts
    therefore use the same canonical form as the layout check; binary assets
    retain an exact byte hash.
    """

    prefixes = ("word/media/", "word/theme/", "word/embeddings/", "customXml/")
    exact = {
        "word/styles.xml",
        "word/numbering.xml",
        "word/fontTable.xml",
        "word/settings.xml",
        "word/webSettings.xml",
    }
    with zipfile.ZipFile(BytesIO(data)) as archive:
        fingerprints: dict[str, str] = {}
        for name in archive.namelist():
            if name not in exact and not name.startswith(prefixes):
                continue
            value = archive.read(name)
            if name.endswith((".xml", ".rels")):
                try:
                    value = _normalized_docx_layout_xml(value)
                except Exception:
                    # An extension can be XML-like without containing a
                    # parsable XML payload.  In that case, retain byte-level
                    # protection rather than silently excluding it.
                    pass
            fingerprints[name] = hashlib.sha256(value).hexdigest()
        return fingerprints


def _validate_template_format_fidelity(
    source_data: bytes,
    rendered_data: bytes,
    *,
    source_file_name: str,
    render_context: dict[str, Any] | None = None,
) -> dict[str, Any]:
    """Verify that a rendering retained the uploaded template's format frame.

    Content is expected to change; styles, sheet/page structure and existing
    table slots are not.  The check deliberately reports structural evidence
    rather than comparing whole files, because OOXML packages contain changing
    timestamps and relationship IDs even for a faithful in-place edit.
    """

    extension = Path(source_file_name).suffix.lower()
    if extension == ".docx":
        from docx import Document

        source = Document(BytesIO(source_data))
        rendered = Document(BytesIO(rendered_data))
        source_style_names = {style.name for style in source.styles}
        rendered_style_names = {style.name for style in rendered.styles}
        with zipfile.ZipFile(BytesIO(source_data)) as source_zip, zipfile.ZipFile(BytesIO(rendered_data)) as rendered_zip:
            source_parts = set(source_zip.namelist())
            rendered_parts = set(rendered_zip.namelist())
        required_parts = {"word/document.xml", "word/styles.xml", "[Content_Types].xml"}
        source_headers_footers = {
            part for part in source_parts if part.startswith("word/header") or part.startswith("word/footer")
        }
        preserved_headers_footers = source_headers_footers <= rendered_parts
        style_coverage = source_style_names <= rendered_style_names
        materialized_docvariable_slot_ids = (
            (render_context or {}).get("__docx_materialized_docvariable_slot_ids__")
            or []
        )
        if isinstance(materialized_docvariable_slot_ids, (str, bytes, bytearray)):
            materialized_docvariable_slot_ids = [materialized_docvariable_slot_ids]
        materialized_docvariable_slot_ids = {
            str(item)
            for item in materialized_docvariable_slot_ids
            if str(item).strip()
        }
        materialized_docvariable_paths = _materialized_docvariable_target_paths(
            source,
            render_context,
            materialized_docvariable_slot_ids,
        )
        materialized_docvariable_target_count = sum(
            len(paths) for paths in materialized_docvariable_paths.values()
        )
        materialized_docvariable_targets_resolved = (
            materialized_docvariable_target_count == len(materialized_docvariable_slot_ids)
        )
        source_layout_hashes = _docx_layout_part_hashes(
            source_data,
            materialized_docvariable_target_paths=materialized_docvariable_paths,
        )
        rendered_layout_hashes = _docx_layout_part_hashes(
            rendered_data,
            materialized_docvariable_target_paths=materialized_docvariable_paths,
        )
        layout_parts_preserved = source_layout_hashes == rendered_layout_hashes
        source_preserve_only_hashes = _docx_preserve_only_part_hashes(source_data)
        rendered_preserve_only_hashes = _docx_preserve_only_part_hashes(rendered_data)
        preserve_only_parts_unchanged = source_preserve_only_hashes == rendered_preserve_only_hashes
        same_body_child_count = len(source._element.body) == len(rendered._element.body)
        same_paragraph_count = len(source.paragraphs) == len(rendered.paragraphs)
        same_table_count = len(source.tables) == len(rendered.tables)
        same_section_count = len(source.sections) == len(rendered.sections)
        passed = (
            required_parts <= rendered_parts
            and style_coverage
            and preserved_headers_footers
            and layout_parts_preserved
            and materialized_docvariable_targets_resolved
            and preserve_only_parts_unchanged
            and same_body_child_count
            and same_paragraph_count
            and same_table_count
            and same_section_count
        )
        return {
            "format_contract_checked": True,
            "format_contract_passed": passed,
            "kind": "docx",
            "source_paragraph_count": len(source.paragraphs),
            "output_paragraph_count": len(rendered.paragraphs),
            "source_table_count": len(source.tables),
            "output_table_count": len(rendered.tables),
            "source_style_count": len(source_style_names),
            "output_style_count": len(rendered_style_names),
            "style_coverage": style_coverage,
            "header_footer_parts_preserved": preserved_headers_footers,
            "source_body_child_count": len(source._element.body),
            "output_body_child_count": len(rendered._element.body),
            "source_section_count": len(source.sections),
            "output_section_count": len(rendered.sections),
            "body_child_count_preserved": same_body_child_count,
            "paragraph_count_preserved": same_paragraph_count,
            "table_count_preserved": same_table_count,
            "section_count_preserved": same_section_count,
            "layout_parts_preserved": layout_parts_preserved,
            "docvariable_materialization_slot_count": len(
                materialized_docvariable_slot_ids
            ),
            "docvariable_materialization_target_count": materialized_docvariable_target_count,
            "docvariable_materialization_targets_resolved": materialized_docvariable_targets_resolved,
            "preserve_only_parts_unchanged": preserve_only_parts_unchanged,
            "checked_layout_part_count": len(source_layout_hashes),
            "checked_preserve_only_part_count": len(source_preserve_only_hashes),
        }
    if extension in {".xlsx", ".xlsm"}:
        from openpyxl import load_workbook

        source = load_workbook(BytesIO(source_data), read_only=False, data_only=False, keep_vba=extension == ".xlsm")
        rendered = load_workbook(BytesIO(rendered_data), read_only=False, data_only=False, keep_vba=extension == ".xlsm")
        same_sheet_order = source.sheetnames == rendered.sheetnames
        source_merges = {
            sheet.title: sorted(str(item) for item in sheet.merged_cells.ranges)
            for sheet in source.worksheets
        }
        rendered_merges = {
            sheet.title: sorted(str(item) for item in sheet.merged_cells.ranges)
            for sheet in rendered.worksheets
        }
        same_merges = source_merges == rendered_merges
        # openpyxl exposes the registered style records even when a cell does
        # not use all of them.  The renderer writes values only, so a loss of
        # registered styles is a real template-fidelity regression.
        style_count_preserved = len(rendered._cell_styles) >= len(source._cell_styles)
        passed = same_sheet_order and same_merges and style_count_preserved
        return {
            "format_contract_checked": True,
            "format_contract_passed": passed,
            "kind": "xlsx",
            "source_sheet_count": len(source.worksheets),
            "output_sheet_count": len(rendered.worksheets),
            "sheet_order_preserved": same_sheet_order,
            "merged_ranges_preserved": same_merges,
            "source_style_count": len(source._cell_styles),
            "output_style_count": len(rendered._cell_styles),
            "style_count_preserved": style_count_preserved,
        }
    return {"format_contract_checked": False}


def _validate_rendered_template_content(
    data: bytes,
    *,
    template_type: str,
    source_file_name: str,
    render_context: dict[str, Any] | None = None,
) -> dict[str, Any]:
    """Reject a successful-looking render that still contains template text."""

    if Path(source_file_name).suffix.lower() != ".docx":
        return {"content_checked": False}
    from docx import Document

    document = Document(BytesIO(data))
    story_parts = _docx_ooxml_text_parts(data)
    body_content = _docx_document_text(document)
    content = "\n".join(story_parts.values()) or body_content
    strict_instruction_cleanliness = bool(
        isinstance(render_context, dict)
        and render_context.get("__strict_client_delivery__")
    )
    marker_content = content if strict_instruction_cleanliness else body_content
    unresolved_markers = len(
        re.findall(
            r"X{2,}|x{2,}|20XX|一般企业模板|\{\{[^}]+\}\}|\[\[[^\]]+\]\]",
            marker_content,
        )
    )
    instruction_source = content if strict_instruction_cleanliness else body_content
    instruction_markers = sum(
        instruction_source.count(marker) for marker in _GENERIC_NOTE_INSTRUCTION_MARKERS
    )
    instruction_markers += instruction_source.count("[或适用]")
    missing_outer_border_tables = 0
    if template_type == "notes":
        missing_outer_border_tables = sum(
            1
            for table in _iter_docx_tables(document)
            if not all(_docx_table_has_outer_borders(table))
        )
    render_contract = (render_context or {}).get("__template_render_contract__") or {}
    requires_annotation_cleanliness = strict_instruction_cleanliness or (
        isinstance(render_contract, dict)
        and render_contract.get("authoring_annotation_policy") == "clear_template_authoring_marks"
    )
    declared_slot_contract = (
        (render_context or {}).get("__attachment_slot_contract__") or {}
    )
    if strict_instruction_cleanliness and not (
        isinstance(declared_slot_contract, dict)
        and declared_slot_contract.get("declared")
        and declared_slot_contract.get("slots")
    ):
        raise ValueError("DOCX 正式交付缺少经确认的 slot_contract")
    authoring_annotation_count = _docx_authoring_annotation_count(document)
    internal_trace_matches = sorted(
        {
            match.group(0)
            for pattern in _CLIENT_DELIVERY_TRACE_PATTERNS
            for match in pattern.finditer(content)
        }
    )
    known_evidence_ids = (
        (render_context or {}).get("__attachment_known_evidence_ids__") or set()
    )
    if isinstance(known_evidence_ids, (str, bytes, bytearray)):
        known_evidence_ids = [known_evidence_ids]
    if isinstance(known_evidence_ids, (list, tuple, set, frozenset)):
        internal_trace_matches.extend(
            f"known-evidence-id:{value}"
            for value in known_evidence_ids
            if len(str(value).strip()) >= 4 and str(value) in content
        )
    internal_trace_matches = sorted(set(internal_trace_matches))
    if (
        unresolved_markers
        or (
            strict_instruction_cleanliness
            and instruction_markers
        )
        or (requires_annotation_cleanliness and authoring_annotation_count)
        or internal_trace_matches
    ):
        raise ValueError(
            f"{template_type} 渲染后仍含模板占位、说明文字、内部追溯信息或作者批注："
            f"占位符 {unresolved_markers} 个，说明文字 {instruction_markers} 个，"
            f"左右边框缺失表格 {missing_outer_border_tables} 个，内部追溯标记 {len(internal_trace_matches)} 个，"
            f"作者批注 {authoring_annotation_count} 个"
        )
    return {
        "content_checked": True,
        "paragraph_count": len(document.paragraphs),
        "table_count": len(document.tables),
        "unresolved_marker_count": unresolved_markers,
        "instruction_marker_count": instruction_markers,
        "strict_instruction_cleanliness": strict_instruction_cleanliness,
        "checked_story_part_count": len(story_parts),
        "missing_outer_border_table_count": missing_outer_border_tables,
        "internal_trace_marker_count": len(internal_trace_matches),
        "authoring_annotation_count": authoring_annotation_count,
    }


def _output_name(
    template_type: str,
    engagement_id: int,
    package_version: int,
    source_name: str,
    workpaper_code: str = "",
    output_suffix: str | None = None,
    *,
    context: dict[str, Any] | None = None,
    duplicate_index: int = 0,
    duplicate_count: int = 1,
) -> str:
    """构建用户可读的附件文件名，保留源文件格式。

    命名规则参考 new_docs 案例目录：
    - 年度审计报告：``{被审计单位}{年度}年度审计报告.docx``
    - 财务报表：``{被审计单位}{年度}年度审计财务报表.xlsx``
    - 财务报表附注：``{被审计单位}{年度}财务报表附注.docx``
    - 管理建议书：``{被审计单位}{年度}管理建议书.docx``
    - 审计工作底稿：``{被审计单位}{年度}审计工作底稿-{底稿编号}.xlsx``
    - 函证：``{被审计单位}{年度}函证-{函证类型}.docx``

    扩展名始终来自实际源模板文件，``.xls`` 模板保持 ``.xls`` 输出，
    绝不会重命名为 ``.xlsx``。
    """

    suffix = output_suffix or Path(source_name).suffix.lower()
    if suffix and not suffix.startswith("."):
        suffix = f".{suffix}"
    source_stem = _safe_filename_part(Path(source_name).stem, "模板文件")
    context = context or {}
    entity_name = _safe_filename_part(context.get("entity_name"), f"项目{engagement_id}")
    year_match = re.search(r"20\d{2}", str(context.get("fiscal_year") or ""))
    fiscal_year = year_match.group(0) if year_match else ""
    year_part = f"{fiscal_year}年度" if fiscal_year else "年度"

    if template_type == "annual_report":
        label = "审计报告"
        base = f"{entity_name}{year_part}{label}"
        if duplicate_count > 1:
            base += f"-{source_stem}"
    elif template_type == "financial_statements":
        label = "财务报表"
        base = f"{entity_name}{year_part}{label}"
        if duplicate_count > 1:
            base += f"-{source_stem}"
    elif template_type == "notes":
        label = "财务报表附注"
        base = f"{entity_name}{year_part}{label}"
        if duplicate_count > 1:
            base += f"-{source_stem}"
    elif template_type == "management_letter":
        label = "管理建议书"
        base = f"{entity_name}{year_part}{label}"
        if duplicate_count > 1:
            base += f"-{source_stem}"
    elif template_type == "audit_workpaper":
        label = "审计工作底稿"
        base = f"{entity_name}{year_part}{label}"
        if workpaper_code:
            normalized_code = _safe_filename_part(workpaper_code, "底稿")
            # Most supplied workpaper names already start with the code, e.g.
            # ``4101 货币资金.xls``.  Keep that descriptive source stem once
            # instead of producing ``...-4101-4101 货币资金.xls``.
            base += f"-{source_stem if source_stem.lower().startswith(normalized_code.lower()) else normalized_code}"
        elif duplicate_count > 1:
            base += f"-{source_stem}"
    elif template_type == "confirmations":
        label = "函证"
        base = f"{entity_name}{year_part}{label}"
        if duplicate_count > 1 or source_stem not in {"模板", "函证模板"}:
            base += f"-{source_stem}"
    else:
        label = _ATTACHMENT_NAME_LABELS.get(template_type, _safe_filename_part(template_type, "审计附件"))
        base = f"{entity_name}{year_part}{label}"
        if duplicate_count > 1:
            base += f"-{source_stem}"

    if duplicate_index > 0 and duplicate_count <= 1:
        base += f"-{duplicate_index + 1}"
    return f"{_safe_filename_part(base, '审计附件')}{suffix}"


def _infer_workpaper_code(file_name: str) -> str:
    match = re.search(r"(?<![A-Za-z0-9])((?:[A-Z]\d{1,4}|\d{3,4})(?:-\d+)?)(?![A-Za-z0-9])", Path(file_name).stem, re.IGNORECASE)
    return match.group(1).upper() if match else ""


def _result_placement_description(template_type: str, source_name: str) -> str:
    """Describe where structured audit values are allowed to land."""

    extension = Path(source_name).suffix.lower()
    if extension == ".docx":
        return "仅在经确认的段落、表格、DOCVARIABLE 或文本槽位原位填充；不重建正文、不新增段落或表格、不写入证据追溯信息"
    if template_type == "annual_report":
        return "模板显式结果字段或字段映射；未定义承载位置时阻断交付"
    if template_type == "financial_statements":
        return "按模板工作表、公式和已声明科目单元格填入冻结结构化事实"
    if template_type == "notes":
        return "模板显式结果字段或字段映射；未定义承载位置时阻断交付"
    if template_type == "management_letter":
        return "模板显式结果字段或字段映射；未定义承载位置时阻断交付"
    if extension in {".md", ".markdown", ".txt", ".csv"}:
        return "模板显式结果字段；不追加通用审计正文"
    if extension == ".pdf":
        return "模板已有 AcroForm 字段；不追加新页面"
    return "模板显式字段或结构化表格位置"


class AttachmentItemGenerationError(RuntimeError):
    """Stage-aware error consumed by the durable attachment runner."""

    def __init__(self, message: str, *, stage: str, retryable: bool = False) -> None:
        super().__init__(message)
        self.stage = stage
        self.retryable = retryable


def _planned_generation_item(
    preflight_plan: dict[str, Any],
    generation_item: dict[str, Any],
) -> dict[str, Any]:
    template_type = str(generation_item.get("attachment_type") or generation_item.get("template_type") or "")
    file_name = str(generation_item.get("source_file_name") or generation_item.get("file_name") or "")
    file_id = int(generation_item.get("template_file_id") or 0)
    source_sha256 = str(generation_item.get("source_sha256") or generation_item.get("source_template_sha256") or "")
    matches = []
    for raw in preflight_plan.get("templates") or []:
        if not isinstance(raw, dict):
            continue
        if str(raw.get("template_type") or "") != template_type:
            continue
        if file_id and int(raw.get("template_file_id") or 0) != file_id:
            continue
        if file_name and str(raw.get("file_name") or "") != file_name:
            continue
        if source_sha256 and str(raw.get("source_template_sha256") or "") != source_sha256:
            continue
        matches.append(raw)
    if len(matches) != 1:
        raise AttachmentItemGenerationError(
            "冻结子任务无法唯一匹配模板文件身份",
            stage="frozen_input_validation",
            retryable=False,
        )
    return dict(matches[0])


def generate_annual_attachment_item(
    engagement_id: int,
    *,
    generation_item: dict[str, Any],
    preflight_plan: dict[str, Any],
    reserved_package_id: int | None = None,
    reserved_package_version: int,
    created_by: str = "ai_agent",
    generation_run_id: str = "",
    settings: Settings | None = None,
) -> dict[str, Any]:
    """Generate exactly one frozen template file without replanning.

    This is the only entry used by durable child workers.  It deliberately
    never calls ``generate_annual_report_draft``, ``_load_latest_report`` or
    ``get_active_template_catalog``.  Repeated attempts therefore consume the
    same report, material evidence, template bytes identity and blueprint.
    """

    resolved = settings or get_settings()
    if int(preflight_plan.get("engagement_id") or 0) != engagement_id:
        raise AttachmentItemGenerationError(
            "冻结计划与当前审计项目不一致",
            stage="frozen_input_validation",
            retryable=False,
        )
    frozen_report = preflight_plan.get("frozen_report") or {}
    snapshot = dict(frozen_report.get("fact_snapshot") or {})
    report_text = str(frozen_report.get("report_text") or "")
    if not snapshot or not report_text:
        raise AttachmentItemGenerationError(
            "冻结审计报告输入不完整",
            stage="frozen_input_validation",
            retryable=False,
        )
    expected_snapshot_hash = str(frozen_report.get("snapshot_sha256") or "")
    if expected_snapshot_hash and expected_snapshot_hash != _stable_snapshot_hash(snapshot):
        raise AttachmentItemGenerationError(
            "冻结审计事实快照哈希不一致",
            stage="frozen_input_validation",
            retryable=False,
        )

    planned = _planned_generation_item(preflight_plan, generation_item)
    template_type = str(planned.get("template_type") or generation_item.get("attachment_type") or "")
    source_name = str(planned.get("file_name") or generation_item.get("source_file_name") or "template")
    source_storage_ref = str(
        generation_item.get("source_storage_ref")
        or planned.get("source_storage_ref")
        or ""
    )
    source_sha256 = str(
        generation_item.get("source_sha256")
        or planned.get("source_template_sha256")
        or ""
    )
    if not source_storage_ref or not source_sha256:
        raise AttachmentItemGenerationError(
            "冻结模板文件引用不完整",
            stage="frozen_input_validation",
            retryable=False,
        )
    try:
        source_bytes = get_minio_service().get_object_bytes(source_storage_ref)
    except Exception as exc:
        raise AttachmentItemGenerationError(
            f"读取冻结模板文件失败：{str(exc)[:300]}",
            stage="template_read",
            retryable=True,
        ) from exc
    if hashlib.sha256(source_bytes).hexdigest() != source_sha256:
        raise AttachmentItemGenerationError(
            "模板对象内容与冻结 SHA-256 不一致",
            stage="frozen_input_validation",
            retryable=False,
        )

    template = dict(planned.get("template_snapshot") or {})
    item_template = generation_item.get("template_snapshot") or {}
    if isinstance(item_template, str):
        try:
            item_template = json.loads(item_template)
        except json.JSONDecodeError:
            item_template = {}
    if isinstance(item_template, dict):
        nested = item_template.get("template")
        if isinstance(nested, dict):
            template.update(nested)
        else:
            template.update(
                {
                    key: value
                    for key, value in item_template.items()
                    if key in {
                        "template_code",
                        "template_type",
                        "version_no",
                        "version_label",
                        "content_hash",
                        "content",
                        "field_schema",
                        "template_contract",
                    }
                }
            )
    template.setdefault("template_type", template_type)

    ordinal = int(generation_item.get("ordinal_no") or 0)
    code = _infer_workpaper_code(source_name) if template_type == "audit_workpaper" else ""
    workpaper = {"code": code, "name": source_name} if code else None
    render_context = _context(
        snapshot,
        report_text,
        int(reserved_package_version),
        template,
        workpaper,
        template_type=template_type,
        engagement_id=engagement_id,
        settings=resolved,
        material_index=dict(preflight_plan.get("material_index") or {}),
        source_file_name=source_name,
    )
    blueprint = planned.get("blueprint") or {}
    extension = Path(source_name).suffix.lower()
    authored: dict[str, Any] = {}
    if extension in {".docx", ".md", ".markdown", ".txt"}:
        frozen_authoring_context = planned.get("frozen_authoring_context") or {}
        if not blueprint or not frozen_authoring_context:
            raise AttachmentItemGenerationError(
                "完整附件缺少冻结蓝图或章节证据",
                stage="authoring_precondition",
                retryable=False,
            )
        try:
            from .attachment_document_author_service import (
                AttachmentDocumentAuthoringUnavailable,
                AttachmentDocumentValidationError,
                attach_authored_document_to_context,
                author_attachment_document,
            )

            authored = author_attachment_document(
                blueprint=blueprint,
                frozen_context=frozen_authoring_context,
                settings=resolved,
            )
            render_context["__attachment_blueprint__"] = dict(blueprint)
            blueprint_structure = blueprint.get("structure") or {}
            if isinstance(blueprint_structure, dict) and isinstance(
                blueprint_structure.get("slot_contract"), dict
            ):
                render_context["__attachment_slot_contract__"] = dict(
                    blueprint_structure["slot_contract"]
                )
            attach_authored_document_to_context(render_context, authored)
        except Exception as exc:
            raise AttachmentItemGenerationError(
                str(exc)[:1000],
                stage="ai_complete_document_authoring",
                retryable=not isinstance(
                    exc,
                    (
                        AttachmentDocumentAuthoringUnavailable,
                        AttachmentDocumentValidationError,
                        ValueError,
                    ),
                ),
            ) from exc
    elif template_type == "financial_statements" and extension in {".xlsx", ".xlsm"}:
        from .attachment_financial_author_service import (
            FinancialAttachmentAuthoringUnavailable,
            FinancialAttachmentValidationError,
        )

        try:
            _prepare_financial_semantic_authoring(source_bytes, extension, render_context)
        except Exception as exc:
            raise AttachmentItemGenerationError(
                str(exc)[:1000],
                stage="ai_financial_semantic_planning",
                retryable=not isinstance(
                    exc,
                    (
                        FinancialAttachmentAuthoringUnavailable,
                        FinancialAttachmentValidationError,
                        ValueError,
                    ),
                ),
            ) from exc

    output_name = str(generation_item.get("output_file_name") or "") or _output_name(
        template_type,
        engagement_id,
        int(reserved_package_version),
        source_name,
        code,
        context=render_context,
        duplicate_index=int(generation_item.get("duplicate_index") or 0),
        duplicate_count=max(int(generation_item.get("duplicate_count") or 1), 1),
    )
    try:
        render_context["__enforce_result_mapping__"] = True
        render_context["__strict_client_delivery__"] = True
        rendered_bytes, fill_status = render_template_file(source_bytes, source_name, render_context)
        source_ext = Path(source_name).suffix.lower()
        output_ext = Path(output_name).suffix.lower()
        if not source_ext or source_ext != output_ext:
            raise ValueError("模板格式与输出格式不一致")
        if fill_status in _BLOCKED_RENDER_STATUSES:
            raise ValueError(f"附件未完成可信编制：{fill_status}")
        if rendered_bytes == source_bytes:
            raise ValueError("生成结果与模板原文件完全相同")
        render_contract = render_context.get("__template_render_contract__") or {}
        if not bool(render_contract.get("audit_result_mapped")):
            raise ValueError("附件没有承载完整 AI 文稿或结构化事实计划")
        composition_manifest = render_context.get("__attachment_composition_manifest__") or {}
        if authored and composition_manifest.get("source_business_body_reused") is not True:
            raise ValueError("附件未证明已在原模板正文中原位填充")
        if authored and composition_manifest.get("source_body_preserved") is not True:
            raise ValueError("附件未证明保留了原模板正文结构")
        if composition_manifest.get("unplaced_section_ids"):
            sections = "、".join(
                str(item) for item in composition_manifest.get("unplaced_section_ids")[:6]
            )
            raise ValueError(f"模板缺少允许写入的正文位，拒绝重建正文：{sections}")
        validation = _validate_rendered_file(rendered_bytes, output_name)
        validation.update(
            _validate_rendered_template_content(
                rendered_bytes,
                template_type=template_type,
                source_file_name=source_name,
                render_context=render_context,
            )
        )
        fidelity = _validate_template_format_fidelity(
            source_bytes,
            rendered_bytes,
            source_file_name=source_name,
            render_context=render_context,
        )
        validation["template_format_fidelity"] = fidelity
        if fidelity.get("format_contract_checked") and not fidelity.get("format_contract_passed"):
            raise ValueError("模板设计系统保真校验未通过")
        quality_validation = evaluate_attachment_quality(
            template_type=template_type,
            rendered_bytes=rendered_bytes,
            output_name=output_name,
            snapshot=snapshot,
            material_index=dict(preflight_plan.get("material_index") or {}),
            render_context=render_context,
            authored=authored,
        )
        if not bool(quality_validation.get("passed")):
            blockers = "；".join(
                str(item) for item in quality_validation.get("blockers") or []
            )
            raise ValueError(f"附件质量门禁未通过：{blockers or '未通过未分类检查'}")
    except Exception as exc:
        if isinstance(exc, AttachmentItemGenerationError):
            raise
        raise AttachmentItemGenerationError(
            str(exc)[:1000],
            stage="compose_and_validate",
            retryable=False,
        ) from exc

    content_type = _content_type_for_file_name(
        output_name,
        str(generation_item.get("source_content_type") or planned.get("source_content_type") or ""),
    )
    run_key = _safe_filename_part(generation_run_id, "run")[:64]
    try:
        uploaded = get_minio_service().upload_artifact(
            project_id=engagement_id,
            file_name=output_name,
            content_type=content_type,
            file_bytes=rendered_bytes,
            storage_file_name=(
                f"package-v{reserved_package_version}-item-{ordinal}-{run_key}-{output_name}"
            ),
        )
    except Exception as exc:
        raise AttachmentItemGenerationError(
            f"附件上传失败：{str(exc)[:300]}",
            stage="artifact_upload",
            retryable=True,
        ) from exc

    reference = ArtifactRef(
        artifact_type=template_type if not code else f"{template_type}_{code}",
        template_version=template_version_ref(template),
        version=int(reserved_package_version),
        storage_ref=uploaded.storage_ref,
        content_type=content_type,
        file_name=output_name,
        status="draft",
    ).to_dict()
    from .attachment_document_author_service import public_authoring_manifest

    declared_slot_contract = render_context.get("__attachment_slot_contract__") or {}
    docx_slot_contract_verified = bool(
        extension != ".docx"
        or (
            isinstance(declared_slot_contract, dict)
            and declared_slot_contract.get("declared")
            and declared_slot_contract.get("slots")
        )
    )
    delivery_contract = {
        "contract_version": "annual-attachment-delivery-contract-v1",
        "frozen_plan_sha256": str(preflight_plan.get("frozen_input_sha256") or ""),
        "template_file_id": planned.get("template_file_id"),
        "source_template_sha256": source_sha256,
        "output_sha256": validation["sha256"],
        "docx_slot_contract_verified": docx_slot_contract_verified,
        "template_fidelity_verified": bool(
            not fidelity.get("format_contract_checked")
            or fidelity.get("format_contract_passed")
        ),
    }
    reference.update(
        {
            "template_file_id": planned.get("template_file_id"),
            "source_template_file_name": source_name,
            "source_template_file_ext": extension,
            "output_file_ext": Path(output_name).suffix.lower(),
            "template_fill_status": fill_status,
            "rendered_from_template": True,
            "template_used_as_design_blueprint": False,
            "source_template_sha256": source_sha256,
            "output_sha256": validation["sha256"],
            "output_size": validation["size"],
            "format_validation": validation,
            "quality_validation": quality_validation,
            "delivery_readiness": (
                "ready" if quality_validation.get("passed") else "needs_review"
            ),
            "delivery_contract": delivery_contract,
            "result_placement": _result_placement_description(template_type, source_name),
            "audit_result_included": True,
            "authoring": public_authoring_manifest(authored) if authored else {},
            "composition": dict(render_context.get("__attachment_composition_manifest__") or {}),
            "financial_semantic_plan": dict(render_context.get("__financial_semantic_manifest__") or {}),
            "generation_ordinal": ordinal,
            "package_version": int(reserved_package_version),
            "template_mapping": {
                "plan_version": preflight_plan.get("plan_version"),
                "render_contract": _template_render_contract_evidence(render_context),
                "material_catalog_sha256": (preflight_plan.get("material_catalog") or {}).get("catalog_sha256"),
                "material_source_sha256": (preflight_plan.get("material_index") or {}).get("source_sha256"),
                "frozen_report_id": frozen_report.get("report_id"),
                "frozen_report_version": frozen_report.get("report_version"),
                "frozen_report_snapshot_sha256": frozen_report.get("snapshot_sha256"),
            },
        }
    )
    if reserved_package_id:
        reference["package_id"] = int(reserved_package_id)
    return reference


def generate_annual_attachment_package(
    engagement_id: int,
    *,
    created_by: str = "ai_agent",
    requested_types: list[str] | None = None,
    settings: Settings | None = None,
    preflight_plan: dict[str, Any] | None = None,
) -> dict[str, Any]:
    """Compose a package from one frozen preflight and its template blueprints.

    根据前置检查冻结的审计结果、项目证据、内容范围和版式原型编制年度
    审计交付附件。叙事附件重新撰写完整正文；结构化附件从冻结事实重建
    业务值。模板正文和示例金额只提供设计参考，不作为项目证据或交付正文。

    核心交付附件包括：年度审计报告、年度审计财务报表、财务报表附注。
    管理建议书是按模板单独生成的可选附件；审计工作底稿和函证属于
    过程资料，需显式请求才会生成。
    """

    resolved = settings or get_settings()
    selected_types = tuple(
        requested_types
        or list((preflight_plan or {}).get("requested_types") or [])
        or DEFAULT_ATTACHMENT_TYPES
    )
    if preflight_plan is None:
        preflight_plan = plan_annual_attachment_package(
            engagement_id,
            requested_types=list(selected_types),
            settings=resolved,
        )
    if int(preflight_plan.get("engagement_id") or 0) != engagement_id:
        raise ValueError("附件生成前置检查属于其他审计项目")
    planned_types = tuple(
        str(item) for item in preflight_plan.get("requested_types") or [] if str(item)
    )
    if planned_types and planned_types != selected_types:
        raise ValueError("附件生成类型与冻结前置检查不一致，请重新执行前置检查")
    planned_template_types = {
        str(item.get("template_type") or "")
        for item in preflight_plan.get("templates") or []
        if isinstance(item, dict) and str(item.get("template_type") or "")
    }
    known_attachment_types = (
        set(DEFAULT_ATTACHMENT_TYPES)
        | set(OPTIONAL_ATTACHMENT_TYPES)
        | planned_template_types
    )
    unknown = [item for item in selected_types if item not in known_attachment_types]
    if unknown:
        raise ValueError(f"不支持的附件类型：{', '.join(unknown)}")
    if str(preflight_plan.get("status") or "blocked") != "ready":
        blockers = list(preflight_plan.get("blockers") or [])
        detail = "；".join(str(item) for item in blockers[:8]) or "模板/资料映射未通过生成门禁"
        raise ValueError(
            f"附件生成前置检查未通过：{detail}。"
            "系统没有生成或返回任何模板原文件，请先修正模板或补齐资料后重新执行。"
        )

    frozen_report = preflight_plan.get("frozen_report") or {}
    frozen_snapshot = frozen_report.get("fact_snapshot") or {}
    frozen_report_text = str(frozen_report.get("report_text") or "")
    if not isinstance(frozen_snapshot, dict) or not frozen_snapshot or not frozen_report_text:
        raise ValueError("附件生成前置检查缺少冻结审计报告输入")
    expected_snapshot_hash = str(frozen_report.get("snapshot_sha256") or "")
    if expected_snapshot_hash and expected_snapshot_hash != _stable_snapshot_hash(
        dict(frozen_snapshot)
    ):
        raise ValueError("附件生成前置检查的冻结审计快照哈希不一致")

    planned_templates = [
        dict(item)
        for item in preflight_plan.get("templates") or []
        if isinstance(item, dict)
        and str(item.get("template_type") or "") in selected_types
    ]
    templates_by_type = {
        attachment_type: [
            item
            for item in planned_templates
            if str(item.get("template_type") or "") == attachment_type
        ]
        for attachment_type in selected_types
    }
    missing = [item for item in selected_types if not templates_by_type[item]]
    if missing:
        usage_labels = {
            "annual_report": "年度审计报告",
            "financial_statements": "财务报表",
            "notes": "财务报表附注",
            "management_letter": "管理建议书（可选）",
            "audit_workpaper": "审计工作底稿",
            "confirmations": "函证",
        }
        labels = "、".join(usage_labels.get(item, item) for item in missing)
        raise ValueError(
            f"以下附件类型没有已激活且包含实际文件的模板版本：{labels}。"
            "请重新执行附件生成前置检查，并确认所选模板文件仍处于冻结计划中。"
        )

    package_version = _next_package_version(engagement_id, resolved)
    template_snapshot: dict[str, Any] = {"fill_plan": preflight_plan}
    for attachment_type in selected_types:
        planned_group = templates_by_type[attachment_type]
        frozen_template = dict(planned_group[0].get("template_snapshot") or {})
        template_snapshot[attachment_type] = {
            "template_code": frozen_template.get("template_code"),
            "template_type": attachment_type,
            "version_no": int(frozen_template.get("version_no") or 0),
            "version_label": str(
                frozen_template.get("version_label")
                or planned_group[0].get("template_version")
                or ""
            ),
            "content_hash": str(frozen_template.get("content_hash") or ""),
            "content": frozen_template.get("content") or {},
            "field_schema": frozen_template.get("field_schema") or {},
            "template_contract": frozen_template.get("template_contract") or {},
            "files": [
                {
                    "id": item.get("template_file_id"),
                    "file_name": item.get("file_name"),
                    "file_ext": (
                        item.get("source_file_ext")
                        or Path(str(item.get("file_name") or "")).suffix.lower()
                    ),
                    "content_type": item.get("source_content_type"),
                    "storage_ref": item.get("source_storage_ref"),
                    "storage_sha256": item.get("source_template_sha256"),
                    "file_size": item.get("source_template_size"),
                }
                for item in planned_group
            ],
        }

    work_items: list[dict[str, Any]] = []
    ordinal = 0
    for attachment_type in selected_types:
        planned_group = templates_by_type[attachment_type]
        for duplicate_index, planned in enumerate(planned_group):
            work_items.append(
                {
                    "ordinal_no": ordinal,
                    "attachment_type": attachment_type,
                    "template_file_id": planned.get("template_file_id"),
                    "template_snapshot": dict(
                        planned.get("template_snapshot") or {}
                    ),
                    "source_file_name": str(planned.get("file_name") or ""),
                    "source_content_type": str(
                        planned.get("source_content_type")
                        or "application/octet-stream"
                    ),
                    "source_storage_ref": str(
                        planned.get("source_storage_ref") or ""
                    ),
                    "source_sha256": str(
                        planned.get("source_template_sha256") or ""
                    ),
                    "duplicate_index": duplicate_index,
                    "duplicate_count": len(planned_group),
                }
            )
            ordinal += 1

    def generate_one(item: dict[str, Any]) -> dict[str, Any]:
        item_ordinal = int(item["ordinal_no"])
        try:
            artifact = generate_annual_attachment_item(
                engagement_id,
                generation_item=item,
                preflight_plan=preflight_plan,
                reserved_package_id=None,
                reserved_package_version=package_version,
                created_by=created_by or "ai_agent",
                generation_run_id=f"direct-package-v{package_version}",
                settings=resolved,
            )
            if not isinstance(artifact, dict) or not artifact:
                raise AttachmentItemGenerationError(
                    "单附件生成入口没有返回有效产物",
                    stage="artifact_result_validation",
                    retryable=False,
                )
            normalized_artifact = dict(artifact)
            normalized_artifact["generation_ordinal"] = item_ordinal
            normalized_artifact.setdefault("package_version", package_version)
            return {
                "ordinal_no": item_ordinal,
                "artifact": normalized_artifact,
            }
        except Exception as exc:
            stage = str(
                getattr(exc, "stage", "") or "attachment_item_generation"
            )
            message = str(exc).strip()[:1000] or "附件生成失败"
            return {
                "ordinal_no": item_ordinal,
                "error": {
                    "ordinal_no": item_ordinal,
                    "attachment_type": str(item.get("attachment_type") or ""),
                    "source_file_name": str(item.get("source_file_name") or ""),
                    "output_file_name": str(item.get("output_file_name") or ""),
                    "stage": stage,
                    "retryable": bool(getattr(exc, "retryable", False)),
                    "message": message,
                },
            }

    max_workers = max(
        1,
        min(
            int(getattr(resolved, "attachment_generation_concurrency", 1) or 1),
            len(work_items),
        ),
    )
    results: list[dict[str, Any]] = []
    with ThreadPoolExecutor(
        max_workers=max_workers,
        thread_name_prefix="annual-attachment",
    ) as executor:
        futures = [executor.submit(generate_one, item) for item in work_items]
        for future in as_completed(futures):
            results.append(future.result())
    results.sort(key=lambda item: int(item.get("ordinal_no") or 0))

    published = [
        dict(item["artifact"]) for item in results if item.get("artifact")
    ]
    error_details = [
        dict(item["error"]) for item in results if item.get("error")
    ]
    errors = [
        (
            f"{item.get('source_file_name') or item.get('attachment_type')}: "
            f"{item.get('message')}"
        )
        for item in error_details
    ]
    template_snapshot["generation"] = {
        "mode": "frozen_single_item_concurrent_composition",
        "item_count": len(work_items),
        "max_workers": max_workers,
        "success_count": len(published),
        "error_count": len(error_details),
        "errors": error_details,
    }
    status = "draft_saved" if not errors else ("partial" if published else "failed")
    package_id = _persist_package(
        engagement_id=engagement_id,
        package_version=package_version,
        status=status,
        template_snapshot=template_snapshot,
        artifacts=published,
        created_by=created_by,
        settings=resolved,
    )
    for index, artifact in enumerate(published):
        artifact["download_url"] = (
            f"/api/annual-audit/{engagement_id}/attachment-packages/"
            f"{package_id}/files/{index}"
        )
        artifact["preview_url"] = (
            f"/api/annual-audit/{engagement_id}/attachment-packages/"
            f"{package_id}/files/{index}/preview"
        )
    return {
        "package_id": package_id,
        "package_version": package_version,
        "status": status,
        "template_snapshot": template_snapshot,
        "artifacts": published,
        "errors": errors,
        "error_details": error_details,
        "preflight_plan": preflight_plan,
    }


__all__ = [
    "DEFAULT_ATTACHMENT_TYPES",
    "OPTIONAL_ATTACHMENT_TYPES",
    "AttachmentItemGenerationError",
    "generate_annual_attachment_item",
    "generate_annual_attachment_package",
    "plan_annual_attachment_package",
    "render_template_file",
]
