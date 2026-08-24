"""Annual-audit HTTP API."""

from __future__ import annotations

import json
import hashlib
import mimetypes
from io import BytesIO
from pathlib import Path
from typing import Any
from urllib.parse import quote

from fastapi import APIRouter, Depends, File, Form, HTTPException, Query, UploadFile
from fastapi.responses import JSONResponse, StreamingResponse
from pydantic import BaseModel, Field

from ai_hunter.app.auth.identity import Identity
from ai_hunter.app.auth.permissions import require_module
from ai_hunter.app.auth.tenancy import require_case_access
from ai_hunter.app.settings import get_settings
from ai_hunter.app.services.minio_service import get_minio_service

from . import document_repository as documents
from . import engagement_repository as engagements
from . import task_repository as tasks
from .analysis_service import data_readiness, run_cash_and_bank, run_sales_receivables
from .execution_service import (
    WorkflowBlockedError,
    archive_audit_engagement,
    build_release_gate,
    bootstrap_execution,
    evaluate_release_gate,
    freeze_policy_binding,
    get_execution_snapshot,
    issue_audit_report,
    record_review_decision,
    resolve_finding,
    update_engagement_profile,
    update_program_item,
    upsert_confirmation,
)
from .knowledge_service import (
    create_knowledge_release,
    create_knowledge_version,
    create_ruleset,
    list_policy_catalog,
    publish_knowledge_release,
    publish_ruleset,
    review_knowledge_version,
    review_rule,
    search_published_knowledge,
    upsert_rule,
)
from .report_service import generate_annual_report_draft
from .storage import AnnualAuditStorageError, mysql_connection
from . import generic_template_repository as generic_templates
from . import template_file_repository as templates


router = APIRouter()


def _require_annual_module(module: str):
    """Resolve a module-authorized identity suitable for annual-audit routes."""

    module_dependency = require_module(module)

    def _dependency(identity: Identity = Depends(module_dependency)) -> Identity:
        if get_settings().auth_enabled:
            if not identity.user_id:
                raise HTTPException(status_code=403, detail="当前身份缺少 user_id")
            if not identity.is_super_admin and not identity.company_id:
                raise HTTPException(status_code=403, detail="当前身份缺少 company_id")
        return identity

    return _dependency


_require_report = _require_annual_module("report")
_require_materials = _require_annual_module("materials")
_require_tasks = _require_annual_module("tasks")
_require_drilldown = _require_annual_module("drilldown")
_require_admin = _require_annual_module("admin")


class CreateEngagementRequest(BaseModel):
    case_name: str = Field(min_length=1, max_length=255)
    case_type: str = "年度财务报表审计"
    entity_name: str = Field(min_length=1, max_length=255)
    entity_uscc: str | None = None
    fiscal_year: int | None = Field(default=None, ge=2000, le=2200)
    company_id: str | None = None


class FullContextRequest(BaseModel):
    case_id: int = Field(gt=0)
    recompute: bool = False
    mining_transfer_prices: dict[str, Any] | None = None


class AnnualAnalysisRequest(BaseModel):
    case_id: int = Field(gt=0)
    recompute: bool = False


class CashAnalysisRequest(AnnualAnalysisRequest):
    large_amount_threshold: float = Field(default=1_000_000, gt=0)


class AnnualReportRequest(AnnualAnalysisRequest):
    corrections: list[str] = Field(default_factory=list)


class ValidateDocCategoryRequest(BaseModel):
    case_id: int = Field(gt=0)
    doc_category: str = Field(min_length=1, max_length=128)
    file_names: list[str] = Field(default_factory=list)
    text_preview: str | None = None


class TaskBatchRequest(BaseModel):
    case_id: int = Field(gt=0)
    tasks: list[dict[str, Any]] = Field(default_factory=list)


class ManageTaskRequest(BaseModel):
    case_id: int = Field(gt=0)
    action: str = Field(min_length=1, max_length=32)
    filter_status: str | None = None
    task_id: int | None = Field(default=None, gt=0)
    new_status: str | None = None
    completion_note: str | None = None
    assigned_to: str | None = None


class EngagementProfileUpdateRequest(BaseModel):
    profile: dict[str, Any] = Field(default_factory=dict)
    acceptance_status: str | None = None
    independence_status: str | None = None
    entity_type: str | None = None
    audit_purpose: str | None = None
    accounting_framework: str | None = None
    firm_name: str | None = None
    engagement_partner: str | None = None
    signing_cpa_primary: str | None = None
    signing_cpa_secondary: str | None = None
    signing_cpa_primary_user_id: str | None = None
    signing_cpa_secondary_user_id: str | None = None
    data_classification: str | None = None
    data_residency: str | None = None
    model_data_policy: str | None = None


class ProgramItemUpdateRequest(BaseModel):
    status: str | None = None
    sample_plan: dict[str, Any] | list[Any] | None = None
    evidence_refs: list[dict[str, Any]] | None = None
    exception_count: int | None = Field(default=None, ge=0)
    alternative_procedures: dict[str, Any] | list[Any] | None = None
    conclusion_text: str | None = None
    not_applicable_reason: str | None = None


class ReviewDecisionRequest(BaseModel):
    review_level: str = Field(min_length=1, max_length=64)
    decision: str = Field(min_length=1, max_length=32)
    decision_note: str | None = None
    scope: dict[str, Any] | list[Any] = Field(default_factory=dict)


class FindingResolutionRequest(BaseModel):
    resolution_status: str = "closed"
    resolution_type: str | None = None
    resolution_note: str | None = None
    evidence_refs: list[dict[str, Any]] = Field(default_factory=list)


class ConfirmationRequest(BaseModel):
    confirmation_id: int | None = Field(default=None, gt=0)
    procedure_code: str = Field(min_length=1, max_length=64)
    counterparty_name: str = Field(min_length=1, max_length=255)
    confirmation_type: str = Field(min_length=1, max_length=64)
    status: str = "planned"
    auditor_controlled_delivery: bool = False
    request_evidence_refs: list[dict[str, Any]] = Field(default_factory=list)
    response_evidence_refs: list[dict[str, Any]] = Field(default_factory=list)
    reliability_assessment: str | None = None
    exception_description: str | None = None
    alternative_procedures: dict[str, Any] | list[Any] = Field(default_factory=dict)
    conclusion_text: str | None = None


class PolicyBindingRequest(BaseModel):
    knowledge_release_id: int = Field(gt=0)
    ruleset_id: int = Field(gt=0)
    reporting_period_date: str | None = None


class IssueAuditReportRequest(BaseModel):
    report_artifact_ref: str = Field(min_length=1, max_length=2048)
    report_artifact_sha256: str = Field(min_length=64, max_length=64)
    opinion_type: str = Field(min_length=1, max_length=64)
    signing_attestation: bool = False
    issuance_note: str | None = None


class ArchiveAuditRequest(BaseModel):
    archive_manifest_ref: str = Field(min_length=1, max_length=2048)
    archive_manifest_sha256: str = Field(min_length=64, max_length=64)
    retention_until: str | None = None


class KnowledgeVersionRequest(BaseModel):
    document_code: str = Field(min_length=1, max_length=128)
    title: str = Field(min_length=1, max_length=512)
    authority_type: str = Field(min_length=1, max_length=64)
    source_issuer: str = Field(min_length=1, max_length=255)
    source_url: str = Field(min_length=1, max_length=2048)
    source_hash: str = Field(min_length=64, max_length=64)
    content_text: str = Field(min_length=1)
    publication_date: str | None = None
    effective_from: str | None = None
    effective_to: str | None = None
    change_summary: str | None = None
    scope: dict[str, Any] | list[Any] = Field(default_factory=dict)
    chunks: list[dict[str, Any]] | None = None


class KnowledgeReviewRequest(BaseModel):
    action: str = Field(min_length=1, max_length=16)
    note: str | None = None


class KnowledgeReleaseRequest(BaseModel):
    release_code: str = Field(min_length=1, max_length=128)
    release_version: str = Field(min_length=1, max_length=64)
    knowledge_version_ids: list[int] = Field(min_length=1)
    effective_from: str | None = None
    effective_to: str | None = None
    approval_note: str | None = None


class RulesetRequest(BaseModel):
    ruleset_code: str = Field(min_length=1, max_length=128)
    version: str = Field(min_length=1, max_length=64)
    name: str = Field(min_length=1, max_length=255)
    scope: dict[str, Any] | list[Any] = Field(default_factory=dict)
    effective_from: str | None = None
    effective_to: str | None = None


class RuleRequest(BaseModel):
    rule_code: str = Field(min_length=1, max_length=128)
    rule_type: str = Field(min_length=1, max_length=32)
    name: str = Field(min_length=1, max_length=512)
    authority_locator: str | None = None
    knowledge_version_id: int | None = Field(default=None, gt=0)
    applicability: dict[str, Any] | list[Any] = Field(default_factory=dict)
    preconditions: dict[str, Any] | list[Any] = Field(default_factory=dict)
    evidence_requirements: dict[str, Any] | list[Any] = Field(default_factory=dict)
    logic: dict[str, Any] | list[Any] = Field(default_factory=dict)
    exception_handling: dict[str, Any] | list[Any] = Field(default_factory=dict)


class PublicationRequest(BaseModel):
    effective_from: str | None = None
    effective_to: str | None = None
    approval_note: str | None = None


class TemplateCreateRequest(BaseModel):
    template_code: str = Field(min_length=1, max_length=128)
    template_type: str = Field(min_length=1, max_length=64)
    name: str = Field(min_length=1, max_length=255)
    description: str = Field(default="", max_length=1024)
    content: dict[str, Any] = Field(default_factory=dict)
    field_schema: dict[str, Any] | list[Any] = Field(default_factory=dict)
    version_label: str = Field(default="v1", min_length=1, max_length=128)


class TemplateVersionRequest(BaseModel):
    content: dict[str, Any] = Field(default_factory=dict)
    field_schema: dict[str, Any] | list[Any] = Field(default_factory=dict)
    version_label: str = Field(default="", max_length=128)


def _storage_http_error(exc: Exception) -> HTTPException:
    return HTTPException(status_code=503, detail=f"年审独立存储不可用：{exc}")


def _template_storage_http_error(exc: Exception) -> HTTPException:
    return HTTPException(status_code=503, detail=f"模板存储不可用：{exc}")


def _workflow_http_error(exc: WorkflowBlockedError) -> HTTPException:
    return HTTPException(
        status_code=409,
        detail={"message": str(exc), "blockers": exc.blockers},
    )


def _has_any_role(identity: Identity, roles: set[str]) -> bool:
    return identity.is_admin or bool(set(identity.roles) & roles)


def _require_project_control(identity: Identity) -> None:
    if not _has_any_role(identity, {"engagement_manager", "engagement_partner", "reviewer"}):
        raise HTTPException(status_code=403, detail="仅项目管理或复核角色可变更项目控制信息")


def _require_review_level(identity: Identity, level: str) -> None:
    allowed_roles = {
        "project_manager": {"engagement_manager", "reviewer", "engagement_partner"},
        "department_manager": {"reviewer", "engagement_partner"},
        "engagement_partner": {"engagement_partner"},
    }
    if not _has_any_role(identity, allowed_roles.get(level, set())):
        raise HTTPException(status_code=403, detail="当前角色无权作出该级复核决定")


def _require_human_signatory(identity: Identity) -> None:
    if not _has_any_role(identity, {"engagement_partner"}):
        raise HTTPException(status_code=403, detail="正式签发仅限项目合伙人身份")


def _creation_payload(request: CreateEngagementRequest, identity: Identity) -> dict[str, Any]:
    """Bind tenant and audit actor fields to the authenticated identity."""

    payload = request.model_dump()
    requested_company_id = str(payload.get("company_id") or "").strip()
    company_id = str(identity.company_id or "").strip()

    if identity.is_super_admin:
        company_id = requested_company_id or company_id
    elif requested_company_id and requested_company_id != company_id:
        raise HTTPException(status_code=403, detail="不能跨机构创建年度审计项目")

    if get_settings().auth_enabled and not company_id:
        raise HTTPException(status_code=403, detail="创建年度审计项目必须绑定 company_id")

    payload.update(
        {
            "company_id": company_id,
            "owner_id": identity.user_id,
            "created_by": identity.user_id,
        }
    )
    return payload


@router.get("/health", tags=["年审基础设施"], summary="年审独立存储健康检查")
def health() -> dict[str, str]:
    try:
        with mysql_connection() as connection:
            with connection.cursor() as cursor:
                cursor.execute("SELECT 1 AS ok")
                cursor.fetchone()
    except AnnualAuditStorageError as exc:
        raise _storage_http_error(exc) from exc
    settings = get_settings()
    return {
        "status": "ok",
        "db": "connected",
        "domain": "annual_audit",
        "llm_provider": settings.resolve_provider("agent"),
    }


@router.get("/api/cases", tags=["年审项目"], summary="按原案件契约查询年审项目")
def list_cases(
    keyword: str | None = None,
    case_type: str | None = None,
    status: str | None = None,
    page: int = Query(1, ge=1),
    page_size: int = Query(20, ge=1, le=100),
    identity: Identity = Depends(_require_report),
) -> dict[str, Any]:
    try:
        return engagements.list_engagements(
            keyword=keyword,
            case_type=case_type,
            status=status,
            company_id=identity.company_id,
            user_id=identity.user_id,
            is_company_admin=identity.is_company_admin,
            is_super_admin=identity.is_super_admin,
            page=page,
            page_size=page_size,
        )
    except AnnualAuditStorageError as exc:
        raise _storage_http_error(exc) from exc


@router.post("/api/ingest/case", tags=["年审项目"], summary="按原建案契约创建年审项目")
def create_case(
    request: CreateEngagementRequest,
    identity: Identity = Depends(_require_report),
) -> dict[str, Any]:
    try:
        return engagements.create_engagement(_creation_payload(request, identity))
    except ValueError as exc:
        raise HTTPException(status_code=422, detail=str(exc)) from exc
    except AnnualAuditStorageError as exc:
        raise _storage_http_error(exc) from exc


@router.get(
    "/api/case/{case_id}/profile",
    tags=["年审项目"],
    summary="按原案件画像契约查询年审项目画像",
)
def get_case_profile(
    case_id: int,
    identity: Identity = Depends(_require_report),
) -> dict[str, Any]:
    require_case_access(case_id, identity)
    try:
        return engagements.get_engagement_profile(case_id)
    except engagements.EngagementNotFoundError as exc:
        raise HTTPException(status_code=404, detail=str(exc)) from exc
    except AnnualAuditStorageError as exc:
        raise _storage_http_error(exc) from exc


@router.get("/api/ingest/doc-categories", tags=["年审资料"])
def list_doc_categories(
    identity: Identity = Depends(_require_materials),
) -> dict[str, Any]:
    try:
        return documents.list_doc_categories()
    except AnnualAuditStorageError as exc:
        raise _storage_http_error(exc) from exc


@router.get("/api/case/{case_id}/doc-categories", tags=["年审资料"])
def get_case_doc_categories(
    case_id: int,
    identity: Identity = Depends(_require_materials),
) -> dict[str, Any]:
    require_case_access(case_id, identity)
    try:
        return documents.get_case_doc_categories(case_id)
    except engagements.EngagementNotFoundError as exc:
        raise HTTPException(status_code=404, detail=str(exc)) from exc
    except AnnualAuditStorageError as exc:
        raise _storage_http_error(exc) from exc


@router.post("/api/ingest/validate-doc-category", tags=["年审资料"])
def validate_doc_category(
    request: ValidateDocCategoryRequest,
    identity: Identity = Depends(_require_materials),
) -> dict[str, Any]:
    require_case_access(request.case_id, identity)
    try:
        return documents.validate_doc_category(request.model_dump())
    except engagements.EngagementNotFoundError as exc:
        raise HTTPException(status_code=404, detail=str(exc)) from exc
    except AnnualAuditStorageError as exc:
        raise _storage_http_error(exc) from exc


@router.post("/api/tasks/create-batch", tags=["年审任务"])
def create_task_batch(
    request: TaskBatchRequest,
    identity: Identity = Depends(_require_tasks),
) -> dict[str, Any]:
    require_case_access(request.case_id, identity)
    try:
        return tasks.create_task_batch(request.case_id, request.tasks)
    except engagements.EngagementNotFoundError as exc:
        raise HTTPException(status_code=404, detail=str(exc)) from exc
    except ValueError as exc:
        raise HTTPException(status_code=422, detail=str(exc)) from exc
    except AnnualAuditStorageError as exc:
        raise _storage_http_error(exc) from exc


@router.post("/api/tasks/manage", tags=["年审任务"])
def manage_tasks(
    request: ManageTaskRequest,
    identity: Identity = Depends(_require_tasks),
) -> dict[str, Any]:
    require_case_access(request.case_id, identity)
    try:
        return tasks.manage_tasks(request.model_dump())
    except engagements.EngagementNotFoundError as exc:
        raise HTTPException(status_code=404, detail=str(exc)) from exc
    except ValueError as exc:
        raise HTTPException(status_code=422, detail=str(exc)) from exc
    except AnnualAuditStorageError as exc:
        raise _storage_http_error(exc) from exc


@router.post(
    "/api/audit/get_full_context",
    tags=["年审项目"],
    summary="按原完整审计契约聚合年审上下文",
)
def get_full_context(
    request: FullContextRequest,
    identity: Identity = Depends(_require_report),
) -> dict[str, Any]:
    require_case_access(request.case_id, identity)
    try:
        return engagements.get_full_context(request.case_id)
    except engagements.EngagementNotFoundError as exc:
        raise HTTPException(status_code=404, detail=str(exc)) from exc
    except AnnualAuditStorageError as exc:
        raise _storage_http_error(exc) from exc


@router.post("/api/annual-audit/readiness", tags=["年审分析"])
def get_data_readiness(
    request: AnnualAnalysisRequest,
    identity: Identity = Depends(_require_drilldown),
) -> dict[str, Any]:
    require_case_access(request.case_id, identity)
    try:
        return data_readiness(request.case_id)
    except engagements.EngagementNotFoundError as exc:
        raise HTTPException(status_code=404, detail=str(exc)) from exc
    except AnnualAuditStorageError as exc:
        raise _storage_http_error(exc) from exc


@router.post("/api/annual-audit/sales-receivables", tags=["年审分析"])
def analyze_sales_receivables(
    request: AnnualAnalysisRequest,
    identity: Identity = Depends(_require_drilldown),
) -> dict[str, Any]:
    require_case_access(request.case_id, identity)
    try:
        return run_sales_receivables(request.case_id, recompute=request.recompute)
    except engagements.EngagementNotFoundError as exc:
        raise HTTPException(status_code=404, detail=str(exc)) from exc
    except AnnualAuditStorageError as exc:
        raise _storage_http_error(exc) from exc


@router.post("/api/annual-audit/cash-and-bank", tags=["年审分析"])
def analyze_cash_and_bank(
    request: CashAnalysisRequest,
    identity: Identity = Depends(_require_drilldown),
) -> dict[str, Any]:
    require_case_access(request.case_id, identity)
    from decimal import Decimal

    try:
        return run_cash_and_bank(
            request.case_id,
            large_amount_threshold=Decimal(str(request.large_amount_threshold)),
            recompute=request.recompute,
        )
    except engagements.EngagementNotFoundError as exc:
        raise HTTPException(status_code=404, detail=str(exc)) from exc
    except AnnualAuditStorageError as exc:
        raise _storage_http_error(exc) from exc


@router.post("/api/annual-audit/report", tags=["骞村鎶ュ憡"])
def generate_report(
    request: AnnualReportRequest,
    identity: Identity = Depends(_require_report),
) -> dict[str, Any]:
    """Run all deterministic annual cycles and publish versioned artifacts."""

    require_case_access(request.case_id, identity)
    try:
        return generate_annual_report_draft(
            request.case_id,
            recompute=request.recompute,
            corrections=request.corrections,
            created_by=identity.user_id,
        )
    except engagements.EngagementNotFoundError as exc:
        raise HTTPException(status_code=404, detail=str(exc)) from exc
    except AnnualAuditStorageError as exc:
        raise _storage_http_error(exc) from exc


@router.get("/api/annual-audit/{case_id}/execution", tags=["年审执行"])
def get_execution(
    case_id: int,
    identity: Identity = Depends(_require_report),
) -> dict[str, Any]:
    """Read the governed programme; missing bootstrap rows are created idempotently."""

    require_case_access(case_id, identity)
    try:
        return bootstrap_execution(case_id, actor_user_id=identity.user_id)
    except engagements.EngagementNotFoundError as exc:
        raise HTTPException(status_code=404, detail=str(exc)) from exc
    except AnnualAuditStorageError as exc:
        raise _storage_http_error(exc) from exc


@router.put("/api/annual-audit/{case_id}/profile", tags=["年审执行"])
def update_profile(
    case_id: int,
    request: EngagementProfileUpdateRequest,
    identity: Identity = Depends(_require_report),
) -> dict[str, Any]:
    require_case_access(case_id, identity)
    _require_project_control(identity)
    try:
        return update_engagement_profile(
            case_id,
            request.model_dump(exclude_none=True),
            actor_user_id=identity.user_id,
        )
    except WorkflowBlockedError as exc:
        raise _workflow_http_error(exc) from exc
    except ValueError as exc:
        raise HTTPException(status_code=422, detail=str(exc)) from exc
    except AnnualAuditStorageError as exc:
        raise _storage_http_error(exc) from exc


@router.put("/api/annual-audit/{case_id}/program/{procedure_code}", tags=["年审执行"])
def update_execution_program_item(
    case_id: int,
    procedure_code: str,
    request: ProgramItemUpdateRequest,
    identity: Identity = Depends(_require_tasks),
) -> dict[str, Any]:
    require_case_access(case_id, identity)
    try:
        return update_program_item(
            case_id,
            procedure_code,
            request.model_dump(exclude_none=True),
            actor_user_id=identity.user_id,
        )
    except WorkflowBlockedError as exc:
        raise _workflow_http_error(exc) from exc
    except ValueError as exc:
        raise HTTPException(status_code=422, detail=str(exc)) from exc
    except AnnualAuditStorageError as exc:
        raise _storage_http_error(exc) from exc


@router.post("/api/annual-audit/{case_id}/reviews", tags=["年审复核"])
def create_review_decision(
    case_id: int,
    request: ReviewDecisionRequest,
    identity: Identity = Depends(_require_report),
) -> dict[str, Any]:
    require_case_access(case_id, identity)
    _require_review_level(identity, request.review_level)
    try:
        return record_review_decision(
            case_id,
            request.model_dump(),
            reviewer_user_id=identity.user_id,
        )
    except WorkflowBlockedError as exc:
        raise _workflow_http_error(exc) from exc
    except ValueError as exc:
        raise HTTPException(status_code=422, detail=str(exc)) from exc
    except AnnualAuditStorageError as exc:
        raise _storage_http_error(exc) from exc


@router.put("/api/annual-audit/{case_id}/findings/{finding_id}/resolution", tags=["年审执行"])
def close_finding(
    case_id: int,
    finding_id: int,
    request: FindingResolutionRequest,
    identity: Identity = Depends(_require_drilldown),
) -> dict[str, Any]:
    require_case_access(case_id, identity)
    _require_project_control(identity)
    try:
        return resolve_finding(
            case_id,
            finding_id,
            request.model_dump(),
            actor_user_id=identity.user_id,
        )
    except WorkflowBlockedError as exc:
        raise _workflow_http_error(exc) from exc
    except ValueError as exc:
        raise HTTPException(status_code=422, detail=str(exc)) from exc
    except AnnualAuditStorageError as exc:
        raise _storage_http_error(exc) from exc


@router.put("/api/annual-audit/{case_id}/confirmations", tags=["年审执行"])
def save_confirmation(
    case_id: int,
    request: ConfirmationRequest,
    identity: Identity = Depends(_require_tasks),
) -> dict[str, Any]:
    require_case_access(case_id, identity)
    try:
        return upsert_confirmation(
            case_id,
            request.model_dump(),
            actor_user_id=identity.user_id,
        )
    except WorkflowBlockedError as exc:
        raise _workflow_http_error(exc) from exc
    except ValueError as exc:
        raise HTTPException(status_code=422, detail=str(exc)) from exc
    except AnnualAuditStorageError as exc:
        raise _storage_http_error(exc) from exc


@router.post("/api/annual-audit/{case_id}/policy-binding", tags=["年审执行"])
def freeze_project_policy_binding(
    case_id: int,
    request: PolicyBindingRequest,
    identity: Identity = Depends(_require_report),
) -> dict[str, Any]:
    require_case_access(case_id, identity)
    _require_project_control(identity)
    try:
        return freeze_policy_binding(
            case_id,
            request.model_dump(exclude_none=True),
            actor_user_id=identity.user_id,
        )
    except WorkflowBlockedError as exc:
        raise _workflow_http_error(exc) from exc
    except ValueError as exc:
        raise HTTPException(status_code=422, detail=str(exc)) from exc
    except AnnualAuditStorageError as exc:
        raise _storage_http_error(exc) from exc


@router.get("/api/annual-audit/{case_id}/release-gate", tags=["年审签发"])
def get_release_gate(
    case_id: int,
    identity: Identity = Depends(_require_report),
) -> dict[str, Any]:
    require_case_access(case_id, identity)
    try:
        return build_release_gate(case_id)
    except engagements.EngagementNotFoundError as exc:
        raise HTTPException(status_code=404, detail=str(exc)) from exc
    except AnnualAuditStorageError as exc:
        raise _storage_http_error(exc) from exc


@router.post("/api/annual-audit/{case_id}/release-gate", tags=["年审签发"])
def record_release_gate(
    case_id: int,
    identity: Identity = Depends(_require_report),
) -> dict[str, Any]:
    require_case_access(case_id, identity)
    _require_project_control(identity)
    try:
        return evaluate_release_gate(case_id, actor_user_id=identity.user_id)
    except engagements.EngagementNotFoundError as exc:
        raise HTTPException(status_code=404, detail=str(exc)) from exc
    except AnnualAuditStorageError as exc:
        raise _storage_http_error(exc) from exc


@router.post("/api/annual-audit/{case_id}/issue", tags=["年审签发"])
def issue_formal_audit_report(
    case_id: int,
    request: IssueAuditReportRequest,
    identity: Identity = Depends(_require_report),
) -> dict[str, Any]:
    require_case_access(case_id, identity)
    _require_human_signatory(identity)
    try:
        return issue_audit_report(
            case_id,
            request.model_dump(),
            signed_by=identity.user_id,
        )
    except WorkflowBlockedError as exc:
        raise _workflow_http_error(exc) from exc
    except ValueError as exc:
        raise HTTPException(status_code=422, detail=str(exc)) from exc
    except AnnualAuditStorageError as exc:
        raise _storage_http_error(exc) from exc


@router.post("/api/annual-audit/{case_id}/archive", tags=["年审归档"])
def archive_formal_audit_engagement(
    case_id: int,
    request: ArchiveAuditRequest,
    identity: Identity = Depends(_require_report),
) -> dict[str, Any]:
    require_case_access(case_id, identity)
    _require_project_control(identity)
    try:
        return archive_audit_engagement(
            case_id,
            request.model_dump(exclude_none=True),
            actor_user_id=identity.user_id,
        )
    except WorkflowBlockedError as exc:
        raise _workflow_http_error(exc) from exc
    except ValueError as exc:
        raise HTTPException(status_code=422, detail=str(exc)) from exc
    except AnnualAuditStorageError as exc:
        raise _storage_http_error(exc) from exc


@router.get("/api/annual-audit/knowledge/search", tags=["年审知识"])
def search_knowledge(
    query: str = Query(min_length=1, max_length=512),
    as_of_date: str | None = None,
    authority_type: list[str] | None = Query(default=None),
    limit: int = Query(12, ge=1, le=50),
    identity: Identity = Depends(_require_report),
) -> dict[str, Any]:
    try:
        return search_published_knowledge(
            query,
            as_of_date=as_of_date,
            authority_types=authority_type,
            limit=limit,
        )
    except ValueError as exc:
        raise HTTPException(status_code=422, detail=str(exc)) from exc
    except AnnualAuditStorageError as exc:
        raise _storage_http_error(exc) from exc


@router.get("/api/annual-audit/knowledge/policy-catalog", tags=["年审知识"])
def get_policy_catalog(
    identity: Identity = Depends(_require_report),
) -> dict[str, Any]:
    try:
        return list_policy_catalog()
    except AnnualAuditStorageError as exc:
        raise _storage_http_error(exc) from exc


@router.post("/api/annual-audit/knowledge/versions", tags=["年审知识治理"])
def create_knowledge(
    request: KnowledgeVersionRequest,
    identity: Identity = Depends(_require_admin),
) -> dict[str, Any]:
    try:
        return create_knowledge_version(request.model_dump(exclude_none=True), actor_user_id=identity.user_id)
    except ValueError as exc:
        raise HTTPException(status_code=422, detail=str(exc)) from exc
    except AnnualAuditStorageError as exc:
        raise _storage_http_error(exc) from exc


@router.post("/api/annual-audit/knowledge/versions/{knowledge_version_id}/review", tags=["年审知识治理"])
def review_knowledge(
    knowledge_version_id: int,
    request: KnowledgeReviewRequest,
    identity: Identity = Depends(_require_admin),
) -> dict[str, Any]:
    try:
        return review_knowledge_version(
            knowledge_version_id,
            request.model_dump(exclude_none=True),
            actor_user_id=identity.user_id,
        )
    except ValueError as exc:
        raise HTTPException(status_code=422, detail=str(exc)) from exc
    except AnnualAuditStorageError as exc:
        raise _storage_http_error(exc) from exc


@router.post("/api/annual-audit/knowledge/releases", tags=["年审知识治理"])
def create_knowledge_release_route(
    request: KnowledgeReleaseRequest,
    identity: Identity = Depends(_require_admin),
) -> dict[str, Any]:
    try:
        return create_knowledge_release(
            request.model_dump(exclude_none=True),
            actor_user_id=identity.user_id,
        )
    except ValueError as exc:
        raise HTTPException(status_code=422, detail=str(exc)) from exc
    except AnnualAuditStorageError as exc:
        raise _storage_http_error(exc) from exc


@router.post("/api/annual-audit/knowledge/releases/{release_id}/publish", tags=["年审知识治理"])
def publish_knowledge_release_route(
    release_id: int,
    request: PublicationRequest,
    identity: Identity = Depends(_require_admin),
) -> dict[str, Any]:
    try:
        return publish_knowledge_release(
            release_id,
            request.model_dump(exclude_none=True),
            actor_user_id=identity.user_id,
        )
    except ValueError as exc:
        raise HTTPException(status_code=422, detail=str(exc)) from exc
    except AnnualAuditStorageError as exc:
        raise _storage_http_error(exc) from exc


@router.post("/api/annual-audit/rulesets", tags=["年审知识治理"])
def create_ruleset_route(
    request: RulesetRequest,
    identity: Identity = Depends(_require_admin),
) -> dict[str, Any]:
    try:
        return create_ruleset(request.model_dump(exclude_none=True), actor_user_id=identity.user_id)
    except ValueError as exc:
        raise HTTPException(status_code=422, detail=str(exc)) from exc
    except AnnualAuditStorageError as exc:
        raise _storage_http_error(exc) from exc


@router.put("/api/annual-audit/rulesets/{ruleset_id}/rules", tags=["年审知识治理"])
def save_rule(
    ruleset_id: int,
    request: RuleRequest,
    identity: Identity = Depends(_require_admin),
) -> dict[str, Any]:
    try:
        return upsert_rule(
            ruleset_id,
            request.model_dump(exclude_none=True),
            actor_user_id=identity.user_id,
        )
    except ValueError as exc:
        raise HTTPException(status_code=422, detail=str(exc)) from exc
    except AnnualAuditStorageError as exc:
        raise _storage_http_error(exc) from exc


@router.post("/api/annual-audit/rules/{rule_id}/review", tags=["年审知识治理"])
def review_rule_route(
    rule_id: int,
    request: KnowledgeReviewRequest,
    identity: Identity = Depends(_require_admin),
) -> dict[str, Any]:
    try:
        return review_rule(
            rule_id,
            request.model_dump(exclude_none=True),
            actor_user_id=identity.user_id,
        )
    except ValueError as exc:
        raise HTTPException(status_code=422, detail=str(exc)) from exc
    except AnnualAuditStorageError as exc:
        raise _storage_http_error(exc) from exc


@router.post("/api/annual-audit/rulesets/{ruleset_id}/publish", tags=["年审知识治理"])
def publish_ruleset_route(
    ruleset_id: int,
    request: PublicationRequest,
    identity: Identity = Depends(_require_admin),
) -> dict[str, Any]:
    try:
        return publish_ruleset(
            ruleset_id,
            request.model_dump(exclude_none=True),
            actor_user_id=identity.user_id,
        )
    except ValueError as exc:
        raise HTTPException(status_code=422, detail=str(exc)) from exc
    except AnnualAuditStorageError as exc:
        raise _storage_http_error(exc) from exc


async def _read_template_uploads(
    files: list[UploadFile],
    *,
    template_usages: list[str] | None = None,
    remarks: list[str] | None = None,
    default_usage: str = "",
    default_remark: str = "",
) -> list[dict[str, Any]]:
    settings = get_settings()
    if not files:
        raise HTTPException(status_code=400, detail="模板版本至少需要上传一个文件")
    if len(files) > settings.max_upload_files:
        raise HTTPException(status_code=400, detail=f"单次最多上传 {settings.max_upload_files} 个模板文件")
    payloads: list[dict[str, Any]] = []
    for index, upload in enumerate(files):
        file_name = Path(upload.filename or "uploaded-template").name
        file_ext = Path(file_name).suffix.lower()
        if file_ext not in generic_templates.SUPPORTED_TEMPLATE_EXTENSIONS:
            raise HTTPException(status_code=422, detail=f"不支持的模板文件格式：{file_ext or file_name}")
        file_bytes = await upload.read()
        if len(file_bytes) > settings.max_upload_file_mb * 1024 * 1024:
            raise HTTPException(status_code=400, detail=f"文件 {file_name} 超过 {settings.max_upload_file_mb}MB 限制")
        usage_values = template_usages or []
        remark_values = remarks or []
        payloads.append(
            {
                "file_name": file_name,
                "file_ext": file_ext,
                "content_type": upload.content_type or mimetypes.guess_type(file_name)[0] or "application/octet-stream",
                "file_bytes": file_bytes,
                "storage_sha256": hashlib.sha256(file_bytes).hexdigest(),
                "template_usage": usage_values[index] if index < len(usage_values) else default_usage,
                "remark": remark_values[index] if index < len(remark_values) else default_remark,
            }
        )
    return payloads


async def _persist_template_uploads(
    template_code: str,
    version_no: int,
    payloads: list[dict[str, Any]],
    *,
    actor: str,
) -> list[dict[str, Any]]:
    service = get_minio_service()
    stored_refs: list[str] = []
    inserted_ids: list[int] = []
    rows: list[dict[str, Any]] = []
    try:
        for payload in payloads:
            stored = service.upload_template_file(
                template_code=template_code,
                version_no=version_no,
                file_name=str(payload["file_name"]),
                content_type=str(payload["content_type"]),
                file_bytes=payload["file_bytes"],
            )
            stored_refs.append(stored.storage_ref)
            row = generic_templates.add_template_file(
                template_code,
                version_no,
                file_name=str(payload["file_name"]),
                file_ext=str(payload["file_ext"]),
                content_type=str(payload["content_type"]),
                storage_ref=stored.storage_ref,
                storage_sha256=str(payload["storage_sha256"]),
                file_size=len(payload["file_bytes"]),
                template_usage=str(payload.get("template_usage") or ""),
                remark=str(payload.get("remark") or ""),
                created_by=actor or "system",
            )
            inserted_ids.append(int(row["id"]))
            rows.append(row)
    except Exception:
        for file_id in inserted_ids:
            try:
                generic_templates.delete_template_file(file_id, settings=get_settings())
            except Exception:
                pass
        for storage_ref in stored_refs:
            service.delete_object(storage_ref)
        raise
    return rows


@router.get("/api/templates", tags=["模板管理"])
def list_template_versions_route(identity: Identity = Depends(_require_admin)) -> dict[str, Any]:
    """Return a flat list of generic template versions for the admin menu."""

    try:
        return generic_templates.list_template_versions()
    except AnnualAuditStorageError as exc:
        raise _template_storage_http_error(exc) from exc


@router.post("/api/templates/versions", tags=["模板管理"])
async def create_template_version_route(
    template_name: str = Form(..., min_length=1, max_length=255),
    business_line: str = Form(..., min_length=1, max_length=128),
    version_label: str = Form(default="", max_length=128),
    description: str = Form(default="", max_length=1024),
    template_code: str = Form(default="", max_length=128),
    template_usages: list[str] | None = Form(default=None),
    remarks: list[str] | None = Form(default=None),
    template_usage: str = Form(default="", max_length=128),
    remark: str = Form(default="", max_length=1024),
    files: list[UploadFile] = File(..., description="模板原始文件"),
    identity: Identity = Depends(_require_admin),
) -> dict[str, Any]:
    record = None
    try:
        payloads = await _read_template_uploads(
            files,
            template_usages=template_usages,
            remarks=remarks,
            default_usage=template_usage,
            default_remark=remark,
        )
        record = generic_templates.create_template_version_record(
            template_name=template_name,
            business_line=business_line,
            version_label=version_label,
            template_code=template_code,
            description=description,
            created_by=identity.user_id or "system",
        )
        await _persist_template_uploads(
            str(record["template_code"]),
            int(record["version_no"]),
            payloads,
            actor=identity.user_id or "system",
        )
        return generic_templates.get_template_version(str(record["template_code"]), int(record["version_no"]))
    except HTTPException:
        raise
    except ValueError as exc:
        if record:
            try:
                generic_templates.delete_template_version(str(record["template_code"]), int(record["version_no"]))
            except Exception:
                pass
        raise HTTPException(status_code=422, detail=str(exc)) from exc
    except AnnualAuditStorageError as exc:
        if record:
            try:
                generic_templates.delete_template_version(str(record["template_code"]), int(record["version_no"]))
            except Exception:
                pass
        raise _template_storage_http_error(exc) from exc
    except Exception as exc:
        if record:
            try:
                generic_templates.delete_template_version(str(record["template_code"]), int(record["version_no"]))
            except Exception:
                pass
        raise _template_storage_http_error(exc) from exc


@router.get("/api/templates/{template_code}/versions/{version_no}", tags=["模板管理"])
def get_template_version_route(
    template_code: str,
    version_no: int,
    identity: Identity = Depends(_require_admin),
) -> dict[str, Any]:
    try:
        return generic_templates.get_template_version(template_code, version_no)
    except ValueError as exc:
        raise HTTPException(status_code=404, detail=str(exc)) from exc
    except AnnualAuditStorageError as exc:
        raise _template_storage_http_error(exc) from exc


@router.post("/api/templates/{template_code}/versions/{version_no}/files", tags=["模板管理"])
async def upload_template_version_files_route(
    template_code: str,
    version_no: int,
    template_usages: list[str] | None = Form(default=None),
    remarks: list[str] | None = Form(default=None),
    template_usage: str = Form(default="", max_length=128),
    remark: str = Form(default="", max_length=1024),
    files: list[UploadFile] = File(..., description="模板原始文件"),
    identity: Identity = Depends(_require_admin),
) -> dict[str, Any]:
    try:
        payloads = await _read_template_uploads(
            files,
            template_usages=template_usages,
            remarks=remarks,
            default_usage=template_usage,
            default_remark=remark,
        )
        rows = await _persist_template_uploads(
            template_code,
            version_no,
            payloads,
            actor=identity.user_id or "system",
        )
        return generic_templates.get_template_version(template_code, version_no) | {"uploaded_files": rows}
    except HTTPException:
        raise
    except ValueError as exc:
        raise HTTPException(status_code=422, detail=str(exc)) from exc
    except AnnualAuditStorageError as exc:
        raise _template_storage_http_error(exc) from exc
    except Exception as exc:
        raise _template_storage_http_error(exc) from exc


@router.delete("/api/templates/{template_code}/versions/{version_no}", tags=["模板管理"])
def delete_template_version_route(
    template_code: str,
    version_no: int,
    identity: Identity = Depends(_require_admin),
) -> dict[str, Any]:
    try:
        result = generic_templates.delete_template_version(template_code, version_no, deleted_by=identity.user_id or "system")
        for storage_ref in result.get("storage_refs") or []:
            try:
                get_minio_service().delete_object(str(storage_ref))
            except Exception:
                pass
        return result
    except ValueError as exc:
        raise HTTPException(status_code=422, detail=str(exc)) from exc
    except AnnualAuditStorageError as exc:
        raise _template_storage_http_error(exc) from exc


@router.delete("/api/templates/files/{file_id}", tags=["模板管理"])
def delete_template_file_route(file_id: int, identity: Identity = Depends(_require_admin)) -> dict[str, Any]:
    try:
        result = generic_templates.delete_template_file(file_id, deleted_by=identity.user_id or "system")
        if result.get("storage_ref"):
            try:
                get_minio_service().delete_object(str(result["storage_ref"]))
            except Exception:
                pass
        return result
    except ValueError as exc:
        raise HTTPException(status_code=422, detail=str(exc)) from exc
    except AnnualAuditStorageError as exc:
        raise _template_storage_http_error(exc) from exc


@router.post("/api/templates/{template_code}/versions/{version_no}/activate", tags=["模板管理"])
def activate_template_version_route(
    template_code: str,
    version_no: int,
    identity: Identity = Depends(_require_admin),
) -> dict[str, Any]:
    try:
        return generic_templates.activate_template_version(template_code, version_no, published_by=identity.user_id or "system")
    except ValueError as exc:
        raise HTTPException(status_code=422, detail=str(exc)) from exc
    except AnnualAuditStorageError as exc:
        raise _template_storage_http_error(exc) from exc


@router.get("/api/annual-audit/templates", tags=["年审模板管理"])
def list_audit_templates(
    template_type: str | None = Query(default=None),
    include_versions: bool = Query(default=True),
    identity: Identity = Depends(_require_admin),
) -> dict[str, Any]:
    """Backward-compatible annual route backed by the generic registry."""

    try:
        catalog = generic_templates.list_template_versions()
        if template_type:
            catalog["versions"] = [
                version
                for version in catalog.get("versions", [])
                if any(
                    str(file.get("template_usage") or "") == template_type
                    for file in (generic_templates.get_template_version(version["template_code"], version["version_no"]).get("files") or [])
                )
            ]
        return catalog
    except AnnualAuditStorageError as exc:
        raise _storage_http_error(exc) from exc


@router.get("/api/annual-audit/templates/active", tags=["年审模板管理"])
def list_active_audit_templates(
    identity: Identity = Depends(_require_report),
) -> dict[str, Any]:
    """Return only the active template version snapshot used by generation."""

    try:
        return {"templates": generic_templates.get_active_template_catalog()}
    except AnnualAuditStorageError as exc:
        raise _storage_http_error(exc) from exc


@router.post("/api/annual-audit/templates", tags=["年审模板管理"])
def create_audit_template(
    request: TemplateCreateRequest,
    identity: Identity = Depends(_require_admin),
) -> dict[str, Any]:
    raise HTTPException(
        status_code=410,
        detail="模板版本必须在创建时同时上传真实文件，请使用 /api/templates/versions",
    )


@router.post("/api/annual-audit/templates/{template_code}/versions", tags=["年审模板管理"])
def create_audit_template_version(
    template_code: str,
    request: TemplateVersionRequest,
    identity: Identity = Depends(_require_admin),
) -> dict[str, Any]:
    raise HTTPException(
        status_code=410,
        detail="模板版本必须在创建时同时上传真实文件，请使用 /api/templates/versions",
    )


@router.post("/api/annual-audit/templates/{template_code}/versions/{version_no}/files", tags=["年审模板管理"])
async def upload_audit_template_files(
    template_code: str,
    version_no: int,
    files: list[UploadFile] = File(..., description="模板原始文件，可上传 Word、Excel、PDF、Markdown 等"),
    identity: Identity = Depends(_require_admin),
) -> dict[str, Any]:
    """Upload real template files into a draft template version."""

    settings = get_settings()
    if len(files) > settings.max_upload_files:
        raise HTTPException(status_code=400, detail=f"单次最多上传 {settings.max_upload_files} 个模板文件")
    uploaded_rows: list[dict[str, Any]] = []
    service = None
    try:
        service = get_minio_service()
        for upload in files:
            file_name = Path(upload.filename or "uploaded-template").name
            file_ext = Path(file_name).suffix.lower()
            if file_ext not in templates.SUPPORTED_TEMPLATE_EXTENSIONS:
                raise HTTPException(status_code=422, detail=f"不支持的模板文件格式：{file_ext or file_name}")
            file_bytes = await upload.read()
            if len(file_bytes) > settings.max_upload_file_mb * 1024 * 1024:
                raise HTTPException(status_code=400, detail=f"文件 {file_name} 超过 {settings.max_upload_file_mb}MB 限制")
            content_type = upload.content_type or mimetypes.guess_type(file_name)[0] or "application/octet-stream"
            digest = hashlib.sha256(file_bytes).hexdigest()
            stored = service.upload_template_file(
                template_code=template_code,
                version_no=version_no,
                file_name=file_name,
                content_type=content_type,
                file_bytes=file_bytes,
            )
            try:
                row = templates.add_template_file(
                    template_code,
                    version_no,
                    file_name=file_name,
                    file_ext=file_ext,
                    content_type=content_type,
                    storage_ref=stored.storage_ref,
                    storage_sha256=digest,
                    file_size=len(file_bytes),
                    created_by=identity.user_id or "system",
                )
            except Exception:
                service.delete_object(stored.storage_ref)
                raise
            uploaded_rows.append(row)
    except HTTPException:
        raise
    except ValueError as exc:
        raise HTTPException(status_code=422, detail=str(exc)) from exc
    except AnnualAuditStorageError as exc:
        raise _storage_http_error(exc) from exc
    except Exception as exc:
        raise _storage_http_error(exc) from exc
    return {"template_code": template_code, "version_no": version_no, "files": uploaded_rows}


@router.delete("/api/annual-audit/templates/files/{file_id}", tags=["年审模板管理"])
def delete_audit_template_file(
    file_id: int,
    identity: Identity = Depends(_require_admin),
) -> dict[str, Any]:
    try:
        deleted = templates.delete_template_file(file_id, deleted_by=identity.user_id or "system")
        storage_ref = str(deleted.get("storage_ref") or "")
        if storage_ref:
            try:
                get_minio_service().delete_object(storage_ref)
            except Exception:
                pass
        return deleted
    except ValueError as exc:
        raise HTTPException(status_code=422, detail=str(exc)) from exc
    except AnnualAuditStorageError as exc:
        raise _storage_http_error(exc) from exc


@router.post(
    "/api/annual-audit/templates/{template_code}/versions/{version_no}/activate",
    tags=["年审模板管理"],
)
def activate_audit_template_version(
    template_code: str,
    version_no: int,
    identity: Identity = Depends(_require_admin),
) -> dict[str, Any]:
    try:
        return templates.activate_template_version(
            template_code,
            version_no,
            published_by=identity.user_id or "system",
        )
    except ValueError as exc:
        raise HTTPException(status_code=422, detail=str(exc)) from exc
    except AnnualAuditStorageError as exc:
        raise _storage_http_error(exc) from exc


@router.get("/api/annual-audit/{case_id}/artifacts", tags=["骞村鎶ュ憡"])
def list_report_artifacts(
    case_id: int,
    identity: Identity = Depends(_require_report),
) -> dict[str, Any]:
    """Return the latest report/workpaper artifact references for one engagement."""

    require_case_access(case_id, identity)
    try:
        engagements.get_engagement(case_id)
        with mysql_connection() as connection:
            with connection.cursor() as cursor:
                cursor.execute(
                    """
                    SELECT id, report_version, template_version, status, artifact_ref, created_at
                    FROM audit_report
                    WHERE engagement_id = %s AND report_type = 'annual_audit_draft'
                    ORDER BY report_version DESC LIMIT 1
                    """,
                    (case_id,),
                )
                report = dict(cursor.fetchone() or {})
                cursor.execute(
                    """
                    SELECT id, workpaper_code, workpaper_name, template_version,
                           workpaper_version, status, artifact_ref, created_at
                    FROM annual_workpaper
                    WHERE engagement_id = %s
                    ORDER BY workpaper_code, workpaper_version DESC
                    """,
                    (case_id,),
                )
                workpapers = [dict(row) for row in cursor.fetchall()]
                cursor.execute(
                    """
                    SELECT id, package_version, status, template_snapshot_json,
                           artifact_refs_json, created_by, created_at
                    FROM annual_audit_attachment_package
                    WHERE engagement_id = %s
                    ORDER BY package_version DESC
                    """,
                    (case_id,),
                )
                packages = []
                for row in cursor.fetchall():
                    package = dict(row)
                    for field in ("template_snapshot_json", "artifact_refs_json"):
                        raw = package.get(field)
                        if isinstance(raw, str):
                            try:
                                package[field.removesuffix("_json")] = json.loads(raw)
                            except json.JSONDecodeError:
                                package[field.removesuffix("_json")] = {}
                        package.pop(field, None)
                    for index, artifact in enumerate(package.get("artifact_refs") or []):
                        if isinstance(artifact, dict):
                            artifact["download_url"] = (
                                f"/api/annual-audit/{case_id}/attachment-packages/"
                                f"{package.get('id')}/files/{index}"
                            )
                            artifact["preview_url"] = (
                                f"/api/annual-audit/{case_id}/attachment-packages/"
                                f"{package.get('id')}/files/{index}/preview"
                            )
                    packages.append(package)
        return {
            "case_id": case_id,
            "report": report,
            "workpapers": workpapers,
            "attachment_packages": packages,
            "artifact_status": "draft_saved" if report.get("artifact_ref") else "not_published",
        }
    except engagements.EngagementNotFoundError as exc:
        raise HTTPException(status_code=404, detail=str(exc)) from exc
    except AnnualAuditStorageError as exc:
        raise _storage_http_error(exc) from exc


@router.get("/api/annual-audit/{case_id}/attachment-packages/{package_id}/files/{artifact_index}", tags=["年审报告"])
def download_annual_attachment(
    case_id: int,
    package_id: int,
    artifact_index: int,
    identity: Identity = Depends(_require_report),
) -> StreamingResponse:
    """Download a generated attachment from its persisted package reference."""

    require_case_access(case_id, identity)
    if artifact_index < 0:
        raise HTTPException(status_code=404, detail="附件不存在")
    try:
        with mysql_connection() as connection:
            with connection.cursor() as cursor:
                cursor.execute(
                    """
                    SELECT artifact_refs_json
                    FROM annual_audit_attachment_package
                    WHERE id = %s AND engagement_id = %s
                    """,
                    (package_id, case_id),
                )
                row = dict(cursor.fetchone() or {})
        if not row:
            raise HTTPException(status_code=404, detail="附件包不存在")
        raw_refs = row.get("artifact_refs_json")
        refs = json.loads(raw_refs) if isinstance(raw_refs, str) else raw_refs
        artifact = refs[artifact_index] if isinstance(refs, list) and artifact_index < len(refs) else None
        if not isinstance(artifact, dict) or not artifact.get("storage_ref"):
            raise HTTPException(status_code=404, detail="附件不存在")
        content = get_minio_service().get_object_bytes(str(artifact["storage_ref"]))
        file_name = Path(str(artifact.get("file_name") or f"attachment-{artifact_index}")).name
        media_type = str(artifact.get("content_type") or "application/octet-stream")
        return StreamingResponse(
            BytesIO(content),
            media_type=media_type,
            headers={"Content-Disposition": f"attachment; filename*=UTF-8''{quote(file_name)}"},
        )
    except HTTPException:
        raise
    except (json.JSONDecodeError, IndexError) as exc:
        raise HTTPException(status_code=404, detail="附件索引无效") from exc
    except AnnualAuditStorageError as exc:
        raise _storage_http_error(exc) from exc
    except Exception as exc:
        raise _storage_http_error(exc) from exc


@router.get("/api/annual-audit/{case_id}/attachment-packages/{package_id}/files/{artifact_index}/preview", tags=["年审报告"])
def preview_annual_attachment(
    case_id: int,
    package_id: int,
    artifact_index: int,
    identity: Identity = Depends(_require_report),
) -> JSONResponse:
    """Return a safe, read-only preview model for a generated attachment."""

    require_case_access(case_id, identity)
    if artifact_index < 0:
        raise HTTPException(status_code=404, detail="附件不存在")
    try:
        with mysql_connection() as connection:
            with connection.cursor() as cursor:
                cursor.execute(
                    """
                    SELECT artifact_refs_json
                    FROM annual_audit_attachment_package
                    WHERE id = %s AND engagement_id = %s
                    """,
                    (package_id, case_id),
                )
                row = dict(cursor.fetchone() or {})
        if not row:
            raise HTTPException(status_code=404, detail="附件包不存在")
        raw_refs = row.get("artifact_refs_json")
        refs = json.loads(raw_refs) if isinstance(raw_refs, str) else raw_refs
        artifact = refs[artifact_index] if isinstance(refs, list) and artifact_index < len(refs) else None
        if not isinstance(artifact, dict) or not artifact.get("storage_ref"):
            raise HTTPException(status_code=404, detail="附件不存在")
        content = get_minio_service().get_object_bytes(str(artifact["storage_ref"]))
        file_name = Path(str(artifact.get("file_name") or f"attachment-{artifact_index}")).name
        extension = Path(file_name).suffix.lower()
        if extension == ".docx":
            from docx import Document

            document = Document(BytesIO(content))
            paragraphs = [" ".join(p.text.split()) for p in document.paragraphs if p.text.strip()]
            tables: list[list[list[str]]] = []
            for table in document.tables[:20]:
                rows: list[list[str]] = []
                for row_item in table.rows[:40]:
                    rows.append([" ".join(cell.text.split())[:500] for cell in row_item.cells[:20]])
                if rows:
                    tables.append(rows)
            return JSONResponse({
                "kind": "document",
                "file_name": file_name,
                "content_type": artifact.get("content_type") or "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "paragraphs": paragraphs[:240],
                "tables": tables,
                "truncated": len(paragraphs) > 240 or len(document.tables) > 20,
            })
        if extension in {".xlsx", ".xlsm"}:
            from openpyxl import load_workbook

            workbook = load_workbook(BytesIO(content), read_only=True, data_only=False)
            sheets: list[dict[str, Any]] = []
            for worksheet in workbook.worksheets[:12]:
                rows: list[list[str]] = []
                for values in worksheet.iter_rows(max_row=60, max_col=20, values_only=True):
                    row_values = ["" if value is None else str(value)[:300] for value in values]
                    if any(row_values):
                        rows.append(row_values)
                sheets.append({"name": worksheet.title, "rows": rows})
            return JSONResponse({
                "kind": "workbook",
                "file_name": file_name,
                "content_type": artifact.get("content_type") or "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                "sheets": sheets,
                "truncated": len(workbook.worksheets) > 12,
            })
        if extension == ".xls":
            # xlrd is read-only by design here: preview must not rewrite the
            # legacy BIFF8 workbook or change the format the user downloads.
            import xlrd

            workbook = xlrd.open_workbook(file_contents=content, on_demand=True)
            sheets: list[dict[str, Any]] = []
            for sheet in workbook.sheets()[:12]:
                rows: list[list[str]] = []
                for row_index in range(min(sheet.nrows, 60)):
                    values = [
                        "" if value is None else str(value)[:300]
                        for value in sheet.row_values(row_index, 0, min(sheet.ncols, 20))
                    ]
                    if any(values):
                        rows.append(values)
                sheets.append({"name": sheet.name, "rows": rows})
            return JSONResponse({
                "kind": "workbook",
                "file_name": file_name,
                "content_type": artifact.get("content_type") or "application/vnd.ms-excel",
                "sheets": sheets,
                "truncated": len(workbook.sheets()) > 12,
            })
        if extension in {".txt", ".md", ".markdown", ".csv"}:
            text = content.decode("utf-8-sig", errors="replace")
            return JSONResponse({
                "kind": "text",
                "file_name": file_name,
                "content_type": artifact.get("content_type") or "text/plain",
                "text": text[:30000],
                "truncated": len(text) > 30000,
            })
        return JSONResponse({
            "kind": "unsupported",
            "file_name": file_name,
            "content_type": artifact.get("content_type") or "application/octet-stream",
            "message": "该文件格式暂不支持在线预览，请点击下载。",
        })
    except HTTPException:
        raise
    except (json.JSONDecodeError, IndexError) as exc:
        raise HTTPException(status_code=404, detail="附件索引无效") from exc
    except AnnualAuditStorageError as exc:
        raise _storage_http_error(exc) from exc
    except Exception as exc:
        raise _storage_http_error(exc) from exc
