"""检查审计报告附件的填充情况。"""
import json
import sys
from pathlib import Path
from io import BytesIO

backend_dir = Path(__file__).parent.parent
sys.path.insert(0, str(backend_dir))

from ai_hunter.app.settings import get_settings
from ai_hunter.app.services.minio_service import get_minio_service
from ai_hunter.annual_audit.storage import mysql_connection

settings = get_settings()
service = get_minio_service()

with mysql_connection(settings) as conn:
    with conn.cursor() as cur:
        cur.execute("""
            SELECT id, package_version, status, artifact_refs_json
            FROM annual_audit_attachment_package
            WHERE engagement_id = 1
            ORDER BY package_version DESC LIMIT 1
        """)
        row = cur.fetchone()

print(f"包版本: {row['package_version']}")
artifacts = json.loads(row['artifact_refs_json']) if isinstance(row['artifact_refs_json'], str) else row['artifact_refs_json']

for art in artifacts:
    file_name = art.get('file_name', '')
    if '审计报告' not in file_name or '附注' in file_name:
        continue
    storage_ref = art.get('storage_ref', '')
    print(f"\n{'='*60}")
    print(f"文件: {file_name}")

    file_bytes = service.get_object_bytes(storage_ref)
    from docx import Document
    doc = Document(BytesIO(file_bytes))
    print(f"段落数: {len(doc.paragraphs)}")
    print(f"表格数: {len(doc.tables)}")

    # 打印所有段落
    print("\n段落内容:")
    for p_idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            print(f"  P{p_idx}: {text[:100]}")
