"""检查附注附件的表格填充情况。"""
import json
import sys
from pathlib import Path
from io import BytesIO

backend_dir = Path(__file__).parent.parent
sys.path.insert(0, str(backend_dir))

from ai_hunter.app.settings import get_settings
from ai_hunter.app.services.minio_service import get_minio_service
from ai_hunter.annual_audit.storage import mysql_connection

settings = get_settings()
service = get_minio_service()

with mysql_connection(settings) as conn:
    with conn.cursor() as cur:
        cur.execute("""
            SELECT id, package_version, status, artifact_refs_json
            FROM annual_audit_attachment_package
            WHERE engagement_id = 1
            ORDER BY package_version DESC LIMIT 1
        """)
        row = cur.fetchone()

print(f"包版本: {row['package_version']}")
artifacts = json.loads(row['artifact_refs_json']) if isinstance(row['artifact_refs_json'], str) else row['artifact_refs_json']

for art in artifacts:
    file_name = art.get('file_name', '')
    if '附注' not in file_name:
        continue
    storage_ref = art.get('storage_ref', '')
    print(f"\n{'='*60}")
    print(f"文件: {file_name}")

    file_bytes = service.get_object_bytes(storage_ref)
    from docx import Document
    doc = Document(BytesIO(file_bytes))
    print(f"表格数: {len(doc.tables)}")

    # 只打印有数据填充的表格
    filled_count = 0
    for t_idx, table in enumerate(doc.tables):
        has_data = False
        for row in table.rows:
            for c_idx, cell in enumerate(row.cells):
                text = cell.text.strip()
                # 检查是否有数字数据（金额）
                if text and any(ch.isdigit() for ch in text) and c_idx > 0:
                    has_data = True
                    break
            if has_data:
                break
        if has_data:
            filled_count += 1
            print(f"\n--- 表格 {t_idx+1} (有数据): {len(table.rows)}行 x {len(table.columns)}列 ---")
            for r_idx, row in enumerate(table.rows):
                if r_idx >= 10:
                    print(f"  ... (共{len(table.rows)}行)")
                    break
                cells_text = []
                for c_idx, cell in enumerate(row.cells):
                    text = cell.text.strip()
                    if len(text) > 25:
                        text = text[:25] + ".."
                    cells_text.append(text)
                line = " | ".join(cells_text)
                print(f"  R{r_idx}: {line}")

    print(f"\n有数据填充的表格数: {filled_count} / {len(doc.tables)}")
