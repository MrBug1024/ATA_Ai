"""检查生成的附件内容，验证表格填充效果。"""
import json
import sys
from pathlib import Path
from io import BytesIO

backend_dir = Path(__file__).parent.parent
sys.path.insert(0, str(backend_dir))

from ai_hunter.app.settings import get_settings
from ai_hunter.app.services.minio_service import get_minio_service
from ai_hunter.annual_audit.storage import mysql_connection

settings = get_settings()
service = get_minio_service()

# 获取最新的附件包
with mysql_connection(settings) as conn:
    with conn.cursor() as cur:
        cur.execute("""
            SELECT id, package_version, status, artifact_refs_json
            FROM annual_audit_attachment_package
            WHERE engagement_id = 1
            ORDER BY package_version DESC LIMIT 1
        """)
        row = cur.fetchone()

print(f"包ID: {row['id']}, 版本: {row['package_version']}, 状态: {row['status']}")
artifacts = json.loads(row['artifact_refs_json']) if isinstance(row['artifact_refs_json'], str) else row['artifact_refs_json']

for art in artifacts:
    file_name = art.get('file_name', '')
    storage_ref = art.get('storage_ref', '')
    fill_status = art.get('template_fill_status', '')
    print(f"\n{'='*60}")
    print(f"文件: {file_name}")
    print(f"填充状态: {fill_status}")
    print(f"存储: {storage_ref}")

    # 下载文件内容
    try:
        file_bytes = service.get_object_bytes(storage_ref)
        print(f"文件大小: {len(file_bytes)} bytes")

        # 如果是docx，检查表格内容
        if file_name.endswith('.docx'):
            from docx import Document
            doc = Document(BytesIO(file_bytes))
            print(f"段落数: {len(doc.paragraphs)}")
            print(f"表格数: {len(doc.tables)}")

            for t_idx, table in enumerate(doc.tables):
                print(f"\n  表格 {t_idx+1}: {len(table.rows)}行 x {len(table.columns)}列")
                # 打印前10行，看看哪些单元格被填充了
                for r_idx, row in enumerate(table.rows):
                    if r_idx >= 15:
                        print(f"  ... (共{len(table.rows)}行)")
                        break
                    cells_text = []
                    for c_idx, cell in enumerate(row.cells):
                        text = cell.text.strip()
                        if len(text) > 30:
                            text = text[:30] + "..."
                        cells_text.append(f"[{c_idx}]{text}")
                    print(f"  行{r_idx}: {' | '.join(cells_text)}")
    except Exception as e:
        print(f"检查失败: {e}")
        import traceback
        traceback.print_exc()
