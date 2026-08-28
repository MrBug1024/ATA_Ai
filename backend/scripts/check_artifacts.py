"""检查最新生成的附件包内容。"""
import json
from io import BytesIO
from docx import Document

from ai_hunter.app.services.minio_service import get_minio_service
from ai_hunter.annual_audit.storage import mysql_connection
from ai_hunter.app.settings import get_settings

settings = get_settings()
service = get_minio_service()

with mysql_connection(settings) as conn:
    with conn.cursor() as cur:
        cur.execute(
            """
            SELECT id, package_version, status, artifact_refs_json
            FROM annual_audit_attachment_package
            WHERE engagement_id = 1
            ORDER BY package_version DESC LIMIT 1
            """,
        )
        row = dict(cur.fetchone() or {})

pv = row.get("package_version")
st = row.get("status")
print(f"附件包版本: {pv}, 状态: {st}")
artifacts = json.loads(row.get("artifact_refs_json") or "[]")
for i, a in enumerate(artifacts):
    fn = a.get("file_name")
    sr = a.get("storage_ref")
    fs = a.get("template_fill_status")
    print(f"\n  [{i}] {fn} (fill_status={fs})")
    try:
        data = service.get_object_bytes(sr)
        doc = Document(BytesIO(data))
        print(f"    段落数: {len(doc.paragraphs)}, 表格数: {len(doc.tables)}")
        # 打印前5个非空段落
        count = 0
        for p in doc.paragraphs:
            if p.text.strip():
                print(f"    P: {p.text[:120]}")
                count += 1
                if count >= 8:
                    break
        # 检查表格是否有数据
        for ti, table in enumerate(doc.tables[:3]):
            print(f"    表格{ti}: {len(table.rows)}行x{len(table.columns)}列")
            for ri, trow in enumerate(table.rows[:5]):
                cells = [cell.text[:30] for cell in trow.cells]
                print(f"      行{ri}: {cells}")
    except Exception as e:
        print(f"    ERROR: {e}")
