from __future__ import annotations

import zipfile
from io import BytesIO
from types import SimpleNamespace

import pytest
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from openpyxl import Workbook
from openpyxl.worksheet.table import Table
from pypdf import PdfWriter
from reportlab.pdfgen import canvas

from ai_hunter.document_templates.compiler import inspect_template, compile_template
from ai_hunter.document_templates.compiler.models import (
    TemplateCompilationError,
    TemplateSecurityLimits,
)


DOCX_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
XLSX_TYPE = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"


@pytest.fixture
def security_settings() -> SimpleNamespace:
    return SimpleNamespace(
        attachment_template_max_mb=10,
        attachment_zip_max_entries=2_000,
        attachment_zip_max_uncompressed_mb=100,
        attachment_zip_max_entry_mb=20,
        attachment_zip_max_ratio=100,
        attachment_xml_max_mb=20,
        attachment_clamav_host="",
        attachment_clamav_port=3310,
        attachment_clamav_timeout_seconds=1,
        attachment_clamav_required=False,
    )


def _docx(text: str, *, bookmark: str = "") -> bytes:
    document = Document()
    paragraph = document.add_paragraph(text)
    if bookmark:
        start = OxmlElement("w:bookmarkStart")
        start.set(qn("w:id"), "1")
        start.set(qn("w:name"), bookmark)
        end = OxmlElement("w:bookmarkEnd")
        end.set(qn("w:id"), "1")
        paragraph._p.insert(0, start)
        paragraph._p.append(end)
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _xlsx(*, external_formula: bool = False) -> bytes:
    workbook = Workbook()
    worksheet = workbook.active
    worksheet.title = "Audit"
    if external_formula:
        worksheet["A1"] = "='[outside.xlsx]Sheet1'!A1"
    output = BytesIO()
    workbook.save(output)
    workbook.close()
    return output.getvalue()


def _pdf(*, encrypted: bool = False) -> bytes:
    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    if encrypted:
        writer.encrypt("secret")
    output = BytesIO()
    writer.write(output)
    return output.getvalue()


def _acroform_pdf() -> bytes:
    output = BytesIO()
    document = canvas.Canvas(output, pagesize=(612, 792))
    document.acroForm.textfield(
        name="company_name", x=72, y=700, width=300, height=24
    )
    document.showPage()
    document.save()
    return output.getvalue()


def _inspect(name: str, content_type: str, data: bytes, settings) -> object:
    return inspect_template(
        file_name=name,
        content_type=content_type,
        file_bytes=data,
        settings=settings,
    )


def _slot(
    slot_id: str,
    target: str,
    *,
    value_type: str = "scalar",
    source: str | None = None,
    options: dict | None = None,
) -> dict:
    return {
        "slot_id": slot_id,
        "target": target,
        "source": source or f"document.{slot_id}",
        "value_type": value_type,
        "required": True,
        "options": options or {},
    }


def test_all_supported_formats_use_real_synthetic_render(security_settings) -> None:
    cases = [
        (
            "audit.docx",
            DOCX_TYPE,
            _docx("Audit: {{ title }}"),
            _slot("title", "docx:jinja:title"),
        ),
        (
            "statements.xlsx",
            XLSX_TYPE,
            _xlsx(),
            _slot("title", "xlsx:cell:Audit!A1"),
        ),
        (
            "notes.md",
            "text/markdown",
            b"Audit: {{ title }}",
            _slot("title", "md:jinja:title"),
        ),
        (
            "form.pdf",
            "application/pdf",
            _pdf(),
            {
                **_slot("title", "pdf:overlay:title"),
                "overflow_policy": "shrink_font",
                "options": {
                    "page": 1,
                    "x": 72,
                    "y": 700,
                    "width": 400,
                    "height": 30,
                    "font_name": "Helvetica",
                    "font_size": 10,
                    "max_lines": 2,
                },
            },
        ),
    ]
    for name, content_type, data, slot in cases:
        inspection = _inspect(name, content_type, data, security_settings)
        assert inspection.ok, inspection.to_dict()
        compiled = compile_template(
            file_bytes=data,
            inspection=inspection,
            document_code="audit_report",
            binding_manifest={"document_code": "audit_report", "slots": [slot]},
        )
        synthetic = compiled.compilation_report["synthetic_render"]
        assert synthetic["status"] == "passed"
        assert synthetic["method"] == "runtime-renderer-and-format-reopen"
        assert compiled.synthetic_content
        assert len(compiled.binding_manifest["binding_sha256"]) == 64


def test_binding_manifest_is_runtime_canonical_and_hash_is_idempotent(
    security_settings,
) -> None:
    data = b"Hello {{ title }}"
    inspection = _inspect("template.md", "text/markdown", data, security_settings)
    legacy = {
        "document_code": "audit_report",
        "slots": [
            {
                "slot_id": "title",
                "target": "md:jinja:title",
                "source": "document.title",
                "value_type": "text",
                "style_policy": "fixed_contract",
                "required": False,
                "missing_policy": "leave_blank",
                "overflow_policy": "block",
            }
        ],
    }
    first = compile_template(
        file_bytes=data,
        inspection=inspection,
        document_code="audit_report",
        binding_manifest=legacy,
    )
    second = compile_template(
        file_bytes=data,
        inspection=inspection,
        document_code="audit_report",
        binding_manifest=first.binding_manifest,
    )
    slot = first.binding_manifest["slots"][0]
    assert slot["value_type"] == "scalar"
    assert slot["style_policy"] == "explicit"
    assert slot["missing_policy"] == "empty"
    assert slot["overflow_policy"] == "error"
    assert first.binding_manifest["binding_sha256"] == second.binding_manifest["binding_sha256"]


def test_semantic_narrative_slot_allows_empty_deterministic_source(
    security_settings,
) -> None:
    data = b"Overview: {{ overview }}"
    inspection = _inspect("template.md", "text/markdown", data, security_settings)
    slot = {
        "slot_id": "overview",
        "target": "md:jinja:overview",
        "source": "",
        "value_type": "narrative_blocks",
        "required": True,
        "missing_policy": "block",
        "options": {
            "composition_mode": "semantic",
            "semantic_instruction": "按批准标签编排简短陈述",
            "allowed_fact_refs": ["entity.legal_name"],
            "fact_ref_labels": {"entity.legal_name": "企业名称"},
        },
    }

    compiled = compile_template(
        file_bytes=data,
        inspection=inspection,
        document_code="audit_report",
        binding_manifest={"document_code": "audit_report", "slots": [slot]},
    )
    assert compiled.binding_manifest["slots"][0]["source"] == ""
    assert compiled.binding_manifest["slots"][0]["options"]["composition_mode"] == "semantic"

    deterministic = {**slot, "options": {}}
    with pytest.raises(TemplateCompilationError) as caught:
        compile_template(
            file_bytes=data,
            inspection=inspection,
            document_code="audit_report",
            binding_manifest={
                "document_code": "audit_report",
                "slots": [deterministic],
            },
        )
    assert caught.value.code == "INVALID_SLOT_SOURCE"


def test_docx_ssti_call_and_non_allowlisted_filter_are_rejected(security_settings) -> None:
    data = _docx("{{ title }} / {{ range(2)|list }}")
    inspection = _inspect("unsafe.docx", DOCX_TYPE, data, security_settings)
    assert not inspection.ok
    assert "unsafe_jinja_expression" in inspection.threats
    report = inspection.format_details["jinja_security"]
    assert "Call" in report["forbidden_constructs"]
    assert report["rejected_filters"] == ["list"]


@pytest.mark.parametrize(
    ("name", "content_type", "data", "target"),
    [
        (
            "loop.docx",
            DOCX_TYPE,
            _docx("{% for row in rows %}{{ row.name }}{% endfor %}"),
            "docx:jinja:rows",
        ),
        (
            "loop.md",
            "text/markdown",
            b"{% for row in rows %}{{ row.name }}{% endfor %}",
            "md:jinja:rows",
        ),
    ],
    ids=["docx", "markdown"],
)
def test_loop_local_mapping_access_requires_declared_columns(
    security_settings, name, content_type, data, target
) -> None:
    inspection = _inspect(name, content_type, data, security_settings)
    assert inspection.ok, inspection.to_dict()
    manifest = {
        "document_code": "audit_report",
        "slots": [
            _slot(
                "rows",
                target,
                source="document.rows[]",
                value_type="table_rows",
                options={"columns": ["name"]},
            )
        ],
    }
    compiled = compile_template(
        file_bytes=data,
        inspection=inspection,
        document_code="audit_report",
        binding_manifest=manifest,
    )
    assert compiled.compilation_report["synthetic_render"]["status"] == "passed"

    manifest["slots"][0]["options"] = {"columns": ["amount"]}
    with pytest.raises(TemplateCompilationError) as caught:
        compile_template(
            file_bytes=data,
            inspection=inspection,
            document_code="audit_report",
            binding_manifest=manifest,
        )
    assert caught.value.code == "JINJA_LOOP_COLUMN_NOT_DECLARED"


def test_top_level_mapping_access_is_rejected_for_markdown(security_settings) -> None:
    inspection = _inspect(
        "unsafe.md",
        "text/markdown",
        b"{{ document.title }} {{ document['name'] }}",
        security_settings,
    )
    assert not inspection.ok
    assert "unsafe_jinja_expression" in inspection.threats


def test_docx_bookmark_is_executable_when_runtime_supports_it(security_settings) -> None:
    data = _docx("replace me", bookmark="company_name")
    inspection = _inspect("bookmark.docx", DOCX_TYPE, data, security_settings)
    assert inspection.ok
    compiled = compile_template(
        file_bytes=data,
        inspection=inspection,
        document_code="audit_report",
        binding_manifest={
            "document_code": "audit_report",
            "slots": [_slot("company_name", "docx:bookmark:company_name")],
        },
    )
    assert compiled.compilation_report["synthetic_render"]["status"] == "passed"


def test_xlsx_table_and_pdf_acroform_contracts_render(security_settings) -> None:
    workbook = Workbook()
    worksheet = workbook.active
    worksheet.title = "Audit"
    worksheet.append(["Name", "Amount"])
    worksheet.append(["prototype", 0])
    worksheet.add_table(Table(displayName="ATA_TABLE", ref="A1:B2"))
    output = BytesIO()
    workbook.save(output)
    workbook.close()
    xlsx_data = output.getvalue()
    xlsx_inspection = _inspect(
        "table.xlsx", XLSX_TYPE, xlsx_data, security_settings
    )
    xlsx_compiled = compile_template(
        file_bytes=xlsx_data,
        inspection=xlsx_inspection,
        document_code="financial_statements",
        binding_manifest={
            "document_code": "financial_statements",
            "slots": [
                {
                    **_slot(
                        "rows",
                        "xlsx:table:ATA_TABLE",
                        source="document.rows[]",
                        value_type="table_rows",
                        options={"column_map": {"Name": "name", "Amount": "amount"}},
                    ),
                    "style_policy": "clone_prototype_row",
                    "overflow_policy": "extend_rows",
                }
            ],
        },
    )
    assert xlsx_compiled.compilation_report["synthetic_render"]["status"] == "passed"

    pdf_data = _acroform_pdf()
    pdf_inspection = _inspect(
        "form.pdf", "application/pdf", pdf_data, security_settings
    )
    pdf_compiled = compile_template(
        file_bytes=pdf_data,
        inspection=pdf_inspection,
        document_code="audit_report",
        binding_manifest={
            "document_code": "audit_report",
            "slots": [
                {
                    **_slot("company_name", "pdf:acroform:company_name"),
                    "overflow_policy": "shrink_font",
                }
            ],
        },
    )
    assert pdf_compiled.renderer_profile == "pdf_acroform"
    assert pdf_compiled.compilation_report["synthetic_render"]["status"] == "passed"


def test_upload_security_rejects_signature_macro_external_formula_and_encryption(
    security_settings,
) -> None:
    signature = _inspect("fake.pdf", "application/pdf", b"not a pdf", security_settings)
    assert {item["code"] for item in signature.errors} >= {"SIGNATURE_MISMATCH"}

    source = _docx("safe")
    output = BytesIO()
    with zipfile.ZipFile(BytesIO(source)) as archive, zipfile.ZipFile(
        output, "w", zipfile.ZIP_DEFLATED
    ) as rewritten:
        for entry in archive.infolist():
            rewritten.writestr(entry, archive.read(entry))
        rewritten.writestr("word/vbaProject.bin", b"macro")
    macro = _inspect("macro.docx", DOCX_TYPE, output.getvalue(), security_settings)
    assert not macro.ok
    assert "macro_or_embedded_content" in macro.threats

    external = _inspect(
        "external.xlsx", XLSX_TYPE, _xlsx(external_formula=True), security_settings
    )
    assert not external.ok
    assert "xlsx_external_formula_reference" in external.threats

    encrypted = _inspect(
        "locked.pdf", "application/pdf", _pdf(encrypted=True), security_settings
    )
    assert not encrypted.ok
    assert "encrypted_or_password_protected_content" in encrypted.threats


def test_zip_bomb_ratio_is_rejected(security_settings) -> None:
    output = BytesIO()
    with zipfile.ZipFile(output, "w", zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", b"<Types/>")
        archive.writestr("_rels/.rels", b"<Relationships/>")
        archive.writestr("word/document.xml", b"<document/>")
        archive.writestr("word/large.xml", b"<x>" + b"A" * 100_000 + b"</x>")
    inspection = inspect_template(
        file_name="bomb.docx",
        content_type=DOCX_TYPE,
        file_bytes=output.getvalue(),
        settings=security_settings,
        antivirus=False,
        limits=TemplateSecurityLimits(
            max_file_bytes=2_000_000,
            max_zip_entries=100,
            max_uncompressed_bytes=2_000_000,
            max_entry_bytes=1_000_000,
            max_compression_ratio=2,
            max_xml_bytes=1_000_000,
        ),
    )
    assert "ZIP_COMPRESSION_RATIO" in {item["code"] for item in inspection.errors}


def test_required_antivirus_fails_closed_when_clamav_is_not_configured(
    security_settings,
) -> None:
    security_settings.attachment_clamav_required = True
    inspection = _inspect(
        "template.md", "text/markdown", b"{{ title }}", security_settings
    )
    assert not inspection.ok
    assert "ANTIVIRUS_REQUIRED" in {item["code"] for item in inspection.errors}
