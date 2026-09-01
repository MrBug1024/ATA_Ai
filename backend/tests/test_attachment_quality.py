from __future__ import annotations

import hashlib
import zipfile
from io import BytesIO
from xml.etree import ElementTree as ET

import pytest
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Cm
from openpyxl import Workbook, load_workbook
from openpyxl.comments import Comment
from openpyxl.workbook.defined_name import DefinedName
from openpyxl.worksheet.table import Table, TableStyleInfo

from ai_hunter.annual_audit.attachments.agent_graph import prepare_attachment
from ai_hunter.annual_audit.attachments.content_schemas import (
    BindingManifest,
    DocumentPayload,
    LiteralSegment,
    RenderResult,
    ScalarSlotPayload,
)
from ai_hunter.annual_audit.attachments.context_service import build_generation_context_snapshot
from ai_hunter.annual_audit.attachments.fact_registry import FactRecord, FactRegistry
from ai_hunter.annual_audit.attachments.provenance import (
    build_prepared_provenance_manifest,
    provenance_bytes,
)
from ai_hunter.annual_audit.attachments.quality_service import (
    AttachmentQualityError,
    require_quality,
    sanitize_filename,
    sanitize_output_filename,
    validate_payload_contract,
    validate_rendered_artifact,
)
from ai_hunter.annual_audit.attachments.renderers.markdown import render_markdown


_DOCX_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)
_XLSX_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
_PACKAGE_REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"


def _markdown_manifest(template: bytes) -> BindingManifest:
    return BindingManifest(
        document_code="audit_report",
        source_template_sha256=hashlib.sha256(template).hexdigest(),
        slots=[
            {
                "slot_id": "company_name",
                "target": "md:variable:company_name",
                "source": "document.entity.legal_name",
                "value_type": "scalar",
                "required": True,
                "missing_policy": "block",
            }
        ],
    )


def _prepared(template: bytes):
    registry = FactRegistry(
        [
            FactRecord(
                fact_key="entity.legal_name",
                value="Example Co",
                display_value="Example Co",
                status="confirmed",
                source_kind="project_master",
                revision=1,
                evidence_refs=[{"source_file_id": 1, "source_page_id": 2}],
            )
        ]
    )
    context = build_generation_context_snapshot(
        engagement={"engagement_id": 1, "project_name": "Example Audit"},
        entity_facts=registry.frozen_snapshot(),
        report={"id": 2, "version": 3, "fact_snapshot": {"entity": "Example Co"}},
        generation_policy_version="annual-attachment-v1",
    )
    manifest = _markdown_manifest(template)
    return manifest, prepare_attachment(manifest, context, registry)


def _ooxml_manifest(template: bytes, extension: str) -> BindingManifest:
    return BindingManifest(
        document_code="audit_report",
        source_template_sha256=hashlib.sha256(template).hexdigest(),
        slots=[
            {
                "slot_id": "company_name",
                "target": f"{extension.removeprefix('.')}:cell:unused",
                "source": "",
                "value_type": "scalar",
                "required": True,
                "missing_policy": "block",
            }
        ],
    )


def _render_result(template: bytes, data: bytes, extension: str) -> RenderResult:
    return RenderResult(
        data=data,
        extension=extension,
        content_type=_DOCX_CONTENT_TYPE if extension == ".docx" else _XLSX_CONTENT_TYPE,
        source_template_sha256=hashlib.sha256(template).hexdigest(),
    )


def _structured_docx() -> bytes:
    document = Document()
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.header.paragraphs[0].text = "Annual audit header"
    header_table = section.header.add_table(rows=1, cols=2, width=Cm(16))
    header_table.cell(0, 0).text = "Entity"
    header_table.cell(0, 1).text = "Reporting period"
    section.footer.paragraphs[0].text = "Annual audit footer"
    style = document.styles.add_style("AuditBody", WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = document.styles["Normal"]
    document.add_paragraph("Annual audit body", style="AuditBody")
    document.add_section(WD_SECTION.NEW_PAGE)
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _docx_with_external_relationship_and_comment(template: bytes) -> bytes:
    source = BytesIO(template)
    output = BytesIO()
    relationship_part = "word/_rels/document.xml.rels"
    with zipfile.ZipFile(source, "r") as archive, zipfile.ZipFile(output, "w") as rendered:
        for info in archive.infolist():
            data = archive.read(info.filename)
            if info.filename == relationship_part:
                root = ET.fromstring(data)
                ET.SubElement(
                    root,
                    f"{{{_PACKAGE_REL_NS}}}Relationship",
                    {
                        "Id": "rIdExternalQualityTest",
                        "Type": (
                            "http://schemas.openxmlformats.org/officeDocument/2006/"
                            "relationships/hyperlink"
                        ),
                        "Target": "https://example.invalid/audit",
                        "TargetMode": "External",
                    },
                )
                data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
            rendered.writestr(info, data)
        rendered.writestr(
            "word/comments.xml",
            (
                '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                '<w:comments xmlns:w="http://schemas.openxmlformats.org/'
                'wordprocessingml/2006/main">'
                '<w:comment w:id="0" w:author="Reviewer"><w:p><w:r>'
                '<w:t>Remove before publication</w:t>'
                "</w:r></w:p></w:comment></w:comments>"
            ).encode("utf-8"),
        )
    return output.getvalue()


def _structured_xlsx() -> bytes:
    workbook = Workbook()
    worksheet = workbook.active
    worksheet.title = "Data"
    worksheet.append(["Account", "Amount", None, "Report title", None])
    worksheet.append(["Cash", 100])
    worksheet.append(["Revenue", 200])
    worksheet.merge_cells("D1:E1")
    worksheet.print_area = "A1:E20"
    worksheet.print_title_rows = "1:1"
    worksheet.page_setup.orientation = "landscape"
    worksheet.page_setup.paperSize = worksheet.PAPERSIZE_A4
    worksheet.freeze_panes = "B2"
    worksheet.oddHeader.center.text = "Annual audit"
    worksheet.oddFooter.right.text = "Page &P of &N"
    table = Table(displayName="AuditTable", ref="A1:B3")
    table.tableStyleInfo = TableStyleInfo(
        name="TableStyleMedium2",
        showFirstColumn=False,
        showLastColumn=False,
        showRowStripes=True,
        showColumnStripes=False,
    )
    worksheet.add_table(table)
    workbook.defined_names.add(
        DefinedName("InputArea", attr_text="'Data'!$A$2:$B$20")
    )
    worksheet.defined_names.add(
        DefinedName("LocalInput", attr_text="'Data'!$B$2", localSheetId=0)
    )
    workbook.create_sheet("Notes")["A1"] = "Review notes"
    output = BytesIO()
    workbook.save(output)
    workbook.close()
    return output.getvalue()


def _save_workbook(workbook: Workbook) -> bytes:
    output = BytesIO()
    workbook.save(output)
    workbook.close()
    return output.getvalue()


def test_filename_sanitization_blocks_traversal_reserved_names_and_invalid_characters() -> None:
    assert sanitize_filename("../../CON.docx") == "_CON.docx"
    assert sanitize_filename(" report<>:\"/\\|?*.XLSX ") == "_.xlsx"
    assert sanitize_output_filename("Example/../Audit", "Report", ".docx") == "AuditReport.docx"
    with pytest.raises(ValueError):
        sanitize_output_filename("Example", "Report", ".exe")


def test_quality_gate_detects_internal_ids_and_unresolved_placeholders() -> None:
    template = b"{{ company_name }}"
    manifest = _markdown_manifest(template)
    result = RenderResult(
        data=b"Example evidence_id=7 {{ still_here }}",
        extension=".md",
        content_type="text/markdown; charset=utf-8",
        source_template_sha256=hashlib.sha256(template).hexdigest(),
    )
    report = validate_rendered_artifact(
        result,
        file_name="ExampleAuditReport.md",
        manifest=manifest,
    )
    assert not report.passed
    codes = {issue.code for issue in report.issues}
    assert "FORBIDDEN_INTERNAL_ID" in codes
    assert "UNRESOLVED_PLACEHOLDER" in codes
    with pytest.raises(AttachmentQualityError):
        require_quality(report)


def test_quality_gate_reports_invalid_container_instead_of_publishing() -> None:
    manifest = BindingManifest(
        document_code="audit_report",
        source_template_sha256="a" * 64,
        slots=[
            {
                "slot_id": "company_name",
                "target": "docx:content-control:company_name",
                "source": "document.entity.legal_name",
                "value_type": "scalar",
                "required": True,
                "missing_policy": "block",
            }
        ],
    )
    result = RenderResult(
        data=b"not a zip",
        extension=".docx",
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        source_template_sha256="a" * 64,
    )
    report = validate_rendered_artifact(
        result,
        file_name="ExampleAuditReport.docx",
        manifest=manifest,
    )
    assert not report.passed
    assert any(issue.code == "FORMAT_REOPEN_FAILED" for issue in report.issues)


def test_payload_gate_requires_fact_ref_for_authoritative_identity_slots() -> None:
    manifest = _markdown_manifest(b"{{ company_name }}")
    payload = DocumentPayload(
        document_code="audit_report",
        slots=[
            ScalarSlotPayload(
                slot_id="company_name",
                segments=[LiteralSegment(value="invented")],
            )
        ],
    )
    issues = validate_payload_contract(payload, manifest)
    assert any(issue.code == "FACT_REF_REQUIRED" for issue in issues)


def test_valid_render_quality_and_provenance_hash_chain() -> None:
    template = b"Company: {{ company_name }}"
    manifest, prepared = _prepared(template)
    rendered = render_markdown(template, prepared.resolved_payload, manifest)
    quality = require_quality(
        validate_rendered_artifact(
            rendered,
            file_name="ExampleAuditReport.md",
            manifest=manifest,
            template_bytes=template,
            payload=prepared.resolved_payload,
        )
    )
    provenance = build_prepared_provenance_manifest(
        artifact_id="artifact-1",
        prepared=prepared,
        render_result=rendered,
        quality_report=quality,
        template_version_id="version-1",
        template_file_id="file-1",
        source_template_sha256=manifest.source_template_sha256,
        compiled_template_sha256=manifest.source_template_sha256,
        renderer={"name": "jinja_sandbox", "version": "test"},
        model={"provider": "none", "model": "deterministic"},
    )
    assert quality.passed
    assert provenance.output_sha256 == hashlib.sha256(rendered.data).hexdigest()
    assert provenance.context_snapshot_sha256 == prepared.context_snapshot_sha256
    assert provenance.slots[0].fact_refs == ["entity.legal_name"]
    assert provenance.facts[0].source_kind == "project_master"
    assert hashlib.sha256(provenance_bytes(provenance)).hexdigest() == provenance.manifest_sha256


def test_docx_quality_blocks_author_comments_and_external_relationships() -> None:
    template = _structured_docx()
    manifest = _ooxml_manifest(template, ".docx")
    rendered = _docx_with_external_relationship_and_comment(template)

    report = validate_rendered_artifact(
        _render_result(template, rendered, ".docx"),
        file_name="AnnualAudit.docx",
        manifest=manifest,
        template_bytes=template,
    )

    assert not report.passed
    codes = {issue.code for issue in report.issues}
    assert "AUTHOR_COMMENT_PRESENT" in codes
    assert "EXTERNAL_LINK_PRESENT" in codes
    with pytest.raises(AttachmentQualityError):
        require_quality(report)


def test_xlsx_quality_blocks_comments_external_links_and_hyperlink_formulas() -> None:
    template = _structured_xlsx()
    workbook = load_workbook(BytesIO(template), keep_links=False)
    worksheet = workbook["Data"]
    worksheet["A2"].comment = Comment("Remove before publication", "Reviewer")
    worksheet["A2"].hyperlink = "https://example.invalid/audit"
    worksheet["C2"] = '=HYPERLINK("file:///tmp/audit.xlsx","external")'
    rendered = _save_workbook(workbook)

    report = validate_rendered_artifact(
        _render_result(template, rendered, ".xlsx"),
        file_name="AnnualAudit.xlsx",
        manifest=_ooxml_manifest(template, ".xlsx"),
        template_bytes=template,
    )

    assert not report.passed
    codes = {issue.code for issue in report.issues}
    assert "AUTHOR_COMMENT_PRESENT" in codes
    assert "EXTERNAL_LINK_PRESENT" in codes
    assert not _ooxml_manifest(template, ".xlsx").allowed_modified_parts
    assert "ooxml_package_safety" in report.checks
    assert "template_structure" in report.checks
    with pytest.raises(AttachmentQualityError):
        require_quality(report)


def test_docx_quality_preserves_section_header_footer_page_and_style_structure() -> None:
    template = _structured_docx()
    manifest = _ooxml_manifest(template, ".docx")
    unchanged = validate_rendered_artifact(
        _render_result(template, template, ".docx"),
        file_name="AnnualAudit.docx",
        manifest=manifest,
        template_bytes=template,
    )
    assert unchanged.passed

    stripped_document = Document()
    stripped_document.add_paragraph("Annual audit body")
    output = BytesIO()
    stripped_document.save(output)
    changed = validate_rendered_artifact(
        _render_result(template, output.getvalue(), ".docx"),
        file_name="AnnualAudit.docx",
        manifest=manifest,
        template_bytes=template,
    )

    assert not changed.passed
    codes = {issue.code for issue in changed.issues}
    assert "DOCX_SECTION_STRUCTURE_CHANGED" in codes
    assert "DOCX_STYLE_STRUCTURE_CHANGED" in codes


def test_docx_header_footer_gate_ignores_text_but_blocks_structural_loss() -> None:
    template = _structured_docx()
    manifest = _ooxml_manifest(template, ".docx")

    text_changed_document = Document(BytesIO(template))
    text_changed_document.sections[0].header.paragraphs[0].text = "Changed entity header"
    text_changed_output = BytesIO()
    text_changed_document.save(text_changed_output)
    text_changed = validate_rendered_artifact(
        _render_result(template, text_changed_output.getvalue(), ".docx"),
        file_name="AnnualAudit.docx",
        manifest=manifest,
        template_bytes=template,
    )
    assert text_changed.passed, text_changed.issues

    structure_changed_document = Document(BytesIO(template))
    header = structure_changed_document.sections[0].header
    header_table = header.tables[0]
    header_table._element.getparent().remove(header_table._element)
    structure_changed_output = BytesIO()
    structure_changed_document.save(structure_changed_output)
    structure_changed = validate_rendered_artifact(
        _render_result(template, structure_changed_output.getvalue(), ".docx"),
        file_name="AnnualAudit.docx",
        manifest=manifest,
        template_bytes=template,
    )

    assert not structure_changed.passed
    assert any(
        issue.code == "DOCX_HEADER_FOOTER_STRUCTURE_CHANGED"
        for issue in structure_changed.issues
    )


def test_xlsx_quality_preserves_stable_template_structure() -> None:
    template = _structured_xlsx()
    manifest = _ooxml_manifest(template, ".xlsx")
    unchanged = validate_rendered_artifact(
        _render_result(template, template, ".xlsx"),
        file_name="AnnualAudit.xlsx",
        manifest=manifest,
        template_bytes=template,
    )
    assert unchanged.passed

    workbook = load_workbook(BytesIO(template), keep_links=False)
    worksheet = workbook["Data"]
    worksheet.unmerge_cells("D1:E1")
    worksheet.print_area = "A1:B2"
    worksheet.page_setup.orientation = "portrait"
    worksheet.freeze_panes = "C3"
    worksheet.oddHeader.center.text = "Changed header"
    del worksheet.tables["AuditTable"]
    del worksheet.defined_names["LocalInput"]
    del workbook.defined_names["InputArea"]
    workbook.remove(workbook["Notes"])
    rendered = _save_workbook(workbook)

    changed = validate_rendered_artifact(
        _render_result(template, rendered, ".xlsx"),
        file_name="AnnualAudit.xlsx",
        manifest=manifest,
        template_bytes=template,
    )

    assert not changed.passed
    codes = {issue.code for issue in changed.issues}
    assert "XLSX_SHEET_STRUCTURE_CHANGED" in codes
    assert "XLSX_MERGE_STRUCTURE_CHANGED" in codes
    assert "XLSX_PRINT_SETTINGS_CHANGED" in codes
    assert "XLSX_PAGE_SETUP_CHANGED" in codes
    assert "XLSX_HEADER_FOOTER_CHANGED" in codes
    assert "XLSX_FREEZE_PANES_CHANGED" in codes
    assert "XLSX_TABLE_STRUCTURE_CHANGED" in codes
    assert "XLSX_DEFINED_NAME_STRUCTURE_CHANGED" in codes


def test_xlsx_quality_allows_table_row_extension_without_value_comparison() -> None:
    template = _structured_xlsx()
    workbook = load_workbook(BytesIO(template), keep_links=False)
    worksheet = workbook["Data"]
    worksheet.append(["Expenses", 50])
    worksheet.tables["AuditTable"].ref = "A1:B4"
    rendered = _save_workbook(workbook)

    report = validate_rendered_artifact(
        _render_result(template, rendered, ".xlsx"),
        file_name="AnnualAudit.xlsx",
        manifest=_ooxml_manifest(template, ".xlsx"),
        template_bytes=template,
    )

    assert report.passed, report.issues
