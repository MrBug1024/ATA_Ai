from __future__ import annotations

from copy import copy
from io import BytesIO
from pathlib import Path

import pytest
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_COLOR_INDEX
from openpyxl import Workbook, load_workbook
from pypdf import PdfReader, PdfWriter

from ai_hunter.annual_audit.file_attachment_service import (
    DEFAULT_ATTACHMENT_TYPES,
    _context,
    _content_type_for_file_name,
    _docx_heading_level,
    _docx_table_has_outer_borders,
    _docx_authoring_annotation_count,
    _extract_case_statement_values,
    _infer_workpaper_code,
    _output_name,
    _template_field_plan,
    _template_text_and_structure,
    _validate_template_format_fidelity,
    _validate_rendered_template_content,
    render_template_file,
)
from ai_hunter.annual_audit.import_service import TabularSheet


def _docx_bytes() -> bytes:
    document = Document()
    document.add_paragraph("一般企业模板")
    document.add_paragraph("二〇二五年度财务报表审计报告")
    document.add_paragraph("我们认为，后附的财务报表在所有重大方面按照企业会计准则的规定编制。")
    table = document.add_table(rows=1, cols=4)
    table.cell(0, 0).text = "应收账款"
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def test_body_style_with_a_level_suffix_is_not_misclassified_as_a_heading():
    document = Document()
    style = document.styles.add_style("测试正文 [2级]", WD_STYLE_TYPE.PARAGRAPH)
    paragraph = document.add_paragraph("这是必须保留的会计政策正文。", style=style)

    assert _docx_heading_level(paragraph) is None


def test_reference_docx_is_filled_and_formal_opinion_is_downgraded():
    context = _context(
        {
            "engagement": {
                "entity_name": "北京示例有限公司",
                "fiscal_year": 2023,
                "period_end": "2023-12-31",
            },
            "sales_receivables": {"receivables": {"total_balance": 1234.5}},
        },
        "当前为事实性审计草稿。",
        1,
        {"version_label": "年度审计模板 v1"},
        template_type="annual_report",
    )

    rendered, status = render_template_file(_docx_bytes(), "审计报告正文.docx", context)
    document = Document(BytesIO(rendered))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)

    assert status == "filled"
    assert "北京示例有限公司" in text
    assert "二〇二三年度" in text
    assert "未形成正式审计意见" in text
    # The generated file must remain the uploaded template's structure.  Do
    # not append a second AI-written report section after the template.
    assert "系统生成的审计结果摘要" not in text


def test_complete_case_notes_reference_replaces_masked_entity_and_fills_case_data():
    document = Document()
    document.add_paragraph("北京****有限公司")
    document.add_paragraph("2023年度会计报表附注")
    table = document.add_table(rows=2, cols=3)
    table.cell(0, 0).text = "项目"
    table.cell(0, 1).text = "期末余额"
    table.cell(0, 2).text = "期初余额"
    table.cell(1, 0).text = "应收账款"
    table.cell(1, 1).text = "案例原始期末值"
    table.cell(1, 2).text = "案例原始期初值"
    source = BytesIO()
    document.save(source)

    context = {
        "entity_name": "北京有限公司",
        "fiscal_year": 2023,
        "issue_year": 2024,
        "period_end": "2023年12月31日",
        "__template_type__": "notes",
        "__case_workpaper_replay__": True,
        "__table_values__": {"应收账款": 999999},
    }
    rendered, status = render_template_file(
        source.getvalue(), "案例26页附注.docx", context
    )
    output = Document(BytesIO(rendered))
    text = "\n".join(paragraph.text for paragraph in output.paragraphs)
    cells = [cell.text for row in output.tables[0].rows for cell in row.cells]

    assert status == "filled"
    assert "北京****有限公司" not in text
    assert "北京有限公司" in text
    assert "案例原始期末值" not in cells
    assert "999,999.00" in cells


def test_workpaper_code_is_inferred_from_case_template_names():
    assert _infer_workpaper_code("3100-2 销售与收款循环(控制测试).xls") == "3100-2"
    assert _infer_workpaper_code("C5-2应收帐款审定表.xlsx") == "C5-2"
    assert _infer_workpaper_code("管理建议书模板.docx") == ""


def test_output_names_are_readable_and_keep_the_source_extension():
    context = _context(
        {"engagement": {"entity_name": "北京有限公司", "fiscal_year": 2023}},
        "审计结果摘要。",
        1,
        {"version_label": "年度审计模板 v1"},
        template_type="annual_report",
    )

    assert _output_name(
        "annual_report",
        7,
        1,
        "审计报告正文.docx",
        context=context,
    ) == "北京有限公司2023年度审计报告.docx"
    assert _output_name(
        "financial_statements",
        7,
        1,
        "2、经审计的财务报表.xls",
        context=context,
    ) == "北京有限公司2023年度财务报表.xls"
    assert _output_name(
        "audit_workpaper",
        7,
        1,
        "4101 货币资金.xls",
        "4101",
        context=context,
        duplicate_count=2,
    ) == "北京有限公司2023年度审计工作底稿-4101 货币资金.xls"


def test_output_content_type_follows_the_actual_extension():
    assert _content_type_for_file_name("财务报表.xls") == "application/vnd.ms-excel"
    assert _content_type_for_file_name("财务报表.xlsx") == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    assert _content_type_for_file_name("审计报告.docx") == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def test_case_legacy_xls_templates_keep_the_original_xls_format():
    template_root = next(
        Path(__file__).parents[2].glob(
            "new_docs/**/10、函证模版/2有价证券询证函.xls"
        )
    )
    context = _context(
        {"engagement": {"entity_name": "北京示例有限公司", "fiscal_year": 2023}},
        "审计结果摘要。",
        1,
        {"version_label": "函证模板 v1"},
        template_type="confirmations",
    )

    rendered, status = render_template_file(
        template_root.read_bytes(), template_root.name, context
    )

    # BIFF8 must never be silently rebuilt as OOXML.  Excel automation may
    # fill and save the workbook, or a service account without Excel may keep
    # the original bytes; both paths preserve the source format.
    assert rendered[:2] != b"PK"
    assert status in {"filled", "copied_no_matching_placeholders", "copied_legacy_unmodified"}
    if status == "copied_legacy_unmodified":
        assert rendered == template_root.read_bytes()


def test_core_delivery_matches_the_three_case_outputs():
    assert DEFAULT_ATTACHMENT_TYPES == (
        "annual_report",
        "financial_statements",
        "notes",
    )


def test_template_plan_rejects_case_output_as_a_template_and_requires_material_mapping():
    plan = _template_field_plan(
        template_type="notes",
        file_name="case-notes-reference.docx",
        text="北京****有限公司\n2023年度会计报表附注\n应收账款",
        structure={"paragraph_count": 3, "table_count": 1},
        template={
            "version_label": "v1",
            "files": [{"remark": "案例产出物，不是空白模板"}],
        },
        context={
            "entity_name": "北京有限公司",
            "fiscal_year": 2023,
            "period_end": "2023年12月31日",
            "audit_result": "审计结果",
            "notes_disclosure": True,
        },
        material_index={
            "status": "ready",
            "file_name": "主底稿.xlsx",
            "source_sha256": "abc",
            "sheet_count": 296,
            "nonempty_row_count": 15342,
            "read_all_sheets": True,
            "labels": {"应收账款": [{"label": "应收账款", "values": [100.0]}]},
        },
    )
    assert plan["status"] == "blocked"
    assert plan["case_output_reference_detected"] is True
    assert any("案例产出物" in item for item in plan["blockers"])


def test_tokenless_templates_are_not_appended_with_a_generic_audit_report():
    context = _context(
        {
            "engagement": {
                "entity_name": "北京示例有限公司",
                "fiscal_year": 2023,
                "period_end": "2023-12-31",
            },
            "readiness": {"counts": {"account_balance_rows": 3}},
        },
        "本次审计结果正文。",
        1,
        {"version_label": "年度审计模板 v1"},
        template_type="annual_report",
    )

    document = Document()
    document.add_paragraph("客户模板正文")
    docx_source = BytesIO()
    document.save(docx_source)
    docx_rendered, docx_status = render_template_file(docx_source.getvalue(), "模板.docx", context)
    docx_text = "\n".join(paragraph.text for paragraph in Document(BytesIO(docx_rendered)).paragraphs)
    assert docx_status == "copied_no_matching_placeholders"
    assert "审计结果及证据复核说明" not in docx_text

    workbook = Workbook()
    workbook.active["A1"] = "客户报表模板"
    xlsx_source = BytesIO()
    workbook.save(xlsx_source)
    xlsx_rendered, xlsx_status = render_template_file(xlsx_source.getvalue(), "模板.xlsx", context)
    rendered_workbook = load_workbook(BytesIO(xlsx_rendered), read_only=True)
    assert xlsx_status == "copied_no_matching_placeholders"
    assert rendered_workbook.sheetnames == ["Sheet"]
    assert xlsx_rendered[:2] == b"PK"

    md_rendered, md_status = render_template_file("# 客户模板\n".encode("utf-8"), "模板.md", context)
    assert md_status == "copied_no_matching_placeholders"
    assert "审计结果及证据复核说明".encode("utf-8") not in md_rendered

    pdf_writer = PdfWriter()
    pdf_writer.add_blank_page(width=595, height=842)
    pdf_source = BytesIO()
    pdf_writer.write(pdf_source)
    pdf_rendered, pdf_status = render_template_file(pdf_source.getvalue(), "模板.pdf", context)
    assert pdf_status == "copied_no_matching_form_fields"
    assert pdf_rendered.startswith(b"%PDF")
    assert len(PdfReader(BytesIO(pdf_rendered)).pages) == 1


def test_explicit_result_token_receives_bounded_template_result_not_raw_report():
    context = _context(
        {"engagement": {"entity_name": "北京示例有限公司", "fiscal_year": 2023}},
        "完整对话报告正文\n" * 100,
        1,
        {"version_label": "年度审计模板 v1"},
        template_type="management_letter",
    )
    document = Document()
    document.add_paragraph("{{audit_result}}")
    source = BytesIO()
    document.save(source)

    rendered, status = render_template_file(source.getvalue(), "管理建议书.docx", context)
    text = "\n".join(paragraph.text for paragraph in Document(BytesIO(rendered)).paragraphs)
    assert status == "filled"
    assert "完整对话报告正文" not in text
    assert "当前已导入数据未触发" in text


def test_docx_marker_replacement_preserves_the_template_run_styles():
    document = Document()
    paragraph = document.add_paragraph()
    prefix = paragraph.add_run("被审计单位：")
    prefix.italic = True
    placeholder = paragraph.add_run("{{entity_name}}")
    placeholder.bold = True
    source = BytesIO()
    document.save(source)

    rendered, status = render_template_file(
        source.getvalue(),
        "通用报告模板.docx",
        {"entity_name": "北京有限公司", "__template_type__": "annual_report"},
    )
    output = Document(BytesIO(rendered))
    runs = output.paragraphs[0].runs

    assert status == "filled"
    assert runs[0].text == "被审计单位："
    assert runs[0].italic is True
    assert runs[1].text == "北京有限公司"
    assert runs[1].bold is True


def test_table_total_is_resolved_from_the_section_fact_not_a_label_substring():
    document = Document()
    document.add_heading("货币资金", level=2)
    table = document.add_table(rows=5, cols=3)
    for column, text in enumerate(("项目", "期末余额", "期初余额")):
        table.cell(0, column).text = text
    for row, label in enumerate(("库存现金", "银行存款", "其他货币资金", "合计"), start=1):
        table.cell(row, 0).text = label
    source = BytesIO()
    document.save(source)

    rendered, status = render_template_file(
        source.getvalue(),
        "一般企业附注.docx",
        {
            "__template_type__": "notes",
            "__case_workpaper_replay__": True,
            "__table_values__": {
                "库存现金": 26743.25,
                "库存现金_期初": 29744.35,
                "银行存款": 2494581.57,
                "银行存款_期初": 4800392.62,
                "货币资金": 2521324.82,
                "货币资金_期初": 4830136.97,
            },
        },
    )
    output = Document(BytesIO(rendered))
    rows = [[cell.text for cell in row.cells] for row in output.tables[0].rows]

    assert status == "filled"
    assert rows[1][1:] == ["26,743.25", "29,744.35"]
    assert rows[2][1:] == ["2,494,581.57", "4,800,392.62"]
    assert rows[3][1:] == ["", ""]
    assert rows[4][1:] == ["2,521,324.82", "4,830,136.97"]


def test_generic_notes_drop_an_optional_unmapped_disclosure_grid():
    document = Document()
    document.add_heading("财务报表项目注释", level=1)
    document.add_heading("货币资金", level=2)
    cash = document.add_table(rows=3, cols=3)
    for column, text in enumerate(("项目", "期末余额", "期初余额")):
        cash.cell(0, column).text = text
    cash.cell(1, 0).text = "库存现金"
    cash.cell(2, 0).text = "合计"
    document.add_heading("交易性金融资产", level=2)
    optional = document.add_table(rows=2, cols=3)
    for column, text in enumerate(("项目", "期末余额", "期初余额")):
        optional.cell(0, column).text = text
    optional.cell(1, 0).text = "以公允价值计量且其变动计入当期损益的金融资产"
    document.add_paragraph("预计将于XX期间确认收入XX元。")
    source = BytesIO()
    document.save(source)

    rendered, status = render_template_file(
        source.getvalue(),
        "一般企业附注.docx",
        {
            "__template_type__": "notes",
            "__case_workpaper_replay__": True,
            "__table_values__": {"库存现金": 26743.25, "货币资金": 26743.25},
        },
    )
    output = Document(BytesIO(rendered))
    text = "\n".join(
        [paragraph.text for paragraph in output.paragraphs]
        + [cell.text for table in output.tables for row in table.rows for cell in row.cells]
    )

    assert status == "filled"
    assert "交易性金融资产" not in text
    assert "XX" not in text
    assert "26,743.25" in text


def test_table_flow_columns_use_distinct_source_facts():
    document = Document()
    document.add_heading("长期待摊费用", level=2)
    table = document.add_table(rows=2, cols=6)
    for column, text in enumerate(("项目", "上年年末余额", "本期增加金额", "本期摊销金额", "其他减少金额", "期末余额")):
        table.cell(0, column).text = text
    table.cell(1, 0).text = "合计"
    source = BytesIO()
    document.save(source)

    rendered, status = render_template_file(
        source.getvalue(),
        "一般企业附注.docx",
        {
            "__template_type__": "notes",
            "__case_workpaper_replay__": True,
            "__table_values__": {
                "长期待摊费用": 62678.47,
                "长期待摊费用_期初": 32793.51,
                "长期待摊费用本期增加": 70283.02,
                "长期待摊费用摊销": 40398.06,
            },
        },
    )
    output = Document(BytesIO(rendered))
    row = [cell.text for cell in output.tables[0].rows[1].cells]

    assert status == "filled"
    assert row == ["合计", "32,793.51", "70,283.02", "40,398.06", "", "62,678.47"]


def test_source_workpaper_layouts_produce_detail_facts_without_a_template_rule():
    values = _extract_case_statement_values(
        [
            TabularSheet(
                name="C2-2交易性金融资产审定表",
                rows=[
                    ["索引号", "项目", "期初金额", "本期借方金额", "本期贷方金额", "期末未审数", "调整数", "审定数"],
                    [None, "5．其他", 5230000, 18170000, 18400000, "=C2+D2-E2", None, "=F2+G2"],
                ],
            ),
            TabularSheet(
                name="C29-2长期待摊费用审定表",
                rows=[
                    ["项目", None, "装修费及其他", "合计"],
                    ["期初余额", None, 32793.51, None],
                    ["本期增加额", None, 70283.02, None],
                    ["本期已摊销额", None, 40398.06, None],
                    ["期末未审数", None, 62678.47, None],
                ],
            ),
            TabularSheet(
                name="C8-1-2其他应收款审定表",
                rows=[
                    ["序号", "往来单位", None, "期初余额", None, "本期借方", "本期贷方", "期末未审数", None, "调整数", None, "审定数"],
                    [None, "简称", None, "原币", "记账本位币", "发生额", "发生额", "原币", "记账本位币", None, None, None],
                    ["合计", None, None, 0, 120377.01, 413706.22, 386078.23, 0, 148005, 0, None, 148005],
                ],
            ),
        ]
    )

    assert values["交易性金融资产|其他"]["opening"] == 5230000
    assert values["交易性金融资产|其他"]["current"] == 5000000
    assert values["长期待摊费用|装修费及其他"]["opening"] == 32793.51
    assert values["长期待摊费用|装修费及其他"]["current"] == 62678.47
    assert values["长期待摊费用|装修费及其他|本期增加"]["current"] == 70283.02
    assert values["长期待摊费用|装修费及其他|摊销"]["current"] == 40398.06
    assert values["其他应收款"]["opening"] == 120377.01
    assert values["其他应收款"]["current"] == 148005


def test_blank_template_detail_row_is_filled_from_a_source_backed_fact():
    document = Document()
    document.add_heading("财务报表项目注释", level=1)
    document.add_heading("长期待摊费用", level=2)
    table = document.add_table(rows=3, cols=6)
    for column, text in enumerate(("项目", "上年年末余额", "本期增加金额", "本期摊销金额", "其他减少金额", "期末余额")):
        table.cell(0, column).text = text
    table.cell(2, 0).text = "合计"
    source = BytesIO()
    document.save(source)

    rendered, status = render_template_file(
        source.getvalue(),
        "一般企业附注.docx",
        {
            "__template_type__": "notes",
            "__case_workpaper_replay__": True,
            "__table_values__": {
                "长期待摊费用": 62678.47,
                "长期待摊费用_期初": 32793.51,
                "长期待摊费用本期增加": 70283.02,
                "长期待摊费用摊销": 40398.06,
                "长期待摊费用|装修费及其他": 62678.47,
                "长期待摊费用|装修费及其他_期初": 32793.51,
                "长期待摊费用|装修费及其他|本期增加": 70283.02,
                "长期待摊费用|装修费及其他|摊销": 40398.06,
            },
        },
    )
    output = Document(BytesIO(rendered))
    row = [cell.text for cell in output.tables[0].rows[1].cells]

    assert status == "filled"
    assert row == ["装修费及其他", "32,793.51", "70,283.02", "40,398.06", "", "62,678.47"]


def test_note_subsection_uses_its_parent_account_fact_scope():
    document = Document()
    document.add_heading("财务报表项目注释", level=1)
    document.add_heading("应收账款", level=2)
    document.add_heading("应收账款按账龄披露", level=3)
    table = document.add_table(rows=3, cols=3)
    for column, text in enumerate(("账龄", "期末余额", "上年年末余额")):
        table.cell(0, column).text = text
    table.cell(2, 0).text = "合计"
    source = BytesIO()
    document.save(source)

    rendered, status = render_template_file(
        source.getvalue(),
        "一般企业附注.docx",
        {
            "__template_type__": "notes",
            "__case_workpaper_replay__": True,
            "__table_values__": {
                "应收账款": 24000,
                "应收账款|1年以内": 24000,
            },
        },
    )
    output = Document(BytesIO(rendered))
    detail = [cell.text for cell in output.tables[0].rows[1].cells]
    total = [cell.text for cell in output.tables[0].rows[2].cells]

    assert status == "filled"
    assert detail == ["1年以内", "24,000.00", ""]
    assert total == ["合计", "24,000.00", ""]


def test_an_unmapped_grid_cannot_show_a_total_without_detail_evidence():
    document = Document()
    document.add_heading("财务报表项目注释", level=1)
    document.add_heading("其他流动资产", level=2)
    table = document.add_table(rows=3, cols=3)
    for column, text in enumerate(("项目", "期末余额", "上年年末余额")):
        table.cell(0, column).text = text
    table.cell(1, 0).text = "债权投资"
    table.cell(2, 0).text = "合计"
    source = BytesIO()
    document.save(source)

    rendered, status = render_template_file(
        source.getvalue(),
        "一般企业附注.docx",
        {
            "__template_type__": "notes",
            "__case_workpaper_replay__": True,
            "__table_values__": {"其他流动资产": 187210.4},
        },
    )
    output = Document(BytesIO(rendered))

    assert status == "filled"
    assert not output.tables


@pytest.mark.parametrize(
    ("template_type", "file_name"),
    [
        ("annual_report", "客户审计报告版式A.docx"),
        ("financial_statements", "客户报表版式A.docx"),
        ("notes", "客户披露版式A.docx"),
        ("future_custom_business", "任意新增业务模板.docx"),
    ],
)
def test_docx_core_table_contract_is_attachment_and_filename_neutral(template_type: str, file_name: str):
    """All document outputs use the same discovered table/fact placement."""

    document = Document()
    document.add_paragraph("被审计单位：{{entity_name}}")
    document.add_heading("货币资金", level=2)
    table = document.add_table(rows=2, cols=3)
    for column, value in enumerate(("项目", "期末余额", "期初余额")):
        table.cell(0, column).text = value
    table.cell(1, 0).text = "货币资金"
    placeholder = table.cell(1, 1).paragraphs[0].add_run("0.00")
    placeholder.bold = True
    source = BytesIO()
    document.save(source)

    context = {
        "entity_name": "北京有限公司",
        "__template_type__": template_type,
        "__table_values__": {"货币资金": 2521324.82, "货币资金_期初": 4830136.97},
    }
    rendered, status = render_template_file(source.getvalue(), file_name, context)
    output = Document(BytesIO(rendered))
    fidelity = _validate_template_format_fidelity(
        source.getvalue(), rendered, source_file_name=file_name
    )

    assert status == "filled"
    assert output.tables[0].cell(1, 1).text == "2,521,324.82"
    assert output.tables[0].cell(1, 2).text == "4,830,136.97"
    assert output.tables[0].cell(1, 1).paragraphs[0].runs[0].bold is True
    assert context["__template_render_contract__"]["contract_version"] == "template-structure-fact-registry-v1"
    assert context["__template_render_contract__"]["renderer"] == "docx"
    assert context["__template_render_contract__"]["mapped_value_cell_count"] == 2
    assert fidelity["format_contract_passed"] is True


@pytest.mark.parametrize(
    ("template_type", "file_name"),
    [
        ("annual_report", "客户审计报告版式A.xlsx"),
        ("financial_statements", "客户报表版式A.xlsx"),
        ("notes", "客户披露版式A.xlsx"),
        ("future_custom_business", "任意新增业务模板.xlsx"),
    ],
)
def test_xlsx_core_table_contract_is_attachment_and_filename_neutral(template_type: str, file_name: str):
    """Workbook templates receive the same fact registry, never substring matching."""

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "资产负债表"
    sheet["A1"] = "{{entity_name}}"
    for column, value in enumerate(("项目", "期末余额", "期初余额"), start=1):
        sheet.cell(row=3, column=column, value=value)
    sheet["A4"] = "货币资金"
    sheet["B4"] = 0
    sheet["C4"] = 0
    sheet["B4"].number_format = "#,##0.00"
    bold_font = copy(sheet["B4"].font)
    bold_font.bold = True
    sheet["B4"].font = bold_font
    source = BytesIO()
    workbook.save(source)

    context = {
        "entity_name": "北京有限公司",
        "__template_type__": template_type,
        "__table_values__": {"货币资金": 2521324.82, "货币资金_期初": 4830136.97},
    }
    rendered, status = render_template_file(source.getvalue(), file_name, context)
    output = load_workbook(BytesIO(rendered), data_only=False)
    fidelity = _validate_template_format_fidelity(
        source.getvalue(), rendered, source_file_name=file_name
    )

    assert status == "filled"
    assert output["资产负债表"]["B4"].value == 2521324.82
    assert output["资产负债表"]["C4"].value == 4830136.97
    assert output["资产负债表"]["B4"].number_format == "#,##0.00"
    assert output["资产负债表"]["B4"].font.bold is True
    assert context["__template_render_contract__"]["contract_version"] == "template-structure-fact-registry-v1"
    assert context["__template_render_contract__"]["renderer"] == "xlsx"
    assert context["__template_render_contract__"]["mapped_value_cell_count"] == 2
    _, structure = _template_text_and_structure(source.getvalue(), file_name)
    assert structure["template_contract"]["data_table_count"] == 1
    assert fidelity["format_contract_passed"] is True


def test_generic_notes_remove_unsupported_highlighted_business_choices_and_authoring_ink():
    """Highlight is a template authoring cue, not client-facing note content."""

    document = Document()
    document.add_paragraph("一般企业附注")
    retained = document.add_paragraph().add_run("本公司采用人民币为记账本位币。")
    retained.font.highlight_color = WD_COLOR_INDEX.GRAY_25
    construction = document.add_paragraph().add_run("（1）建筑施工业")
    construction.font.highlight_color = WD_COLOR_INDEX.YELLOW
    construction_text = document.add_paragraph().add_run("本公司采用投入法按履约进度确认建造合同收入。")
    construction_text.font.highlight_color = WD_COLOR_INDEX.YELLOW
    trading = document.add_paragraph().add_run("（2）贸易业务")
    trading.font.highlight_color = WD_COLOR_INDEX.YELLOW
    source = BytesIO()
    document.save(source)

    context = {
        "entity_name": "北京有限公司",
        "__template_type__": "notes",
        "__case_workpaper_replay__": True,
        "__table_values__": {"货币资金": 2521324.82},
        "__case_material_index__": {
            "entity_profile": {
                "main_activity": "向客户提供质量认证审核",
                "business_scope": "贸易代理及相关许可经营范围",
            },
        },
    }
    rendered, status = render_template_file(source.getvalue(), "一般企业附注.docx", context)
    output = Document(BytesIO(rendered))
    text = "\n".join(paragraph.text for paragraph in output.paragraphs)
    content_validation = _validate_rendered_template_content(
        rendered,
        template_type="notes",
        source_file_name="一般企业附注.docx",
        render_context=context,
    )

    assert status == "filled"
    assert "建筑施工" not in text
    assert "贸易业务" not in text
    assert "人民币为记账本位币" in text
    assert _docx_authoring_annotation_count(output) == 0
    assert content_validation["authoring_annotation_count"] == 0


def test_note_tables_receive_direct_left_and_right_outer_borders():
    document = Document()
    for label in ("货币资金", "应收账款"):
        table = document.add_table(rows=2, cols=3)
        for column, text in enumerate(("项目", "期末余额", "期初余额")):
            table.cell(0, column).text = text
        table.cell(1, 0).text = label
        table.cell(1, 1).text = "模板示例值"
        table.cell(1, 2).text = "模板示例值"
    source = BytesIO()
    document.save(source)

    context = {
        "__template_type__": "notes",
        "__table_values__": {"货币资金": 100, "应收账款": 200},
    }
    rendered, status = render_template_file(source.getvalue(), "客户附注模板.docx", context)
    output = Document(BytesIO(rendered))

    assert status == "filled"
    assert len(output.tables) == 2
    assert all(all(_docx_table_has_outer_borders(table)) for table in output.tables)
    assert context["__notes_outer_border_table_count__"] == 2
    assert context["__notes_outer_border_side_count__"] == 4


def test_xlsx_clears_sample_amounts_and_calculates_scoped_total():
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "资产负债表"
    sheet.append(["项目", "期末余额", "期初余额"])
    sheet.append(["库存现金", 999, 888])
    sheet.append(["银行存款", 999, 888])
    sheet.append(["货币资金合计", 999, 888])
    sheet.append(["其他项目", 777, 666])
    source = BytesIO()
    workbook.save(source)

    context = {
        "__template_type__": "financial_statements",
        "__table_values__": {
            "库存现金": 120,
            "库存现金_期初": 100,
            "银行存款": 80,
            "银行存款_期初": 80,
        },
    }
    rendered, status = render_template_file(source.getvalue(), "客户报表模板.xlsx", context)
    output = load_workbook(BytesIO(rendered), data_only=False)
    sheet = output["资产负债表"]

    assert status == "filled"
    assert sheet["B2"].value == 120
    assert sheet["C2"].value == 100
    assert sheet["B3"].value == 80
    assert sheet["C3"].value == 80
    assert sheet["B4"].value == 200
    assert sheet["C4"].value == 180
    assert sheet["B5"].value is None
    assert sheet["C5"].value is None
    assert context["__xlsx_table_render_stats__"][0]["cleared_template_default_cells"] == 8


def test_package_enforced_render_requires_a_real_result_mapping():
    document = Document()
    document.add_paragraph("被审计单位：{{entity_name}}")
    source = BytesIO()
    document.save(source)

    context = {
        "__template_type__": "annual_report",
        "entity_name": "北京有限公司",
        "__enforce_result_mapping__": True,
    }
    _, status = render_template_file(source.getvalue(), "审计报告模板.docx", context)

    assert status == "metadata_filled_no_result_mapping"
    assert context["__template_render_contract__"]["audit_result_mapped"] is False


@pytest.mark.parametrize("template_type", ["annual_report", "financial_statements", "notes"])
def test_core_docx_templates_clear_authoring_highlight_without_rebuilding_layout(template_type: str):
    document = Document()
    highlighted = document.add_paragraph().add_run("[日期]")
    highlighted.font.highlight_color = WD_COLOR_INDEX.YELLOW
    source = BytesIO()
    document.save(source)
    context = {"entity_name": "北京有限公司", "__template_type__": template_type}

    rendered, status = render_template_file(source.getvalue(), "客户格式模板.docx", context)
    output = Document(BytesIO(rendered))
    validation = _validate_rendered_template_content(
        rendered,
        template_type=template_type,
        source_file_name="客户格式模板.docx",
        render_context=context,
    )

    assert status == "filled"
    assert output.paragraphs[0].text == "待补充日期"
    assert _docx_authoring_annotation_count(output) == 0
    assert validation["authoring_annotation_count"] == 0
