from __future__ import annotations

import hashlib
import zipfile
from io import BytesIO
from pathlib import Path

import pytest
from docx import Document
from openpyxl import Workbook, load_workbook
from openpyxl.utils import get_column_letter
from openpyxl.workbook.defined_name import DefinedName
from openpyxl.worksheet.table import Table, TableStyleInfo
from pypdf import PdfReader
from reportlab.pdfgen import canvas

from ai_hunter.annual_audit.attachments.content_schemas import (
    AttachmentContractError,
    BindingManifest,
    ResolvedDocumentPayload,
    ResolvedSlotPayload,
)
from ai_hunter.annual_audit.attachments.renderers.docx import render_docx
from ai_hunter.annual_audit.attachments.renderers.markdown import render_markdown
from ai_hunter.annual_audit.attachments.renderers.pdf import render_pdf
from ai_hunter.annual_audit.attachments.renderers import pdf as pdf_renderer
from ai_hunter.annual_audit.attachments.renderers.xlsx import render_xlsx


WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _manifest(extension: str, document_code: str, template: bytes, slots: list[dict]):
    return BindingManifest(
        document_code=document_code,
        source_template_sha256=hashlib.sha256(template).hexdigest(),
        slots=slots,
    )


def _payload(document_code: str, slots: list[ResolvedSlotPayload]):
    return ResolvedDocumentPayload(document_code=document_code, slots=slots)


def _docx_content_control_template() -> bytes:
    document = Document()
    document.add_paragraph("Fixed heading")
    output = BytesIO()
    document.save(output)
    source = BytesIO(output.getvalue())
    target = BytesIO()
    with zipfile.ZipFile(source, "r") as archive, zipfile.ZipFile(target, "w") as rendered:
        for info in archive.infolist():
            data = archive.read(info.filename)
            if info.filename == "word/document.xml":
                marker = b"<w:sectPr"
                control = (
                    f'<w:sdt xmlns:w="{WORD_NS}"><w:sdtPr>'
                    '<w:tag w:val="company_name"/></w:sdtPr><w:sdtContent>'
                    '<w:p><w:r><w:t>placeholder</w:t></w:r></w:p>'
                    '</w:sdtContent></w:sdt>'
                ).encode("utf-8")
                data = data.replace(marker, control + marker, 1)
            rendered.writestr(info, data)
    return target.getvalue()


def _xlsx_template() -> bytes:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Data"
    sheet["A1"] = "placeholder"
    workbook.defined_names.add(DefinedName("CompanyName", attr_text="'Data'!$A$1"))
    sheet.append([])
    sheet.append(["Name", "Amount", "Formula"])
    sheet.append(["prototype", 0, "=B4*2"])
    table = Table(displayName="StatementRows", ref="A3:C4")
    table.tableStyleInfo = TableStyleInfo(name="TableStyleMedium2", showRowStripes=True)
    sheet.add_table(table)
    output = BytesIO()
    workbook.save(output)
    workbook.close()
    return output.getvalue()


def _blank_pdf(*, with_form: bool = False) -> bytes:
    output = BytesIO()
    drawing = canvas.Canvas(output, pagesize=(300, 800))
    drawing.drawString(20, 770, "Fixed PDF heading")
    if with_form:
        drawing.acroForm.textfield(
            name="company_name",
            x=40,
            y=700,
            width=180,
            height=24,
            value="",
            forceBorder=True,
        )
    drawing.showPage()
    drawing.save()
    return output.getvalue()


def test_docx_content_control_renders_explicit_target_and_reopens() -> None:
    template = _docx_content_control_template()
    manifest = _manifest(
        ".docx",
        "audit_report",
        template,
        [
            {
                "slot_id": "company_name",
                "target": "docx:content-control:company_name",
                "source": "document.entity.legal_name",
                "value_type": "scalar",
                "required": True,
                "missing_policy": "block",
            }
        ],
    )
    payload = _payload(
        "audit_report",
        [
            ResolvedSlotPayload(
                slot_id="company_name",
                kind="scalar",
                value="Example Co",
                fact_refs=["entity.legal_name"],
            )
        ],
    )
    result = render_docx(template, payload, manifest)
    Document(BytesIO(result.data))
    with zipfile.ZipFile(BytesIO(result.data)) as archive:
        xml = archive.read("word/document.xml").decode("utf-8")
    assert "Example Co" in xml
    assert "placeholder" not in xml
    assert "word/document.xml" in result.modified_parts


def test_xlsx_named_range_table_numeric_formula_and_formula_injection() -> None:
    template = _xlsx_template()
    manifest = _manifest(
        ".xlsx",
        "financial_statements",
        template,
        [
            {
                "slot_id": "company_name",
                "target": "xlsx:named-range:CompanyName",
                "source": "document.entity.legal_name",
                "value_type": "scalar",
                "required": True,
                "missing_policy": "block",
            },
            {
                "slot_id": "rows",
                "target": "xlsx:table:StatementRows",
                "source": "document.financial_statements.rows",
                "value_type": "table_rows",
                "required": True,
                "missing_policy": "block",
                "overflow_policy": "extend_rows",
                "style_policy": "clone_prototype_row",
                "options": {
                    "column_map": {"Name": "name", "Amount": "amount", "Formula": "formula"},
                    "formula_columns": ["Formula"],
                },
            },
        ],
    )
    payload = _payload(
        "financial_statements",
        [
            ResolvedSlotPayload(
                slot_id="company_name",
                kind="scalar",
                value="Example Co",
                fact_refs=["entity.legal_name"],
            ),
            ResolvedSlotPayload(
                slot_id="rows",
                kind="table_rows",
                value=[
                    {"name": "cash", "amount": 12.5},
                    {"name": "=HYPERLINK(\"https://bad\")", "amount": 8},
                ],
                fact_refs=["financial.balance_sheet"],
            ),
        ],
    )
    result = render_xlsx(template, payload, manifest)
    workbook = load_workbook(BytesIO(result.data), data_only=False)
    sheet = workbook["Data"]
    assert sheet["A1"].value == "Example Co"
    assert sheet["B4"].value == 12.5
    assert sheet["A5"].value.startswith("'=")
    assert sheet["C4"].value == "=B4*2"
    assert sheet["C5"].value == "=B5*2"
    assert sheet.tables["StatementRows"].ref == "A3:C5"
    workbook.close()


def test_xlsx_table_extension_refuses_to_overwrite_following_content() -> None:
    source = load_workbook(BytesIO(_xlsx_template()), data_only=False)
    source["Data"]["A5"] = "审计汇总"
    output = BytesIO()
    source.save(output)
    source.close()
    template = output.getvalue()
    manifest = _manifest(
        ".xlsx",
        "financial_statements",
        template,
        [
            {
                "slot_id": "rows",
                "target": "xlsx:table:StatementRows",
                "source": "document.financial_statements.rows",
                "value_type": "table_rows",
                "required": True,
                "missing_policy": "block",
                "overflow_policy": "extend_rows",
                "style_policy": "clone_prototype_row",
                "options": {
                    "column_map": {
                        "Name": "name",
                        "Amount": "amount",
                        "Formula": "formula",
                    },
                    "formula_columns": ["Formula"],
                },
            }
        ],
    )
    payload = _payload(
        "financial_statements",
        [
            ResolvedSlotPayload(
                slot_id="rows",
                kind="table_rows",
                value=[
                    {"name": "cash", "amount": 12.5},
                    {"name": "receivables", "amount": 8},
                ],
                fact_refs=["financial.balance_sheet"],
            )
        ],
    )

    with pytest.raises(AttachmentContractError, match="existing content at A5"):
        render_xlsx(template, payload, manifest)


def test_markdown_uses_compiler_variable_target_and_strict_sandbox() -> None:
    template = (
        "# {{ title }}\n"
        "{% for row in rows %}{{ row['name'] }}: {{ row['amount'] }}\n{% endfor %}"
    ).encode()
    manifest = _manifest(
        ".md",
        "notes",
        template,
        [
            {
                "slot_id": "title",
                "target": "md:variable:title",
                "source": "document.report.title",
                "value_type": "scalar",
                "required": True,
                "missing_policy": "block",
            },
            {
                "slot_id": "rows",
                "target": "md:variable:rows",
                "source": "document.financial_statements.rows",
                "value_type": "table_rows",
                "required": True,
                "missing_policy": "block",
            },
        ],
    )
    payload = _payload(
        "notes",
        [
            ResolvedSlotPayload(slot_id="title", kind="scalar", value="Notes"),
            ResolvedSlotPayload(
                slot_id="rows",
                kind="table_rows",
                value=[{"name": "cash", "amount": 12.5}],
            ),
        ],
    )
    result = render_markdown(template, payload, manifest)
    assert result.data.decode() == "# Notes\ncash: 12.5\n"

    unsafe = b"{{ title.__class__ }}"
    unsafe_manifest = _manifest(
        ".md",
        "notes",
        unsafe,
        [
            {
                "slot_id": "title",
                "target": "md:variable:title",
                "source": "document.report.title",
                "value_type": "scalar",
                "required": True,
                "missing_policy": "block",
            }
        ],
    )
    with pytest.raises(AttachmentContractError):
        render_markdown(
            unsafe,
            _payload("notes", [ResolvedSlotPayload(slot_id="title", kind="scalar", value="x")]),
            unsafe_manifest,
        )


def test_pdf_overlay_contract_renders_and_reopens() -> None:
    template = _blank_pdf()
    manifest = _manifest(
        ".pdf",
        "confirmation",
        template,
        [
            {
                "slot_id": "company_name",
                "target": "pdf:overlay:1:40:700:200:30",
                "source": "document.entity.legal_name",
                "value_type": "scalar",
                "required": True,
                "missing_policy": "block",
                "overflow_policy": "shrink_font",
                "style_policy": "explicit",
                "options": {
                    "font_name": "Helvetica",
                    "font_size": 10,
                    "minimum_font_size": 6,
                    "max_lines": 2,
                },
            }
        ],
    )
    payload = _payload(
        "confirmation",
        [ResolvedSlotPayload(slot_id="company_name", kind="scalar", value="Example Co")],
    )
    result = render_pdf(template, payload, manifest)
    reader = PdfReader(BytesIO(result.data), strict=True)
    assert len(reader.pages) == 1
    assert "Example Co" in (reader.pages[0].extract_text() or "")


def test_pdf_overlay_blocks_non_ascii_text_without_approved_cjk_font() -> None:
    template = _blank_pdf()
    manifest = _manifest(
        ".pdf",
        "confirmation",
        template,
        [
            {
                "slot_id": "company_name",
                "target": "pdf:overlay:1:40:700:200:30",
                "source": "document.entity.legal_name",
                "value_type": "scalar",
                "required": True,
                "missing_policy": "block",
                "overflow_policy": "shrink_font",
                "style_policy": "explicit",
                "options": {"font_name": "Helvetica", "font_size": 10, "max_lines": 2},
            }
        ],
    )
    payload = _payload(
        "confirmation",
        [ResolvedSlotPayload(slot_id="company_name", kind="scalar", value="示例公司")],
    )

    with pytest.raises(AttachmentContractError, match="approved font"):
        render_pdf(template, payload, manifest)


def test_pdf_renderer_registers_approved_cjk_font_from_fixed_root(monkeypatch) -> None:
    import reportlab.pdfbase.ttfonts as ttfonts

    font_root = Path(__file__).resolve().parents[1]
    font_path = font_root / "pyproject.toml"
    registered: dict[str, object] = {}

    class Metrics:
        @staticmethod
        def getFont(name: str):
            if name not in registered:
                raise KeyError(name)
            return registered[name]

        @staticmethod
        def registerFont(font: object) -> None:
            registered[pdf_renderer.PDF_CJK_FONT_NAME] = font

    captured: dict[str, object] = {}

    def fake_ttfont(name: str, filename: str, **kwargs):
        captured.update(name=name, filename=filename, **kwargs)
        return object()

    monkeypatch.setenv("ATTACHMENT_PDF_FONT_ROOT", str(font_root))
    monkeypatch.setattr(pdf_renderer, "_PDF_CJK_FONT_RELATIVE_PATH", Path("pyproject.toml"))
    monkeypatch.setattr(ttfonts, "TTFont", fake_ttfont)

    pdf_renderer._ensure_contract_font(pdf_renderer.PDF_CJK_FONT_NAME, Metrics)

    assert captured == {
        "name": pdf_renderer.PDF_CJK_FONT_NAME,
        "filename": str(font_path.resolve()),
        "subfontIndex": 2,
    }


def test_pdf_acroform_is_filled_and_flattened() -> None:
    template = _blank_pdf(with_form=True)
    manifest = _manifest(
        ".pdf",
        "confirmation",
        template,
        [
            {
                "slot_id": "company_name",
                "target": "pdf:acroform:company_name",
                "source": "document.entity.legal_name",
                "value_type": "scalar",
                "required": True,
                "missing_policy": "block",
                "overflow_policy": "shrink_font",
                "style_policy": "explicit",
                "options": {"font_name": "Helvetica", "font_size": 10, "max_lines": 1},
            }
        ],
    )
    payload = _payload(
        "confirmation",
        [ResolvedSlotPayload(slot_id="company_name", kind="scalar", value="Example Co")],
    )
    result = render_pdf(template, payload, manifest)
    reader = PdfReader(BytesIO(result.data), strict=True)
    assert not reader.get_fields()
    assert "Example Co" in (reader.pages[0].extract_text() or "")
